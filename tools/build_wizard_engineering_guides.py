from __future__ import annotations

import argparse
import json
import re
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_guides import (
    BLUE,
    GRAY,
    GUIDES_DIR,
    add_bullets,
    add_callout,
    add_cover,
    add_header_footer,
    add_license_appendix,
    add_paragraphs,
    add_steps,
    add_table,
    configure_page,
    finish_document,
    set_page_number_format,
    set_picture_alt_text,
    set_table_geometry,
    setup_styles,
)
from build_user_guide import add_screen, mark_first_rows_as_accessibility_headers


ROOT = Path(__file__).resolve().parents[1]
SCREEN_DIR = GUIDES_DIR / "User Guide Screenshots"
TOC_DEFAULTS = ROOT / "__work" / "wizard-engineering-toc-pages.json"


@dataclass
class PascalUnit:
    path: Path
    name: str
    subsystem: str
    line_count: int
    internal_dependencies: list[str]
    public_symbols: list[str]
    direct_consumers: list[str]
    test_consumers: list[str]


SUBSYSTEM_PURPOSE = {
    "components": "designer-visible language-manager components, lifecycle adapters, and language selectors",
    "core": "framework-neutral catalogs, persistence, locale data, terminology, workspace ownership, and pack construction",
    "design": "RAD Studio Tool Palette registration and design-time package exposure",
    "integration": "component-kit generation, controlled Delphi project integration, build, and deployment",
    "provider": "DeepL/Google configuration, credential protection, batching, placeholders, retries, and HTTPS translation",
    "review": "localization findings, layout proposals, text measurement, and application-owned string review",
    "runtime": "offline pack loading, preference, translation application, layout restoration, templates, and splash handling",
    "scan": "project discovery, DFM/FMX and Pascal text extraction, context, rules, quality, and stable-key catalog merging",
    "studio": "the FMX operator interface, Setup Wizard state machine, Localization Review, and Studio self-translation",
    "validation": "structural catalog checks that protect runtime export",
}

SUBSYSTEM_PLAIN_LANGUAGE = {
    "components": "This layer gives a Delphi application the nonvisual language manager and the visible language selector used at run time.",
    "core": "This layer defines the shared data and file rules used by the Studio, the scanners, and the deployed runtime.",
    "design": "This layer makes the language components appear in the RAD Studio Tool Palette and Object Inspector.",
    "integration": "This layer prepares the files a target Delphi project needs and copies approved outputs to approved destinations.",
    "provider": "This layer is the only part that talks to DeepL or Google and handles provider credentials and limits.",
    "review": "This layer turns translation and layout concerns into findings a developer can inspect and decide.",
    "runtime": "This layer loads local language packs and applies them inside the finished offline VCL or FMX application.",
    "scan": "This layer reads saved Delphi forms and Pascal source and turns eligible text into stable translation records.",
    "studio": "This layer is the developer-facing FMX application: its forms, buttons, pages, Wizard, and workflow state.",
    "validation": "This layer blocks structurally unsafe catalogs and packs without pretending to judge human language quality.",
}


UNIT_DESCRIPTIONS = {
    "DAT.Components.Core": "Defines the framework-neutral component contract shared by VCL and FMX managers: published configuration, language-change events, form identity mapping, error policy, and the process-wide translation functions used by dynamic text and HTML producers.",
    "DAT.Components.FMX.LanguageSelector": "Implements the FMX TDATFMXLanguageComboBox. It binds to one FMX language manager, lists admitted packs, optionally shows locale codes, and changes language without duplicating manager logic.",
    "DAT.Components.FMX": "Implements the production FMX language manager and its application lifecycle integration. It discovers forms, preserves state and designer geometry, applies packs, refreshes tabs, scroll boxes, browsers, grids, and dynamic content, and enforces reversible LTR/RTL behavior.",
    "DAT.Components.FMX.Spike": "Retains the focused FMX lifecycle spike used to prove form discovery and exactly-once application behavior before those contracts were promoted into the production FMX manager.",
    "DAT.Components.VCL.LanguageSelector": "Implements the VCL TDATVCLLanguageComboBox. It connects through a published LanguageManager property, refreshes available packs, and selects the locale while remaining fully designer-editable.",
    "DAT.Components.VCL": "Implements the production VCL language manager. It observes application idle, modal, and active-form transitions; translates open and later forms; preserves live control state; restores source geometry; and applies VCL-specific RTL and layout contracts.",
    "DAT.Components.VCL.Spike": "Retains the VCL lifecycle proof harness that established idle discovery, notification cleanup, and exactly-once form application before the production VCL component was finalized.",
    "DAT.Core.AITranslation": "Coordinates automatic translation at catalog level. It decides eligibility, invokes the provider client, preserves protected work, records origin/status, and produces review information instead of treating machine output as approved.",
    "DAT.Core.AtomicFile": "Provides atomic UTF-8 persistence with validation, temporary files, previous-version recovery, corruption quarantine, and controlled replacement. Catalogs, packs, preferences, settings, glossaries, and deployment data rely on it to avoid partial JSON writes.",
    "DAT.Core.BuildInfo": "Supplies the compiled Studio build stamp and display description used in the title, Wizard header, reports, and diagnostics so artifacts can be traced to the producing build.",
    "DAT.Core.CatalogJson": "Serializes and validates development catalogs and implements CSV import/export planning. It preserves stable keys, statuses, provenance, contexts, locale facts, and source checksums across round trips.",
    "DAT.Core.Diagnostics": "Centralizes structured information, warning, and error diagnostics so scanners, providers, integration, and runtime code report consistent severity and actionable context.",
    "DAT.Core.Glossary": "Owns the project glossary model, matching rules, suggestions, catalog application, validation, and atomic persistence. It keeps product terminology explicit and reusable across provider runs.",
    "DAT.Core.Hyphenation": "Loads language-aware hyphenation dictionaries and applies conservative word/text hyphenation. Runtime layout uses it only when ordinary wrapping and readable font fitting are insufficient.",
    "DAT.Core.Interchange": "Implements TMX translation-memory and TBX terminology interchange so reviewed material can move between the Studio and standards-aware localization tools.",
    "DAT.Core.LocaleFacts": "Reads Windows/Delphi locale facts such as native name, direction, date/time formats, decimal and thousands separators, and currency data for catalog metadata and runtime formatting.",
    "DAT.Core.ProjectDetection": "Identifies a selected DPR/DPROJ, determines VCL or FMX, application ID, targets, forms, units, and output metadata, and rejects ambiguous or unsupported projects before scanning.",
    "DAT.Core.RuntimePack": "Builds the compact offline runtime JSON pack from a validated development catalog. It filters non-runtime records, carries locale/layout metadata, and writes checksum-compatible source and target packs.",
    "DAT.Core.SharedDictionary": "Maintains the shared vetted wording dictionary used to repair or standardize recurring interface terms without overwriting reviewed or approved project-specific work.",
    "DAT.Core.Terminology": "Resolves project glossary, shared dictionary, calendar terminology, protected tokens, and contextual wording in a deterministic priority order before or after provider translation.",
    "DAT.Core.TranslationMemory": "Stores and retrieves reviewed translations by exact, normalized, and similarity match. It preserves application and context provenance so later catalogs can reuse trusted wording without blind cross-key approval.",
    "DAT.Core.TranslationWorkspace": "Defines the authoritative per-application workspace below Local AppData and returns paths for development catalogs, runtime packs, glossaries, deployment settings, and recovery artifacts.",
    "DAT.Core.Types": "Defines the central domain model: framework and status enumerations, project and locale profiles, translation entries, catalogs, ownership roles, and string conversions used by every pipeline stage.",
    "DAT.Design.FMX.Register": "Registers the FMX language manager and selector on the DAT Localization Tool Palette page. It is compiled only into the FMX design-time package, never into a deployed application.",
    "DAT.Design.VCL.Register": "Registers the VCL language manager and selector on the DAT Localization Tool Palette page. It is the VCL design-time bridge between package installation and Object Inspector use.",
    "DAT.Integration.BuildDeploy": "Builds selected Delphi platform/configuration targets and deploys matching packs to detected or authorized application folders while reporting command, output, timeout, and destination results.",
    "DAT.Integration.ComponentPackage": "Generates the recommended non-mutating component integration kit: the complete framework-appropriate ComponentSource set, packs, manifest, README, deployment script, and completion report.",
    "DAT.Integration.DelphiSource": "Implements the explicitly advanced source-integration editor with preview, authorization, transaction backup, bounded changes, and restoration; it is not used by the recommended component workflow.",
    "DAT.Integration.MenuResource": "Plans and applies controlled language-menu resource changes for advanced integration, including safe detection and reversal of the generated menu block.",
    "DAT.Integration.Package": "Builds the broader integration plan and coordinates package/source actions, generated artifacts, reset behavior, and safety transactions across supported integration modes.",
    "DAT.Provider.Batching": "Groups entries by compatible context and splits them into bounded request batches so provider limits, latency, cancellation, and context quality remain predictable.",
    "DAT.Provider.CalendarTerms": "Supplies locale-specific calendar vocabulary and protects weekday/month meanings from ambiguous short-word provider output.",
    "DAT.Provider.Client": "Implements the bounded HTTPS client for DeepL API Free/Pro and Google Cloud Translation Basic v2, including request bodies, response parsing, contexts, cancellation, timeout, retries, and normalized errors.",
    "DAT.Provider.CredentialStore": "Stores, retrieves, and removes provider secrets as Windows Generic Credentials. It deliberately separates secrets from provider-settings JSON, catalogs, packs, source, logs, and exports.",
    "DAT.Provider.Glossary": "Converts vetted project terminology into provider-facing glossary behavior where supported and local pre/post-processing where the provider API does not offer a glossary feature.",
    "DAT.Provider.LanguageCodes": "Maps canonical locale codes to the distinct source/target language codes accepted by DeepL and Google, including provider-specific aliases and unsupported-locale checks.",
    "DAT.Provider.Placeholders": "Protects Delphi format placeholders, accelerators, markup, identifiers, and other immutable tokens before a provider call and verifies exact restoration afterward.",
    "DAT.Provider.Retry": "Defines bounded retry policy for transient provider and network failures, including backoff and retryable status classification without retrying permanent authentication errors indefinitely.",
    "DAT.Provider.Settings": "Owns non-secret provider configuration under Local AppData: selected provider, DeepL plan, remember policy, timeout, and batch size. API key text is intentionally excluded.",
    "DAT.Provider.Types": "Defines provider and DeepL-plan enumerations, display/string conversions, and the structured provider exception carrying status and retry context.",
    "DAT.Review.ApplicationStrings": "Classifies strings owned by the application rather than by static form resources, providing focused review of live status, generated messages, templates, and runtime text contracts.",
    "DAT.Review.CodeGeometry": "Extracts control geometry assignments from Pascal so layout review can distinguish designer-owned positions from deliberately code-positioned controls that must not be moved automatically.",
    "DAT.Review.Localization": "Runs linguistic and layout review, creates findings and conservative per-language layout proposals, records accept/reject decisions, and prepares the HTML/JSON review package.",
    "DAT.Review.TextMeasurement": "Defines the platform-neutral text-measurement interface and fitting calculations used to compare translated text against available control geometry.",
    "DAT.Review.TextMeasurement.FMX": "Provides FMX canvas-based text measurement so font, wrapping, and control-size proposals reflect FireMonkey rendering characteristics.",
    "DAT.Review.TextMeasurement.GDI": "Provides Windows GDI text measurement for VCL controls and report/header fitting using the actual font metrics available on Windows.",
    "DAT.Runtime.FMX": "Applies runtime packs to FMX object trees while preserving live state, snapshots, designer geometry, control order, styled controls, grids, menus, scrolling, wrapping, and language direction.",
    "DAT.Runtime.LanguagePack": "Parses and validates deployed JSON packs, indexes stable keys/source text/templates/layout rules, supports dynamic and HTML translation, and exposes locale and compatibility metadata to the runtime manager.",
    "DAT.Runtime.LayoutOverrides": "Persists developer-entered per-language control overrides separately from generated packs so hand-tuned positions, sizes, and font choices remain explicit and recoverable.",
    "DAT.Runtime.Manager": "Owns offline pack discovery, admission, checksum/source-pack compatibility, selected locale, formatting settings, translation lookup, preference fallback, and the active source/target runtime pair.",
    "DAT.Runtime.Preference": "Reads and atomically writes the selected language code in the configured Local AppData or executable-folder preference location, with safe fallback from stale or invalid values.",
    "DAT.Runtime.SplashTranslation": "Defines the framework-neutral splash translation contract and keys so early startup text can be translated before the main form and standard lifecycle hooks exist.",
    "DAT.Runtime.SplashTranslation.FMX": "Implements early FMX splash translation and coordinates the selected runtime before ordinary FMX form discovery begins.",
    "DAT.Runtime.SplashTranslation.VCL": "Implements early VCL splash translation while respecting VCL handle creation and startup ordering.",
    "DAT.Runtime.TemplateRewrite": "Provides the framework-neutral contract for translating templated or owner-drawn text whose final string is produced after normal property translation.",
    "DAT.Runtime.TemplateRewrite.FMX": "Applies FMX-specific template and caption rewrites after styles or control templates regenerate visible text.",
    "DAT.Runtime.TemplateRewrite.VCL": "Intercepts and reapplies VCL caption/template text that Windows messages or owner-drawn controls can recreate after the initial translation pass.",
    "DAT.Runtime.VCL": "Applies packs to VCL windows and controls while preserving handles, focus, selections, list/grid state, designer layout, status panels, menus, dialogs, wrapping, and reversible RTL state.",
    "DAT.Scan.CatalogMerge": "Merges a new scan into an existing catalog by stable key and source text, detects collisions, preserves valid translations, recovers semantic contracts, and marks removed records obsolete.",
    "DAT.Scan.Context": "Infers UI role, semantic concept, description, confidence, and surrounding context for each scanned string so provider output and human review are less ambiguous.",
    "DAT.Scan.DomainProfile": "Builds an application-specific vocabulary and sense profile from project text, helping distinguish overloaded terms such as Play, Close, Schedule, and application-specific nouns.",
    "DAT.Scan.FormText": "Parses text DFM/FMX resources, walks component/property structure, applies eligibility rules, repairs encodings, and emits stable form.component.property scan items.",
    "DAT.Scan.PascalResources": "Parses Pascal resourcestring declarations and approved runtime assignments while avoiding arbitrary code/data strings that cannot be safely localized.",
    "DAT.Scan.Project": "Orchestrates complete project scanning across detected form and source files, progress/cancellation, context analysis, quality checks, and result provenance.",
    "DAT.Scan.Quality": "Analyzes scan output for suspicious keys, duplicate occurrences, encoding problems, dynamic/runtime ownership, missing context, and other conditions requiring operator attention.",
    "DAT.Scan.Rules": "Centralizes the universal include/exclude/property rules used by both VCL and FMX scanning so a project is not governed by application-specific exceptions.",
    "DAT.Scan.TextCodec": "Loads Delphi text safely across BOM/code-page variants, repairs known mojibake, and decodes Delphi string expressions without executing source code.",
    "DAT.Scan.Types": "Defines scan items, diagnostics, kinds, progress/cancellation callbacks, source snapshots, counts, and the result object passed into catalog merge and UI reporting.",
    "DAT.Studio.LocalizationReview": "Implements the Localization Review FMX window: findings, glossary terms, translation suggestions, layout proposals, decision persistence, HTML package generation, and controlled return to Wizard processing.",
    "DAT.Studio.MainForm": "Implements the landing screen and seven-page Maintenance Studio. It coordinates project detection, scanning, catalog editing, provider translation, validation, export, integration, settings, keyboard defaults, and status reporting.",
    "DAT.Studio.SetupWizard": "Implements the eight-step Setup Wizard state machine, prerequisite validation, scan/translation/review/final-processing sequence, cancellation boundaries, safety backup, kit generation, deployment, and Finish behavior.",
    "DAT.Studio.Translation": "Self-localizes the Studio interface by loading its own packs, applying them to FMX forms, maintaining the interface-language menu, and keeping Studio language state separate from target-project catalogs.",
    "DAT.Validation.Catalog": "Performs blocking structural validation and non-blocking review warnings for locale metadata, required text, placeholders, accelerators, duplicate keys, source changes, runtime wiring, and export readiness.",
}


TEST_PURPOSE = {
    "ApplicationStringSmokeTests": "Verifies application-owned and dynamic string contracts that are not ordinary designer properties.",
    "AtomicPersistenceSmokeTests": "Exercises atomic save, previous-file recovery, corruption quarantine, and interrupted-write behavior.",
    "CalendarTermSmokeTests": "Checks locale-aware calendar terminology and ambiguous weekday/month handling.",
    "ContextBatchingSmokeTests": "Checks that provider batches preserve compatible contexts and obey configured size limits.",
    "ContextSmokeTests": "Checks semantic role, concept, description, and confidence inference for representative UI text.",
    "FormScanContracts": "Runs the formal DFM/FMX scan contract fixtures and compares the canonical expected results.",
    "FoundationSmokeTests": "Covers foundational types, JSON, locale, workspace, validation, and runtime-pack construction contracts.",
    "StudioFormSmokeTests": "Streams and creates the production Studio forms to detect missing components, invalid properties, and form-resource regressions.",
}


def read_pascal_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def pascal_program_name(path: Path, text: str) -> str:
    match = re.search(r"(?im)^\s*(?:unit|program|library|package)\s+([A-Za-z0-9_.]+)", text)
    return match.group(1) if match else path.stem


def interface_text(text: str) -> str:
    match = re.search(r"(?is)\binterface\b(.*?)\bimplementation\b", text)
    return match.group(1) if match else text


def internal_uses(text: str) -> list[str]:
    names: set[str] = set()
    for match in re.finditer(r"(?is)\buses\b(.*?);", text):
        clause = match.group(1)
        # Unit source paths follow ``in 'path\\Unit.pas'`` in DPR/DPK files.
        # Removing string literals before extracting identifiers prevents those
        # filenames (and backup suffixes) from being reported as extra units.
        clause = re.sub(r"'(?:''|[^'])*'", "", clause)
        clause = re.sub(r"\{.*?\}|\(\*.*?\*\)|//[^\r\n]*", " ", clause, flags=re.S)
        names.update(re.findall(r"\bDAT\.[A-Za-z0-9_]+(?:\.[A-Za-z0-9_]+)*", clause))
    return sorted(names)


def public_symbols(text: str) -> list[str]:
    public = interface_text(text)
    symbols: list[str] = []
    patterns = [
        r"(?im)^\s{0,4}(T[A-Za-z0-9_]+)\s*=\s*(?:class|record|interface|\()",
        r"(?im)^\s{0,2}(E[A-Za-z0-9_]+)\s*=\s*class",
        r"(?im)^\s{0,2}(?:class\s+)?(?:function|procedure)\s+([A-Za-z0-9_]+)",
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, public):
            name = match.group(1)
            if name not in symbols:
                symbols.append(name)
    return symbols


def collect_pascal_units() -> list[PascalUnit]:
    paths = sorted((ROOT / "source").rglob("*.pas"))
    raw: dict[str, tuple[Path, str, list[str], list[str]]] = {}
    for path in paths:
        text = read_pascal_text(path)
        name = pascal_program_name(path, text)
        raw[name] = (path, text, internal_uses(text), public_symbols(text))

    consumers: dict[str, list[str]] = defaultdict(list)
    test_consumers: dict[str, list[str]] = defaultdict(list)
    for consumer_name, (_, _, dependencies, _) in raw.items():
        for dependency in dependencies:
            if dependency != consumer_name:
                consumers[dependency].append(consumer_name)
    for test_path in sorted((ROOT / "tools" / "tests").rglob("*.dpr")):
        text = read_pascal_text(test_path)
        test_name = pascal_program_name(test_path, text)
        for dependency in internal_uses(text):
            test_consumers[dependency].append(test_name)

    result: list[PascalUnit] = []
    for name, (path, text, dependencies, symbols) in raw.items():
        relative = path.relative_to(ROOT)
        subsystem = relative.parts[1]
        result.append(PascalUnit(
            path=relative,
            name=name,
            subsystem=subsystem,
            line_count=len(text.splitlines()),
            internal_dependencies=[item for item in dependencies if item != name],
            public_symbols=symbols,
            direct_consumers=sorted(consumers.get(name, [])),
            test_consumers=sorted(test_consumers.get(name, [])),
        ))
    return sorted(result, key=lambda item: (item.subsystem, item.name.lower()))


def add_static_toc(
    document: Document,
    guide_title: str,
    contents: list[tuple[str, int]],
) -> None:
    toc_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(toc_section)
    set_page_number_format(toc_section, "lowerRoman", 1)
    add_header_footer(toc_section, guide_title, True)
    document.add_paragraph("Table of Contents", style="TOC Heading")
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for heading, page_number in contents:
        cells = table.add_row().cells
        left = cells[0].paragraphs[0]
        left.paragraph_format.space_after = Pt(1)
        left.paragraph_format.keep_together = True
        left_run = left.add_run(heading)
        left_run.font.name = "Aptos"
        left_run.font.size = Pt(9)
        left_run.font.color.rgb = RGBColor.from_string(BLUE)
        right = cells[1].paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right.paragraph_format.space_after = Pt(1)
        right_run = right.add_run(str(page_number))
        right_run.font.name = "Aptos"
        right_run.font.size = Pt(9)
        right_run.font.bold = True
        right_run.font.color.rgb = RGBColor.from_string(GRAY)
    if table.rows:
        properties = table.rows[0]._tr.get_or_add_trPr()
        header_marker = OxmlElement("w:tblHeader")
        header_marker.set(qn("w:val"), "false")
        properties.append(header_marker)
    set_table_geometry(table, [8200, 1160], indent=0)
    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(content_section)
    set_page_number_format(content_section, "decimal", 1)
    add_header_footer(content_section, guide_title, True)


def add_control_table(document: Document, rows: list[list[str]]) -> None:
    add_table(document, ["Control or area", "Purpose", "Required operator action"], rows)


def add_focus_screen(
    document: Document,
    file_name: str,
    caption: str,
    alt_text: str,
    crop: tuple[int, int, int, int],
    aspect_ratio: float,
) -> None:
    add_screen(
        document,
        file_name,
        caption,
        alt_text,
        crop=crop,
        aspect_ratio=aspect_ratio,
        width=5.4,
    )


def load_toc_pages(path: Path | None) -> dict[str, dict[str, int]]:
    candidate = path or TOC_DEFAULTS
    if not candidate.exists():
        return {}
    return json.loads(candidate.read_text(encoding="utf-8"))


def toc_entries(headings: list[str], pages: dict[str, int]) -> list[tuple[str, int]]:
    return [(heading, int(pages.get(heading, 1))) for heading in headings]


def add_key_value(document: Document, label: str, value: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    lead = paragraph.add_run(label + ": ")
    lead.bold = True
    lead.font.color.rgb = RGBColor.from_string(BLUE)
    paragraph.add_run(value)
    return paragraph


def add_unit_entry(document: Document, unit: PascalUnit) -> None:
    heading = document.add_paragraph(style="Heading 3")
    heading.paragraph_format.keep_with_next = True
    heading.add_run(unit.name)
    file_paragraph = add_key_value(document, "File", str(unit.path))
    file_paragraph.paragraph_format.keep_with_next = True
    add_key_value(document, "Where it fits", SUBSYSTEM_PLAIN_LANGUAGE[unit.subsystem])
    add_key_value(document, "What it is", UNIT_DESCRIPTIONS[unit.name])
    reason = (
        f"This unit exists to keep {SUBSYSTEM_PURPOSE[unit.subsystem]} in the "
        f"{unit.subsystem} layer instead of coupling that responsibility to unrelated UI or framework code."
    )
    add_key_value(document, "Why it is separate", reason)
    symbols = ", ".join(unit.public_symbols[:14]) or "No separate public type or routine; the unit supplies registrations or implementation bindings."
    if len(unit.public_symbols) > 14:
        symbols += f"; plus {len(unit.public_symbols) - 14} additional public declarations"
    add_key_value(document, "Public surface", symbols)
    dependencies = ", ".join(unit.internal_dependencies) or "No DAT unit dependency; this is a leaf/foundation unit."
    add_key_value(document, "Direct DAT dependencies", dependencies)
    consumers = ", ".join(unit.direct_consumers) or "No production DAT unit imports it directly; it is an executable/package entry point or optional adapter."
    add_key_value(document, "Direct production consumers", consumers)
    tests = ", ".join(unit.test_consumers) or "Coverage is indirect through subsystem and integration tests."
    add_key_value(document, "Named test consumers", tests)
    risk = (
        f"The file contains {unit.line_count:,} lines. Changes require the {unit.subsystem} contract tests, "
        "a clean package/application build, and a runtime regression pass for both VCL and FMX whenever shared behavior changes."
    )
    add_key_value(document, "Change and validation risk", risk)


def test_program_purpose(name: str) -> str:
    if name in TEST_PURPOSE:
        return TEST_PURPOSE[name]
    spaced = re.sub(r"(?<!^)([A-Z])", r" \1", name)
    return f"Exercises the {spaced.replace(' Smoke Tests', '').replace(' Tests', '')} contract as an isolated Delphi console or GUI regression program."


WIZARD_HEADINGS = [
    "1. Purpose, Audience, and Final Result",
    "2. Before Opening the Wizard",
    "3. Wizard Navigation, Keyboard, and Safety",
    "4. Step 1 - Welcome",
    "5. Step 2 - Delphi Project",
    "6. Step 3 - Deployment",
    "7. Step 4 - Languages",
    "8. Step 5 - Translation Service",
    "9. Step 6 - Scan Project",
    "10. Step 7 - Review and Authorize",
    "11. Localization Review During Final Processing",
    "12. Step 8 - Process and Finish",
    "13. Component Kit and dependencies Folder",
    "14. Delphi Compiler Search Path",
    "15. Files and Folders Created or Used",
    "16. Build and Deployment Procedure",
    "17. Complete Runtime Verification",
    "18. Updating, Resuming, and Adding Languages",
    "19. Failure Recovery and Troubleshooting",
    "20. Security and Privacy",
    "21. Printable Completion Checklist",
    "22. Quick Reference",
    "23. Appendix A - License Information",
]


def build_wizard_guide(toc_pages: dict[str, int]) -> Path:
    document = Document()
    setup_styles(document)
    title = "Delphi App Translation Studio - Setup Wizard Guide"
    add_cover(
        document,
        "Setup Wizard Guide",
        "A Detailed, Screen-by-Screen Path from Delphi Source to Offline Language Packs",
        last_changed="August 31, 2026",
    )
    add_static_toc(document, title, toc_entries(WIZARD_HEADINGS, toc_pages))

    document.add_heading(WIZARD_HEADINGS[0], level=1)
    add_paragraphs(document, [
        "Welcome to the Translation Setup Wizard. This guide walks beside you through all eight steps, explains what each choice means, and tells you what to look for before you move on. Use it for a first VCL or FireMonkey translation, for a later update after source changes, or when you want to add another language.",
        "You do not need to memorize the whole process. Follow the screens in order on your first run, and return to the individual step sections whenever you need a reminder. Each section includes the complete screen, a closer view of the important controls, and a plain-language explanation of what happens next.",
        "During a normal run, the Wizard reads your selected Delphi project, creates or merges its development catalog, translates only eligible unresolved material, opens Localization Review, validates the result, exports the canonical source and translated packs, prepares the component integration kit, and deploys packs to the folders you approved. It leaves the target Pascal, DFM, FMX, DPR, and DPROJ files unchanged.",
        "Your finished application remains offline. Only the Studio contacts DeepL or Google, and it does so from the developer computer while translation work is under way. The application you ship reads validated local JSON packs and never receives the provider API key.",
    ])
    add_callout(document, "Start with a safe copy.", "Keep a pristine backup and perform the first localization run on a separate test copy of the target project. This gives you an easy way back while you learn the workflow. The Wizard also creates its own timestamped safety ZIP before final processing.")
    add_table(document, ["Completion product", "What it is", "Where it is used"], [
        ["Development catalog", "The full editable key/source/translation/context/review model.", r"Studio workspace under %LOCALAPPDATA%."],
        ["Canonical source pack", "The validated English or other authored-source runtime pack.", r"Beside every target EXE under Localization\Languages."],
        ["Translated pack", "The validated compact target-locale runtime pack.", r"Loaded by the VCL/FMX runtime manager."],
        ["Localization Review package", "HTML review, findings, glossary, layout proposals, and decisions.", r"Studio export\localization-review."],
        ["Component integration kit", "Complete framework-appropriate runtime/component source, packs, manifest, README, deployment script, and report.", r"Studio export\component-integration."],
        ["Safety ZIP", "Timestamped pre-processing backup of the selected project.", "Reported in Wizard progress and completion report."],
    ])

    document.add_heading(WIZARD_HEADINGS[1], level=1)
    add_paragraphs(document, [
        "A few minutes of preparation makes the Wizard much easier to use. Complete the checks below once, then keep this section nearby as your setup checklist.",
    ])
    document.add_heading("2.1 Obtain and build the Studio", level=2)
    add_steps(document, [
        "Download or clone the complete repository from https://github.com/tmartindub/DelphiAppTranslationStudio. Do not download isolated PAS, DPK, BPL, or JSON files.",
        r"Extract it to a short writable path, for example C:\DelphiProjects\Delphi App Translation.",
        "Open DelphiAppTranslationStudio.dproj in RAD Studio 13 Florence and build Win32 Release for the simplest first run.",
        r"The verified RAD Studio environment is C:\Program Files (x86)\Embarcadero\Studio\37.0\bin\rsvars.bat and the Win32 compiler is dcc32.exe in the same folder.",
        r"Run bin\Win32\Release\DelphiAppTranslationStudio.exe from the repository copy you just built.",
    ])
    document.add_heading("2.2 Prepare the target application", level=2)
    add_steps(document, [
        "Create a pristine backup and a separate test copy. Build and briefly run the test copy before localization work begins.",
        "Build the DAT runtime and design packages supplied with the Studio. In RAD Studio choose Component > Install Packages > Add and install the matching Win32 Release design BPL.",
        "Open the target primary form in the Form Designer. Place one TDATVCLLanguageManager or TDATFMXLanguageManager and one matching language combo box.",
        "In Object Inspector set ApplicationId to the exact Delphi project name, LanguagesFolder to Localization\\Languages, SourceLanguage to the authored locale, and connect the combo box LanguageManager property.",
        "Add and position any visible Language label or menu item in the designer. Choose File > Save All so the first scan sees every intended static string.",
        "Close the target project before Wizard final processing. The Wizard verifies this explicitly at Review and Authorize.",
    ])
    add_callout(document, "Keep the setup visible in Delphi.", "The supplied manager and selector are normal designer components. Place and configure them in the Form Designer and Object Inspector so another developer can see and maintain the setup later. The Wizard does not inject hidden runtime UI construction or silently rewrite the form.")
    document.add_heading("2.3 Provider and network readiness", level=2)
    add_bullets(document, [
        "DeepL is the recommended default. Use a DeepL API Free or API Pro key, not merely a consumer Translator subscription.",
        "Google Cloud Translation Basic v2 is supported when the API is enabled, billing/quota are valid, and the key restrictions permit this developer computer.",
        "A stored key is kept in Windows Credential Manager. A session-only key disappears when the Studio closes.",
        "Test the provider connection before final processing. Do not treat a timeout, 401/403, or 429 result as a successful setup.",
    ])

    document.add_heading(WIZARD_HEADINGS[2], level=1)
    add_screen(document, "07-wizard-welcome.png", "Figure 3-1. Complete Wizard frame and navigation controls.", "Translation Setup Wizard Welcome screen with step rail, content panel, Back, Next, and Cancel controls")
    add_control_table(document, [
        ["Left step rail", "Shows the eight-step sequence and the current step.", "Select only a previously reached step; unreached steps remain disabled."],
        ["Next", "Validates the current step and advances.", "Press Enter when Next is the default button."],
        ["Back", "Returns to the previous reached step.", "Use before final processing; it is hidden or disabled when backward movement is unsafe or meaningless."],
        ["Cancel", "Closes an ordinary Wizard session or a stopped run.", "Press Esc before final processing. After successful completion use Finish instead."],
        ["Begin Final Processing", "Replaces Next on Review and Authorize.", "Select only after the project-closed and authorization checks are true."],
        ["Finish", "Closes a successful completed run.", "Press Enter after reading the completion status."],
        ["Footer status", "Reports the last completed action and next requirement.", "Read it whenever a button remains disabled."],
    ])
    add_callout(document, "You stay in control until processing begins.", "Before you select Begin Final Processing, Cancel leaves the target unchanged. Once processing starts, the Wizard briefly blocks unsafe navigation while the active operation finishes or stops safely. After a successful run, Back and Cancel disappear because Finish is the only action you need.")

    document.add_heading(WIZARD_HEADINGS[3], level=1)
    add_screen(document, "07-wizard-welcome.png", "Figure 4-1. Step 1 - Welcome.", "Welcome page explaining the Wizard and safety boundary")
    add_focus_screen(document, "07-wizard-welcome.png", "Figure 4-2. Welcome-page navigation buttons.", "Back, Next, and Cancel buttons on the Welcome page", (55700, 75700, 0, 0), 5.2)
    add_paragraphs(document, [
        "The Welcome page is a quiet starting point. Nothing has been scanned, translated, backed up, exported, or deployed yet, so you can read the overview without affecting a project.",
        "Read the safety notice, then press Enter or select Next. The notice explains that final processing creates a backup and Studio-owned artifacts while your target Pascal, form, DPR, and DPROJ files remain read-only.",
    ])
    add_control_table(document, [
        ["Welcome explanation", "Summarizes project identification, languages, provider verification, scanning, packs, and component integration.", "Confirm that this is the intended guided workflow."],
        ["Safety notice", "Defines the read-only target-source contract and required backup.", "Do not continue if the selected target copy is not disposable or protected."],
        ["Next", "Opens Delphi Project.", "Default action; Enter should activate it."],
    ])

    document.add_heading(WIZARD_HEADINGS[4], level=1)
    add_screen(document, "08-wizard-project.png", "Figure 5-1. Step 2 - Delphi Project.", "Delphi Project page with project path, Browse, detected framework, application ID, targets, forms, and source units")
    add_focus_screen(document, "08-wizard-project.png", "Figure 5-2. Project selection and detected identity.", "Project path, Browse button, and detected project metadata", (6200, 10900, 1700, 36300), 2.65)
    add_paragraphs(document, [
        "This is where the Wizard learns which application you want to translate. Select Browse and choose the project's DPROJ file when one is available. A DPROJ gives the Wizard the clearest framework, platform, configuration, source, form, and output information. Detection is read-only; it does not compile or run the target.",
        "Pause for a moment after detection and check the project name, framework, targets, forms, and source-unit count. ApplicationId is especially important: it is the stable identity shared by the catalog, runtime packs, manager, workspace, and saved language preference. It must match the manager's ApplicationId property exactly.",
    ])
    add_control_table(document, [
        ["Project path", "Holds the selected DPR/DPROJ.", "Verify the test copy path character by character."],
        ["Browse", "Opens the project-file picker and reruns detection.", "Choose the target DPROJ whenever available."],
        ["Application ID", "Displays the detected stable application identity.", "Copy it into the target manager ApplicationId property exactly."],
        ["Copy ID", "Copies ApplicationId to the clipboard.", "Use when configuring Object Inspector."],
        ["Framework / targets", "Reports VCL or FMX and Win32/Win64 availability.", "Stop if the framework is wrong or expected targets are missing."],
        ["Form resources / source units", "Reports project discovery breadth.", "A suspiciously low count normally means the wrong or incomplete project was selected."],
        ["Next", "Accepts detected identity and opens Deployment.", "Enabled only after a supported project is identified."],
    ])

    document.add_heading(WIZARD_HEADINGS[5], level=1)
    add_screen(document, "09-wizard-deployment.png", "Figure 6-1. Step 3 - Deployment.", "Deployment page listing detected build outputs and optional application folders")
    add_focus_screen(document, "09-wizard-deployment.png", "Figure 6-2. Optional destination controls.", "Deployment destination list with Add and Remove controls", (0, 14700, 6200, 15700), 2.05)
    add_paragraphs(document, [
        "Most developers can leave this page alone. The Wizard already detects normal Delphi build-output folders, so you only need to add a destination for a separate installed, portable, network, USB, or test-run copy of the application.",
        "When you do add one, choose the folder that contains the executable - not its Localization or Languages child folder. If a removable or network destination is unavailable later, the Wizard reports and skips it without invalidating the language pack.",
    ])
    add_control_table(document, [
        ["Detected outputs", "Shows build-output folders inferred from the project.", "Verify they correspond to the configurations you intend to run."],
        ["Destination list", "Stores additional authorized application folders.", "Leave empty when normal build outputs are sufficient."],
        ["Add Application Folder", "Adds one executable-containing folder.", "Use the full path, including drive letter."],
        ["Remove Selected", "Removes an incorrect or obsolete optional destination.", "Select the row first."],
        ["Deployment summary", "Explains what will be deployed and when.", "Read before continuing; no source file is modified."],
    ])

    document.add_heading(WIZARD_HEADINGS[6], level=1)
    add_screen(document, "10-wizard-languages.png", "Figure 7-1. Step 4 - Languages.", "Languages page with source and target locale selectors and locale facts")
    add_focus_screen(document, "10-wizard-languages.png", "Figure 7-2. Source and target locale selection.", "Source-language and target-language selectors", (0, 13400, 0, 37600), 3.10)
    add_paragraphs(document, [
        "Choose the language already written in the forms and source as the Source language, then choose the new language you want to create as the Target language. The locale code also controls the pack name, native display name, reading direction, and regional formatting facts shown on the page.",
        "Check the native name and direction before continuing. If you change the project or either language later, the Wizard clears the downstream scan state for this session. That protective reset keeps a catalog built for one application or locale from being exported as another.",
    ])
    add_control_table(document, [
        ["Source language", "Identifies the authored source locale.", "Normally English (United States) [en-US] for an English Delphi application."],
        ["Target language", "Selects the pack to create or update.", "Choose one locale for this Wizard run."],
        ["Native name", "Shows the end-user language label.", "Verify spelling and script."],
        ["Direction", "Reports LTR or RTL.", "Arabic, Hebrew, and Urdu must report RTL."],
        ["Locale facts", "Shows date/time, separators, currency, and related regional metadata.", "Review for the intended region rather than only the base language."],
    ])

    document.add_heading(WIZARD_HEADINGS[7], level=1)
    add_screen(document, "11-wizard-translation-service.png", "Figure 8-1. Step 5 - Translation Service.", "Translation Service page showing DeepL provider, plan, masked key, remember option, save, and connection test")
    add_focus_screen(document, "11-wizard-translation-service.png", "Figure 8-2. Provider, key, and connection controls.", "DeepL provider and plan selectors, masked API key, save, and test buttons", (0, 10900, 0, 14400), 2.03)
    add_paragraphs(document, [
        "DeepL is the recommended first choice and the first-run default. If you previously saved a provider and plan, the Studio restores those choices. For security, it never puts a stored secret back into the API-key box, so a blank box beside a saved-key message is normal.",
        "If this is a new or replacement key, paste it, choose whether to remember it securely, and select Save / Replace Key. Then select Test Connection. A successful test confirms that the key and endpoint work; it does not guarantee translation quality or enough remaining quota for a large catalog.",
    ])
    add_control_table(document, [
        ["Provider", "Selects DeepL or Google Cloud Translation.", "Use DeepL unless project/account requirements call for Google."],
        ["DeepL plan", "Selects API Free or API Pro endpoint.", "Match the key's actual account plan; ignored for Google."],
        ["API key", "Accepts a new or replacement secret and masks it.", "Leave blank when the saved-key status says a usable credential exists."],
        ["Remember securely", "Selects Windows Credential Manager instead of session-only memory.", "Use only on a trusted developer computer."],
        ["Save / Replace Key", "Stores provider settings and the entered secret according to the remember policy.", "Use after pasting or rotating a key."],
        ["Test Connection", "Performs one small provider request with timeout/cancellation.", "Continue only after a success result."],
        ["Credential status", "Reports saved/session availability without revealing the key.", "Use it to distinguish an empty edit from a missing credential."],
    ])
    add_callout(document, "Your key is not stored in the project.", r"Non-secret settings are in %LOCALAPPDATA%\DelphiAppTranslationStudio\provider-settings.json. A remembered key is stored as a Windows Generic Credential named DelphiAppTranslationStudio/Providers/DeepL or DelphiAppTranslationStudio/Providers/Google Cloud Translation. It is never copied into the target source, catalog, pack, kit, or guide.")

    document.add_heading(WIZARD_HEADINGS[8], level=1)
    add_screen(document, "12-wizard-scan.png", "Figure 9-1. Step 6 - Scan Project.", "Scan Project page showing unique catalog entries, raw observations, recovered semantic contracts, duplicate occurrences, and scan items")
    add_focus_screen(document, "12-wizard-scan.png", "Figure 9-2. Canonical scan-count explanation.", "Scan totals and representative stable-key rows", (0, 17300, 0, 15700), 2.26)
    add_paragraphs(document, [
        "Select Scan Project and let the Wizard inventory the text that belongs in the translation catalog. The Wizard and Maintenance Studio use the same canonical scanner, so the unique catalog-entry total should agree when both are looking at the same saved project snapshot.",
        "The supporting numbers explain how the total was formed. Raw observations are direct discoveries, recovered semantic contracts restore known dynamic application text, and equivalent duplicate occurrences are represented once. The scanner reads saved text DFM/FMX resources, Pascal resourcestrings, eligible runtime assignments, and explicitly authorized project resource manifests; it does not sweep up arbitrary identifiers, user data, database values, or every string literal in source code.",
    ])
    add_control_table(document, [
        ["Scan Project", "Runs canonical project discovery and scan.", "Use after every saved source/form text change."],
        ["Unique catalog entries", "The stable-key records that enter catalog merge.", "Use this as the Wizard/Maintenance comparison count."],
        ["Raw observations", "Direct scanner observations before recovery/collapse.", "Use to explain how the unique total was formed."],
        ["Recovered semantic contracts", "Known dynamic/application text restored through stable semantic rules.", "Review when runtime text coverage changes."],
        ["Duplicate occurrences collapsed", "Equivalent repeated observations represented once.", "A nonzero value is expected when identical ownership contracts repeat."],
        ["Scan list", "Shows stable key and source text.", "Inspect representative forms, components, properties, resourcestrings, and dynamic contracts."],
    ])

    document.add_heading(WIZARD_HEADINGS[9], level=1)
    add_screen(document, "13-wizard-review.png", "Figure 10-1. Step 7 - Review and Authorize content area.", "Review and Authorize page summarizing project, language, provider, scan, translation, backup, and authorization")
    add_paragraphs(document, [
        "This page gives you one last chance to confirm the whole job before anything time-consuming begins. Read the summary from top to bottom and make sure it names the exact test project, framework, source and target locales, provider, scan state, unresolved work, backup, and component workflow you intended.",
        "Close the target project in RAD Studio before you authorize the run. An open DPROJ or form can leave unsaved work in memory and make the selected source snapshot unreliable, even though the Wizard itself does not rewrite those files.",
    ])
    add_control_table(document, [
        ["Review memo", "Prints the complete proposed operation in plain text.", "Read the path, application ID, locale, provider, counts, output, and backup line by line."],
        ["Project closed confirmation", "Confirms RAD Studio no longer owns an in-memory target-project state.", "Close the project, then select it."],
        ["Safety backup", "Requires the pre-processing ZIP.", "Leave enabled; final processing will not start without it."],
        ["Authorization", "Records that the developer reviewed the operation.", "Select only after all summary values are correct."],
        ["Begin Final Processing", "Starts backup, translation, review, validation, export, kit generation, and deployment.", "This is not a generic Next button; it crosses the controlled execution boundary."],
    ])

    document.add_heading(WIZARD_HEADINGS[10], level=1)
    add_paragraphs(document, [
        "After provider translation, the Wizard opens Localization Review automatically. Take your time here: this is where you review machine output, glossary candidates, application-owned strings, layout findings, and language-specific proposals while all translated text is available.",
        "When your decisions are saved, close Localization Review normally. You are not abandoning the run. Control returns to the waiting Wizard, which continues with validation, export, kit generation, deployment, and the completion report.",
    ])
    add_table(document, ["Review area", "Engineering meaning", "Developer decision"], [
        ["Findings", "Information, warning, and high-risk localization issues.", "Resolve high-risk issues before treating the language as release-ready."],
        ["Project glossary", "Approved source/target terminology and matching rules.", "Add, edit, delete, and save only product-correct terms."],
        ["Suggestions", "Translation-memory or terminology proposals with confidence/context.", "Accept, approve high-confidence, or reject explicitly."],
        ["Layout proposals", "Conservative width/height/font/wrap/position changes scoped to locale and control.", "Accept only after inspecting current versus proposed geometry and ownership."],
        ["Code-positioned controls", "Controls whose geometry is intentionally assigned in Pascal.", "Leave them alone unless the application design is changed deliberately."],
        ["Review package", "HTML and JSON audit artifacts for later review.", "Generate/open it when a printable or browser-based audit trail is useful."],
        ["Close", "Returns accepted decisions to Wizard processing.", "Close only after saving the decisions that should enter the pack."],
    ])
    add_callout(document, "Layout decision rule.", "Use natural word wrapping first, then intelligent hyphenation for long unbroken words, then a modest readable font reduction only when necessary. Never shift neighboring columns or controls merely to make one translated heading fit.")

    document.add_heading(WIZARD_HEADINGS[11], level=1)
    add_screen(document, "14-wizard-processing.png", "Figure 12-1. Step 8 - Process and Finish status area.", "Processing and completion page with operation text and progress memo", aspect_ratio=1.51)
    add_paragraphs(document, [
        "You can follow the run in the progress memo. It records the safety backup, catalog save, provider work, return from review, glossary application, validation, source and target pack export, kit generation, deployment, and completion report in the order they occur.",
        "When the final line reports success, select Finish. If you see STOPPED instead, the Wizard has ended the operation safely. Read the last successful line and the reason that follows, keep the log and generated artifacts, correct that specific cause, and run the Wizard again. Do not treat partial output as a finished release.",
    ])
    add_control_table(document, [
        ["Progress memo", "Timestamped operation and diagnostic history.", "Read from the last line upward when a run stops."],
        ["Optional application folder", "Deploys packs to one new/temporary executable folder.", "Use only when the folder was not already configured in Deployment."],
        ["Rebuild before deploying", "Requests a selected build before optional deployment.", "Normally leave clear; select only when another compile is actually required."],
        ["Platform / configuration", "Selects a supported Win32/Win64 Debug/Release build.", "Do not select unsupported targets."],
        ["Deploy to Application Folders", "Deploys the just-built result and packs to configured destinations.", "Use when optional destination deployment is part of this run."],
        ["Finish", "Closes a successful completed Wizard.", "Default action after reviewing success."],
        ["Cancel after STOPPED", "Closes a stopped run after diagnostics are preserved.", "Use only when Finish is unavailable because processing did not complete."],
    ])

    document.add_heading(WIZARD_HEADINGS[12], level=1)
    add_paragraphs(document, [
        r"The Wizard places the generated kit under <Studio>\export\component-integration\<ApplicationId>. Inside it, ComponentSource contains the complete current unit set for the selected framework. The Studio deliberately does not create or fill a dependencies folder inside your target project without your direction.",
        r"For a portable, repeatable build, copy that complete source set into <Target Project>\dependencies\DelphiAppTranslation\source. This optional project-local folder keeps the target independent of the Studio repository path and lets Git record the exact runtime source that ships with the application.",
    ])
    add_steps(document, [
        r"Open <Studio>\export\component-integration\<ApplicationId>\ComponentSource.",
        r"Create <Target Project>\dependencies\DelphiAppTranslation\source if the project will vendor the runtime.",
        "Copy every PAS unit from ComponentSource into that source folder. Do not select only the five files that happened to be named in one compiler error.",
        "When the Studio/runtime changes, regenerate the kit and refresh the entire vendored set together so interfaces and implementations remain compatible.",
        "Commit the dependency source with the target test project when licensing and repository policy permit. Never copy provider credentials or development catalogs with it.",
    ])
    add_table(document, ["Framework scope", "Complete unit set"], [
        ["Shared by VCL and FMX", "DAT.Core.AtomicFile; DAT.Core.Diagnostics; DAT.Runtime.LanguagePack; DAT.Runtime.Preference; DAT.Runtime.Manager; DAT.Runtime.LayoutOverrides; DAT.Runtime.SplashTranslation; DAT.Runtime.TemplateRewrite; DAT.Components.Core"],
        ["VCL-specific additions", "DAT.Runtime.VCL; DAT.Runtime.SplashTranslation.VCL; DAT.Runtime.TemplateRewrite.VCL; DAT.Components.VCL; DAT.Components.VCL.LanguageSelector"],
        ["FMX-specific additions", "DAT.Runtime.FMX; DAT.Runtime.SplashTranslation.FMX; DAT.Runtime.TemplateRewrite.FMX; DAT.Components.FMX; DAT.Components.FMX.LanguageSelector"],
    ])

    document.add_heading(WIZARD_HEADINGS[13], level=1)
    add_paragraphs(document, [
        "The Wizard can pass ComponentSource to a build it starts itself, but it does not edit your DPROJ. To make ordinary builds from RAD Studio work every time, add one permanent Search path entry for the current ComponentSource or, preferably, the project-local dependency folder.",
    ])
    add_steps(document, [
        "Open the target project in RAD Studio and choose Project > Options.",
        "At the top choose All configurations and All platforms unless the project intentionally uses distinct paths.",
        "Open Building > Delphi Compiler and locate Search path.",
        r"Append .\dependencies\DelphiAppTranslation\source when using the recommended project-local folder. An absolute example is C:\DelphiProjects\MyApplication\dependencies\DelphiAppTranslation\source.",
        "Preserve every existing path and the inherited $(DCC_UnitSearchPath) value. Do not replace the field with only the DAT path.",
        "Save the option, clean, and build every configuration/platform that will ship.",
    ])
    add_callout(document, "Verified toolchain path.", r"RAD Studio environment: C:\Program Files (x86)\Embarcadero\Studio\37.0\bin\rsvars.bat. Win32 compiler: C:\Program Files (x86)\Embarcadero\Studio\37.0\bin\dcc32.exe.")

    document.add_heading(WIZARD_HEADINGS[14], level=1)
    add_table(document, ["Path or file", "Contents and reason"], [
        [r"%LOCALAPPDATA%\DelphiAppTranslationStudio\provider-settings.json", "Contains provider, plan, remember policy, timeout, and batch size, but never the API key. It restores non-secret Studio provider behavior."],
        [r"%LOCALAPPDATA%\DelphiAppTranslationStudio\language.ini", "Contains the Studio interface locale. It keeps the Studio UI language independent of target projects."],
        [r"%LOCALAPPDATA%\DelphiAppTranslationStudio\Workspaces\<ApplicationId>\Development", "Contains full development catalogs. It preserves stable keys, source, translations, status, context, and scan provenance."],
        [r"...\Languages", "Contains canonical source and translated runtime packs. It provides the workspace copy used for kit generation and deployment."],
        [r"...\Glossaries", "Contains project terminology JSON. It preserves approved product wording across runs."],
        [r"...\Deployment", "Contains additional application destinations. It remembers authorized pack-deployment folders."],
        [r"%LOCALAPPDATA%\<ApplicationId>\language.ini", "Contains the target application's selected locale unless PreferenceLocation is customized. It restores the end user's language at startup."],
        [r"<Studio>\export\localization-review\<ApplicationId>\<locale>", "Contains HTML/JSON review artifacts. It creates a durable linguistic and layout audit trail."],
        [r"<Studio>\export\component-integration\<ApplicationId>", "Contains ComponentSource, packs, README, manifest, script, and completion report. It supplies the complete target integration set without source injection."],
        [r"<Target EXE>\Localization\Languages", "Contains deployed runtime packs. This is the default path resolved by LanguagesFolder."],
    ])
    set_table_geometry(document.tables[-1], [3600, 5760])
    add_callout(document, "Atomic recovery files.", "A .tmp file is an interrupted candidate, .previous is the last known-good version, and .corrupt-<timestamp> is quarantined invalid content. Do not delete recovery files until the active JSON has been validated and backed up.")

    document.add_heading(WIZARD_HEADINGS[15], level=1)
    add_steps(document, [
        "Confirm the complete dependency source and Search path before compiling.",
        "Build Win32 Debug and Win32 Release. Build Win64 only when the target application will ship it and the target's dependencies support it.",
        r"Copy the canonical source pack and every translated pack to each executable's Localization\Languages folder, or use the Wizard deployment action.",
        "Run the executable from the output folder you just built, not an older copy elsewhere on disk.",
        "Repeat pack deployment after any catalog, glossary, layout-decision, or source-pack change.",
    ])
    add_table(document, ["Build", "Minimum deployment check", "Runtime check"], [
        ["Win32 Debug", "EXE plus source/target packs in its output tree.", "Detailed first-pass diagnostics and all language switching."],
        ["Win32 Release", "Independent Release output receives identical pack set.", "Release candidate behavior and startup preference."],
        ["Win64 Debug", "Only when the application supports Win64.", "Architecture-specific runtime and layout behavior."],
        ["Win64 Release", "Only when it will ship.", "Final Win64 release candidate."],
        ["Installed/portable folder", "Executable identity verified before copying packs.", "Run in place, including removable/network availability rules."],
    ])

    document.add_heading(WIZARD_HEADINGS[16], level=1)
    add_steps(document, [
        "Start in canonical English/source language and inspect every form, dialog, menu, tab, grid, report, HTML page, message, and dynamic status area.",
        "Switch to each translated LTR language. Confirm static controls, dynamic strings, templates, HTML/report content, accelerator keys, placeholders, and persisted data.",
        "Switch between two non-English languages without returning to English. No previous-language text or image may remain.",
        "For Arabic, Hebrew, and Urdu, confirm paragraph direction and control alignment are RTL while identifiers, file paths, numbers, WinAPI, FMX, and other protected LTR tokens remain readable.",
        "Switch back to English. The current screen must repaint immediately; no blank content panel or delayed click is acceptable.",
        "Close and restart in each representative locale. Confirm the preference loads quickly and an invalid/missing pack falls back safely to the canonical source pack.",
        "Inspect long headings and paragraphs at normal DPI and any supported scaling. Use natural space wrapping, intelligent hyphenation, and only modest font fitting; reject clipping, overlap, premature wrapping, or shifted neighboring columns.",
        "Repeat the complete matrix for each shipped platform/configuration and record any exception by stable key, form, component, locale, and build.",
    ])

    document.add_heading(WIZARD_HEADINGS[17], level=1)
    add_paragraphs(document, [
        "A later scan merges into the same application/locale catalog. Matching stable key and source text preserve existing work; changed source becomes attention-required; new entries are added; missing entries become obsolete rather than disappearing silently.",
        "The provider receives only eligible unresolved records. Reviewed, approved, excluded, obsolete, and unchanged translated records are not retranslated merely because the Wizard runs again.",
    ])
    add_steps(document, [
        "Make source/form changes in the test project and choose Save All.",
        "Run the Wizard with the same ApplicationId and target locale, then scan again.",
        "Review new, changed, unchanged, recovered, and obsolete counts before authorization.",
        "Translate unresolved work, complete Localization Review, validate, export, regenerate the kit, and redeploy.",
        "Refresh the complete vendored dependency set only when runtime/component source changed; do not replace it for ordinary catalog-only updates.",
        "For another language, select that target locale and repeat scan/translation/review/export. Deploy the canonical source pack and all translated packs together.",
    ])

    document.add_heading(WIZARD_HEADINGS[18], level=1)
    add_table(document, ["Symptom", "Most likely cause", "Required correction"], [
        ["Project is not recognized", "Wrong/partial DPR/DPROJ or unsupported metadata.", "Select the saved DPROJ from the complete test copy and verify it opens in RAD Studio."],
        ["Scan is zero or unexpectedly small", "Wrong project, unsaved forms, binary/unreadable resources, or incomplete source closure.", "Save All, verify text resources and detected counts, then rescan."],
        ["Wizard and Maintenance totals differ", "Different project/source snapshot, old catalog merge basis, or comparing unique entries with raw observations.", "Use the same project and compare the same unique-entry headline and supporting counts."],
        ["Connection test never completes", "Network/provider request did not honor timeout/cancellation or an old build is running.", "Run the current build, verify timeout/settings, provider endpoint, firewall, and account; do not kill processing before diagnostics."],
        ["401/403", "Invalid key, wrong DeepL plan, disabled Google API, billing, or restriction problem.", "Correct account/endpoint/key restrictions and retest."],
        ["429", "Quota or rate limit.", "Wait for the provider window, reduce work/batch pressure, or correct quota/billing."],
        ["Selector is empty", "Packs missing, invalid, wrong ApplicationId/framework/version, or wrong executable folder.", "Inspect the actual running EXE folder and validate source/target pack headers."],
        ["Language switch is slow", "Repeated disk discovery, unnecessary full-form work, synchronous provider/network logic, or lifecycle loop.", "Profile runtime switching; it must use admitted in-memory packs and bounded open-form application only."],
        ["Mixed languages remain", "Previous-language dynamic/template text was not restored before applying the new pack.", "Verify source snapshot/dynamic reverse mapping and language-change refresh contracts."],
        ["English switch shows a blank page", "Current page was cleared without immediate rebuild/repaint.", "Verify the language-changed handler rebuilds current dynamic/HTML content synchronously."],
        ["RTL text is left aligned", "A paragraph/control direction contract is absent or overridden.", "Verify pack direction, framework applicator, and per-control ownership; preserve protected LTR tokens."],
        ["Compiler cannot find DAT units", "ComponentSource/dependency folder is missing or Search path is wrong for the active configuration.", "Copy the full unit set and correct All configurations/All platforms Search path."],
        ["STOPPED in final log", "Controlled backup, translation, review, validation, export, kit, build, or deployment failure.", "Preserve the log/report, correct the specific last error, and rerun; do not treat partial output as complete."],
    ])

    document.add_heading(WIZARD_HEADINGS[19], level=1)
    add_bullets(document, [
        "Provider API keys never belong in PAS/DFM/FMX/DPR/DPROJ, JSON catalogs, runtime packs, kits, logs, screenshots, documentation, email, issue trackers, or Git.",
        "Remembered keys are Windows Generic Credentials for the current user; session-only keys exist only in process memory.",
        "Provider-settings JSON contains no secret. Review it before source distribution only to confirm it contains non-secret configuration.",
        "The deployed target is offline and should not contain provider client code or credentials.",
        "Atomic persistence and required safety ZIP protect recoverability, but Git remains the authoritative engineering history for Studio and target projects.",
        "A component kit contains source and language packs. Review catalog/pack licensing and proprietary text before publishing it.",
    ])

    document.add_heading(WIZARD_HEADINGS[20], level=1)
    add_bullets(document, [
        "[ ] Complete repository downloaded or cloned; Studio builds and starts.",
        "[ ] Pristine target backup exists; disposable test copy builds before localization.",
        "[ ] Correct VCL/FMX Win32 Release design BPL installed through Install Packages.",
        "[ ] One manager and one connected selector saved on the primary form.",
        "[ ] ApplicationId, LanguagesFolder, and SourceLanguage verified in Object Inspector.",
        "[ ] Target project saved and closed before final processing.",
        "[ ] DeepL/Google key stored with intended policy and Test Connection passed.",
        "[ ] Project, source locale, target locale, native name, direction, and locale facts verified.",
        "[ ] Scan headline and supporting raw/recovered/duplicate counts reviewed.",
        "[ ] Review and Authorize summary, safety ZIP, project-closed check, and authorization confirmed.",
        "[ ] Localization Review findings, glossary, suggestions, and layout decisions completed.",
        "[ ] Progress ends successfully; completion report and component kit exist.",
        "[ ] Full ComponentSource copied to optional project dependencies folder when vendoring.",
        "[ ] Delphi Search path includes the exact dependency source for all shipped targets.",
        "[ ] Canonical source pack and every target pack deployed beside each actual EXE.",
        "[ ] LTR, RTL, non-English-to-non-English, switch-back, restart, dynamic, HTML, and layout matrix passes.",
    ])

    document.add_page_break()
    document.add_heading(WIZARD_HEADINGS[21], level=1)
    add_table(document, ["Purpose", "Exact or relative path"], [
        ["Studio repository", r"C:\DelphiProjects\Delphi App Translation (example)"],
        ["RAD environment", r"C:\Program Files (x86)\Embarcadero\Studio\37.0\bin\rsvars.bat"],
        ["Win32 compiler", r"C:\Program Files (x86)\Embarcadero\Studio\37.0\bin\dcc32.exe"],
        ["Studio settings", r"%LOCALAPPDATA%\DelphiAppTranslationStudio"],
        ["Project workspace", r"%LOCALAPPDATA%\DelphiAppTranslationStudio\Workspaces\<ApplicationId>"],
        ["Generated kit", r"<Studio>\export\component-integration\<ApplicationId>"],
        ["Vendored dependencies", r"<Target>\dependencies\DelphiAppTranslation\source"],
        ["Search path entry", r".\dependencies\DelphiAppTranslation\source"],
        ["Runtime packs", r"<Target EXE>\Localization\Languages"],
        ["Target preference", r"%LOCALAPPDATA%\<ApplicationId>\language.ini"],
    ])

    add_license_appendix(document, WIZARD_HEADINGS[22])

    mark_first_rows_as_accessibility_headers(document)
    path = GUIDES_DIR / "Delphi App Translation Studio Setup Wizard Guide.docx"
    finish_document(document, path)
    return path


ENGINEERING_HEADINGS = [
    "1. Guide Purpose, Plain-Language Architecture, and Vocabulary",
    "2. Product Scope, Non-Negotiable Invariants, and Trust Boundaries",
    "3. Repository, Build, and Distribution Layout",
    "4. End-to-End Architecture and Ownership",
    "5. Core Domain Model and Catalog Lifecycle",
    "6. Workspace, Persistence, and Recovery",
    "7. Project Detection and Canonical Scanning",
    "8. Context, Terminology, Translation Memory, and Glossary",
    "9. Provider Pipeline and Credential Security",
    "10. Localization Review, Measurement, Layout, and Hyphenation",
    "11. Validation and Runtime Pack Export",
    "12. Offline Runtime Manager and Pack Admission",
    "13. VCL Runtime Application and Lifecycle",
    "14. FMX Runtime Application and Lifecycle",
    "15. Components, Design Packages, and Designer Contract",
    "16. ComponentSource and Target dependencies",
    "17. Integration, Build, Deployment, and Project Boundaries",
    "18. Studio Forms, State Machines, and Self-Localization",
    "19. Directionality, HTML, Dynamic Text, and Reversible Layout",
    "20. Diagnostics, Failure Containment, and Performance",
    "21. Security and Privacy Engineering",
    "22. Test Architecture and Release Gates",
    "23. Production Pascal Unit Reference",
    "24. Test, Spike, and Contract Pascal Program Reference",
    "25. Delphi Project and Package File Reference",
    "26. JSON Schema and Form Resource Reference",
    "27. Build, Verification, and Documentation Tool Reference",
    "28. Dependency Graph Reading Guide",
    "29. Safe Change Procedure",
    "30. Engineering Release Checklist",
    "31. Quick Path and Identifier Reference",
    "32. Appendix A - License Information",
]


def parse_dpk(path: Path) -> tuple[list[str], list[str]]:
    text = read_pascal_text(path)
    requires_match = re.search(r"(?is)\brequires\s+(.*?);", text)
    contains_match = re.search(r"(?is)\bcontains\s+(.*?);", text)
    def names(match: re.Match[str] | None) -> list[str]:
        if not match:
            return []
        return [item.strip().strip("'") for item in re.split(r",\s*", match.group(1)) if item.strip()]
    return names(requires_match), names(contains_match)


def schema_summary(path: Path) -> tuple[str, list[str], list[str]]:
    data = json.loads(path.read_text(encoding="utf-8-sig"))
    description = str(data.get("description") or data.get("title") or "JSON contract used by the Studio.")
    properties = sorted((data.get("properties") or {}).keys())
    required = [str(item) for item in data.get("required", [])]
    return description, properties, required


def add_test_program_entry(document: Document, path: Path) -> None:
    relative = path.relative_to(ROOT)
    text = read_pascal_text(path)
    name = pascal_program_name(path, text)
    heading = document.add_paragraph(style="Heading 3")
    heading.paragraph_format.keep_with_next = True
    heading.add_run(name)
    file_paragraph = add_key_value(document, "File", str(relative))
    file_paragraph.paragraph_format.keep_with_next = True
    add_key_value(document, "What it verifies", test_program_purpose(name))
    dependencies = ", ".join(internal_uses(text)) or "No DAT unit imported directly."
    add_key_value(document, "DAT units exercised directly", dependencies)
    add_key_value(document, "Why it is important", "This executable isolates one contract so a regression produces a focused failure rather than being hidden inside a full application run.")
    add_key_value(document, "Required use", "Compile and run it through the maintained verification scripts or documented build matrix whenever its named subsystem changes.")


def add_support_pascal_entry(document: Document, path: Path) -> None:
    relative = path.relative_to(ROOT)
    text = read_pascal_text(path)
    name = pascal_program_name(path, text)
    add_key_value(document, "File", str(relative))
    add_key_value(document, "Program or unit", name)
    add_key_value(document, "Purpose", "Support source or fixture used by lifecycle, design-streaming, scan, or layout contract tests; it is not shipped as an application runtime dependency.")
    add_key_value(document, "DAT dependencies", ", ".join(internal_uses(text)) or "None; this fixture supplies target behavior or source text for another test.")


def add_package_entry(document: Document, path: Path) -> None:
    relative = path.relative_to(ROOT)
    requires, contains = parse_dpk(path)
    heading = document.add_heading(path.name, level=3)
    heading.paragraph_format.keep_with_next = True
    file_paragraph = add_key_value(document, "File", str(relative))
    file_paragraph.paragraph_format.keep_with_next = True
    add_key_value(document, "Package kind", "Design-time" if "design" in relative.parts else "Runtime")
    add_key_value(document, "Requires", ", ".join(requires) or "No explicit package requirement parsed.")
    add_key_value(document, "Contains", ", ".join(contains) or "No contained unit parsed.")
    add_key_value(document, "Why it exists", "Separates framework-neutral, framework-runtime, and IDE-only registration code so deployed applications never depend on design-time units and RAD Studio can stream the published components correctly.")


def add_important_file(document: Document, path: str, what: str, why: str) -> None:
    heading = document.add_heading(Path(path).name, level=3)
    heading.paragraph_format.keep_with_next = True
    file_paragraph = add_key_value(document, "File", path)
    file_paragraph.paragraph_format.keep_with_next = True
    add_key_value(document, "What it is", what)
    add_key_value(document, "Why it matters", why)


ENGINEERING_CHAPTER_SUMMARIES = {
    1: ("What rules may never be bypassed?", "This chapter defines the safety boundary. The Studio may read selected source and create its own artifacts. The recommended workflow does not rewrite the target project's Pascal, form, DPR, or DPROJ files."),
    2: ("Where does each kind of file belong?", "This chapter maps repository folders to responsibilities. Use it before adding a unit or generated artifact so production, test, documentation, and output files do not become mixed."),
    3: ("How does text move from source to a running translation?", "This chapter follows the pipeline from project detection through scanning, translation, review, validation, export, deployment, and run-time application. Each stage has one owner and one output contract."),
    4: ("What information does one translation record carry?", "The catalog stores identity, source and target text, context, status, origin, and review facts. Status describes workflow state; origin records where wording came from. They are not interchangeable."),
    5: ("Where is development state stored, and how is it protected?", "Project workspaces live under Local AppData, outside the target project. Important JSON and preference files use validated atomic replacement so an interruption cannot silently replace a good file with a partial one."),
    6: ("What exactly does a scan do?", "Detection finds the project closure. The scanner reads saved form resources and approved Pascal text. It produces stable records and counts; it does not run the target application or translate arbitrary data."),
    7: ("How does the Studio choose the best wording it already knows?", "Context explains meaning. Terminology protects required words. Translation memory reuses reviewed work. A glossary records project-specific choices. These steps reduce provider ambiguity before new machine translation is requested."),
    8: ("What crosses the network?", "Only eligible source text and safe context go to the selected provider. Credentials stay in Windows Credential Manager or session memory. Placeholders are protected, requests are bounded, and malformed responses are rejected."),
    9: ("Who decides whether wording and layout are acceptable?", "Automated review finds likely problems and can propose conservative layout changes. The developer decides. Structural validation can prove that a pack is safe to load, but it cannot prove that a sentence is the best human translation."),
    10: ("When may a runtime pack be created?", "Only a structurally valid catalog can be exported. Source and target packs must agree on application identity, framework, locale relationship, schema, and checksum contract."),
    11: ("Why does a pack appear—or not appear—in the selector?", "The runtime admits a pack only after compatibility checks pass. The canonical source pack is the fallback anchor and is required for reliable switching back to the authored language."),
    12: ("What is special about VCL?", "The VCL manager follows Windows form and handle lifecycle. Its applicator restores the designer baseline, applies the selected pack, and preserves focus, selection, lists, grids, menus, and other live state."),
    13: ("What is special about FireMonkey?", "The FMX manager handles styled controls, tabs, scrolling, browsers, grids, and later repainting. Template repair is important because an FMX style can recreate visible text after the first assignment."),
    14: ("Why are there separate runtime and design packages?", "Runtime packages contain code an application may ship. Win32 design packages contain RAD Studio registration code. Keeping them separate prevents deployed applications from depending on IDE-only units."),
    15: ("Why must ComponentSource be copied as one set?", "The generated source closure contains matching shared and framework-specific units. Copying only files named by the first compiler error can combine incompatible interfaces from different versions."),
    16: ("What may integration change?", "The normal component workflow generates a kit and deploys packs without changing target source. Advanced source editing exists only behind explicit preview, authorization, backup, and restore boundaries."),
    17: ("Which form owns which workflow state?", "MainForm owns Maintenance Studio. SetupWizard owns the guided state machine. LocalizationReview owns findings and decisions. Their FMX resources remain the editable design authority."),
    18: ("How are direction, HTML, and changing text kept correct?", "Every switch begins from the source baseline. Dynamic text and HTML use explicit translation keys and refresh contracts. RTL is applied at the correct control or paragraph boundary while protected identifiers remain LTR."),
    19: ("How does the system fail without hanging or corrupting state?", "Network, build, scan, save, and language-switch operations have explicit bounds and recovery behavior. Errors must return the interface to a usable state and must not expose secrets."),
    20: ("What information is trusted?", "Selected source is parsed, provider responses are validated, files are written atomically, and packs are admitted before use. Secrets, arbitrary JSON, and unverified output folders are never trusted merely because they exist."),
    21: ("What proves a change is ready?", "Small focused test programs isolate individual contracts. Release scripts then combine those results with builds, real pilot applications, documentation checks, and a Git review."),
    22: ("How should the production unit reference be used?", "Find the unit by subsystem, read its plain-language role and purpose, then inspect its direct dependencies, consumers, tests, and change risk before editing it."),
    23: ("What are the test and fixture files for?", "Tests are intentionally separate programs so a failure names one contract. Fixtures provide controlled forms, source, and behavior; they are evidence, not production dependencies."),
    24: ("Which files start builds and packages?", "The DPR starts a program. The DPROJ supplies RAD Studio build metadata. DPK files define packages and their contained units. A change to one can alter what is built even when Pascal implementation code is unchanged."),
    25: ("Why do schemas and form resources need their own review?", "Schemas define serialized compatibility. FMX and DFM resources define designer-owned layout and published properties. Both are contracts, not incidental generated text."),
    26: ("Which tool should be run for a given engineering task?", "This chapter explains the maintained build, verification, distribution, and documentation scripts. Prefer these repeatable entry points to improvised one-off commands."),
    27: ("How do dependencies show change impact?", "A direct dependency is imported by a unit. A direct consumer imports that unit. Trace both directions, then consider package and generated-source closure before deciding that a change is local."),
    28: ("What is the safest order for making a change?", "Define the contract, trace ownership, make the smallest universal correction, run focused checks, run the full gate, test real VCL and FMX applications, update documentation, and review the final Git diff."),
    29: ("What must be true before release?", "The checklist converts architecture promises into evidence: builds, tests, pack compatibility, state preservation, direction, dynamic content, deployment, security, distribution, and documentation."),
    30: ("Where are the most-used engineering paths?", "This is a lookup page for the project, compiler, source layers, schemas, tests, workspaces, generated kits, deployed packs, and credential targets."),
}


def add_engineering_chapter(document: Document, index: int) -> None:
    document.add_heading(ENGINEERING_HEADINGS[index], level=1)
    question, explanation = ENGINEERING_CHAPTER_SUMMARIES[index]
    add_callout(document, question, explanation)


def build_engineering_guide(toc_pages: dict[str, int]) -> Path:
    units = collect_pascal_units()
    missing = sorted(unit.name for unit in units if unit.name not in UNIT_DESCRIPTIONS)
    if missing:
        raise RuntimeError("Missing curated unit descriptions: " + ", ".join(missing))

    document = Document()
    setup_styles(document)
    title = "Delphi App Translation Studio - Engineering Guide"
    add_cover(
        document,
        "Engineering Guide",
        "Architecture, Source Units, Dependencies, Formats, Runtime Contracts, Tests, and Release Controls",
    )
    add_static_toc(document, title, toc_entries(ENGINEERING_HEADINGS, toc_pages))

    document.add_heading(ENGINEERING_HEADINGS[0], level=1)
    add_paragraphs(document, [
        "This guide explains how Delphi App Translation Studio is built and why its parts are separated. It is for Delphi developers who maintain the Studio or the runtime components used by translated VCL and FireMonkey applications.",
        "Begin with this chapter even if you already know Delphi. It defines the product's vocabulary and gives one simple model of the whole system. Later chapters add implementation detail. Chapters 23 through 27 are reference chapters for individual source files, packages, schemas, forms, and tools.",
    ])
    document.add_heading("1.1 Engineering Boundary: Localization Is a Development Workflow", level=2)
    add_paragraphs(document, [
        "Localization happens while a Delphi source project is under active development. It is not post-processing applied to a finished EXE.",
        "Developers can rescan after source changes and inspect the result inside the real application. They can also make deliberate wording or layout adjustments when an application needs them.",
        "Automation has a defined boundary. Stable keys preserve identity. Catalog merge preserves unchanged work. Provider batching limits network requests. Structural validation protects exported packs. Baseline restoration makes language changes reversible.",
        "Those mechanisms cannot infer every sentence assembled from live data. They also cannot make every terminology, layout, or bidirectional-text decision for the application owner.",
        "Release review must therefore cover dynamic text, specialized terms, custom controls, third-party controls, protected identifiers, right-to-left text, mixed-direction paragraphs, clipping, wrapping, and repeated language changes.",
        "A small and simple project may need almost no manual adjustment. The architecture must still leave a clear path for developer review and correction.",
        "The Studio’s engineering objective is to perform the safe translation work automatically and as accurately as possible, preserve application behavior and source ownership, and provide precise evidence for anything that still requires developer judgment.",
    ])
    add_table(document, ["Term", "Definition", "Ownership consequence"], [
        ["ApplicationId", "Stable exact identity shared by selected Delphi project, workspace, catalog, pack, runtime manager, and preference.", "Never derive it from a translated caption or arbitrary folder alias."],
        ["Development catalog", "Full editable translation model containing stable keys, source/target text, status, origin, context, locale facts, checksums, scan provenance, and review data.", "Studio-owned; not deployed directly as the compact runtime contract."],
        ["Canonical source pack", "Validated runtime pack containing the authored source language.", "Required for switch-back, compatibility comparison, and safe fallback."],
        ["Runtime pack", "Compact validated JSON used by the offline target application.", "Must match ApplicationId, framework, schema/version, locale, checksum, and source-pack contract."],
        ["Stable key", "Deterministic identity such as form.component.property or unit.resourcestring.", "Preserves translation across scans when ownership and source remain compatible."],
        ["Semantic contract", "Explicit application-owned key for dynamic text, templates, HTML, or other content not represented by one designer property.", "Allows safe scanning/application without translating arbitrary live data."],
        ["Manager", "One nonvisual VCL or FMX component that owns runtime discovery, selected locale, lifecycle, events, errors, and pack application.", "Exactly one primary manager normally supervises the application."],
        ["Selector", "Designer-visible VCL/FMX combo box connected to a manager.", "Lists only admitted packs and delegates language switching to the manager."],
        ["Applicator", "Framework-specific walker that applies text, layout, direction, and restoration to a VCL or FMX object tree.", "Contains framework behavior; provider and catalog code remain framework-neutral."],
        ["Designer baseline", "Snapshot of source-language text, geometry, alignment, direction, and relevant control state.", "Restored before every language transition so changes are reversible and do not accumulate."],
        ["Localization Review", "Operator window and artifact package for linguistic findings, glossary, suggestions, layout proposals, and decisions.", "Human decisions remain separate from structural validation."],
        ["ComponentSource", "Complete current subset of runtime/component PAS units generated for one target framework.", "Must be copied as one compatible set; partial copying creates interface drift."],
        ["Workspace", r"Per-ApplicationId Studio state under %LOCALAPPDATA%\DelphiAppTranslationStudio\Workspaces.", "Keeps development data outside target source and executable folders."],
        ["Pack admission", "Runtime validation performed before a pack appears as selectable.", "Invalid, wrong-app, wrong-framework, wrong-version, or incompatible packs never become active."],
    ])

    document.add_heading("1.2 The system in plain language", level=2)
    add_paragraphs(document, [
        "The product has two separate halves. The Studio is the developer tool. It scans source, contacts a translation service, supports review, and creates files. The runtime is the code placed in the target Delphi application. It works offline and reads only validated local language packs.",
        "A development catalog is the Studio's working record. It contains much more than translated words: it also records identity, context, review state, source changes, and checksums. A runtime pack is the smaller file produced from that catalog for the finished application.",
        "The target project remains the design authority. Forms stay editable in the Delphi designer. The runtime manager changes language while the application runs, but it must preserve the form's original text, layout, control state, and direction so switching back is reliable.",
    ])
    add_table(document, ["Part", "Plain-language responsibility", "Must not do"], [
        ["Project detector and scanner", "Find the selected Delphi project and collect text that is safe to translate.", "Run the target source or guess that arbitrary data is interface text."],
        ["Catalog and workspace", "Keep the durable development record outside the target source tree.", "Replace a valid file with a partial or corrupt write."],
        ["Provider pipeline", "Send eligible unresolved text to DeepL or Google and validate the response.", "Store a provider key in source, packs, logs, or catalogs."],
        ["Localization Review", "Show wording and layout concerns that need developer judgment.", "Silently approve language quality or move developer-owned controls."],
        ["Pack builder", "Create one compatible source pack and the translated runtime packs.", "Export a structurally invalid or mismatched catalog."],
        ["Runtime manager and applicator", "Load admitted packs and update open or later forms without losing live state.", "Contact a provider, rescan source, or translate the previous translation."],
        ["Component and integration kit", "Give a VCL or FMX project one compatible set of source units and deployment files.", "Mix old and new dependency units or rewrite target source without authorization."],
    ])

    document.add_heading("1.3 Follow one caption through the system", level=2)
    add_steps(document, [
        "A caption is saved in a DFM or FMX form. The scanner reads the saved form and assigns a stable key such as form.component.property.",
        "Catalog merge compares that key and source text with the prior catalog. Unchanged reviewed work is preserved. New or changed text is marked for attention.",
        "The terminology and provider pipeline prepares eligible text, protects placeholders, sends a bounded request, and records the returned text as a machine draft.",
        "Localization Review lets the developer inspect wording and any layout finding. Structural validation then checks identity, required text, placeholders, and pack compatibility.",
        "The pack builder writes the canonical source pack and target-language pack. Deployment places compatible packs under the executable's Localization\\Languages folder.",
        "At run time, the manager restores the designer baseline and asks the VCL or FMX applicator to apply the selected pack. A later switch repeats the process from the baseline, not from the prior translation.",
    ])

    document.add_heading("1.4 How to read file and dependency entries", level=2)
    add_table(document, ["Reference label", "What it means"], [
        ["Public surface", "Types, classes, interfaces, constants, and routines that another unit can use."],
        ["Direct DAT dependency", "A DAT unit named directly in this unit's uses clause. It does not include dependencies used only through another unit."],
        ["Direct production consumer", "A production DAT unit that names this unit directly in its own uses clause."],
        ["Named test consumer", "A test program that imports the unit directly. Other tests may still exercise it indirectly."],
        ["ComponentSource closure", "The complete group of shared and framework-specific PAS files needed by a target project. It must be refreshed as one set."],
        ["Change and validation risk", "The minimum subsystem tests and builds required after a change. Shared behavior normally requires both VCL and FMX verification."],
    ])
    add_callout(document, "A practical reading order.", "To understand behavior, read Chapters 2 through 22 in order. To change one unit, first read its Chapter 23 entry, then the architecture chapter for its layer, then the tests and consumers named by that entry.")

    add_engineering_chapter(document, 1)
    add_bullets(document, [
        "Target Pascal, DFM, FMX, DPR, and DPROJ remain read-only in the recommended Wizard/component workflow.",
        "Forms and published component properties remain designer-authored and editable in Object Inspector; runtime-only UI construction is not the normal integration mechanism.",
        "The development Studio may use DeepL or Google online; the deployed VCL/FMX application is offline and contains no provider credential.",
        "A language transition is reversible: restore the designer/source baseline first, then apply the selected pack once. Never translate the previous translation.",
        "Live application data, identifiers, object names, file paths, API names, numeric values, selection/focus, and connection state are not language-pack text unless an explicit semantic contract says otherwise.",
        "VCL and FMX share catalogs, packs, provider, validation, and manager core but use separate applicators, lifecycle hooks, direction rules, text measurement, and design packages.",
        "Stable keys and source text preserve unchanged work. Provider calls are limited to eligible unresolved entries; reviewed and approved work is never silently overwritten.",
        "All persistence that can affect catalogs, packs, preferences, settings, glossaries, deployment, or layout decisions is atomic and recoverable.",
        "RTL changes paragraph/control direction and alignment while preserving protected LTR tokens and explicit application-owned layout.",
        "No project-specific V5/V6 or named-application exception belongs in universal scanner, runtime, or layout code.",
    ])
    add_callout(document, "Trust boundary.", "The Studio trusts only explicitly selected project source and validated provider responses; the runtime trusts only admitted local packs. Neither side treats arbitrary JSON, user data, provider text, or an unverified executable folder as safe input.")

    add_engineering_chapter(document, 2)
    add_table(document, ["Folder", "Engineering ownership", "What must not be mixed into it"], [
        ["source\\core", SUBSYSTEM_PURPOSE["core"], "FMX/VCL controls, HTTP UI, or target-specific exceptions."],
        ["source\\scan", SUBSYSTEM_PURPOSE["scan"], "Runtime mutation or provider secrets."],
        ["source\\provider", SUBSYSTEM_PURPOSE["provider"], "Target runtime or design-time registration."],
        ["source\\review", SUBSYSTEM_PURPOSE["review"], "Silent application of unapproved geometry."],
        ["source\\validation", SUBSYSTEM_PURPOSE["validation"], "Linguistic approval decisions."],
        ["source\\runtime", SUBSYSTEM_PURPOSE["runtime"], "Provider networking or Studio forms."],
        ["source\\components", SUBSYSTEM_PURPOSE["components"], "Studio workflow or catalog editing."],
        ["source\\design", SUBSYSTEM_PURPOSE["design"], "Code required by deployed applications."],
        ["source\\integration", SUBSYSTEM_PURPOSE["integration"], "Unbounded source rewriting."],
        ["source\\studio", SUBSYSTEM_PURPOSE["studio"], "Core algorithms that tests/runtime need without FMX UI."],
        ["packages", "Runtime and design package projects.", "Application-specific compiled BPL copies."],
        ["bin / dcu", "Platform/configuration output.", "Authoritative source or hand-edited files."],
        ["docs\\guides / docs\\pdf", "Editable DOCX guides and print PDFs.", "Runtime dependencies or credentials."],
        ["export", "Generated review, pack, kit, and source-distribution artifacts.", "Developer secrets or canonical source history."],
        ["tools / tools\\tests", "Build, verification, docs, contracts, test executables, and fixtures.", "Production code reachable only from tests."],
    ])
    add_paragraphs(document, [
        "DelphiAppTranslationStudio.dpr/dproj are the FMX Studio entry point and project metadata. Output is partitioned by platform/configuration. Package projects separate framework-neutral runtime, VCL runtime, FMX runtime, and IDE-only VCL/FMX registration.",
        "Source distributions must contain complete source/package/schema/tool/document sets and must exclude secrets, Local AppData workspaces, provider keys, transient test workspaces, DCUs, EXEs, and unreviewed proprietary catalogs unless explicitly intended.",
    ])

    add_engineering_chapter(document, 3)
    add_table(document, ["Stage", "Input", "Owner", "Output / next boundary"], [
        ["Detect", "Selected DPR/DPROJ", "DAT.Core.ProjectDetection", "Framework, ApplicationId, targets, forms, units, output metadata"],
        ["Scan", "Saved DFM/FMX/PAS and authorized resources", "DAT.Scan.*", "TProjectScanResult with items, diagnostics, snapshots, and counts"],
        ["Merge", "Scan result plus prior catalog", "DAT.Scan.CatalogMerge", "Preserved/new/changed/obsolete development catalog"],
        ["Resolve", "Entries, domain profile, glossary, memory, terminology", "DAT.Core.* and DAT.Provider.*", "Eligible provider inputs and deterministic repairs"],
        ["Translate", "Bounded text/context batches and credential", "DAT.Provider.Client / DAT.Core.AITranslation", "Machine drafts with status/origin and diagnostics"],
        ["Review", "Translated catalog and source geometry", "DAT.Review.* / LocalizationReview form", "Findings, glossary, suggestions, layout decisions, HTML/JSON package"],
        ["Validate", "Complete development catalog", "DAT.Validation.Catalog", "Blocking errors and review warnings"],
        ["Export", "Validated catalog", "DAT.Core.RuntimePack", "Canonical source pack and translated runtime pack"],
        ["Integrate", "Project profile and packs", "DAT.Integration.ComponentPackage", "ComponentSource, manifest, README, deployment script, completion report"],
        ["Admit", "Pack files beside target EXE", "DAT.Runtime.LanguagePack / Manager", "Selectable compatible language descriptors"],
        ["Apply", "Selected pack and open/later forms", "DAT.Runtime.VCL/FMX plus components", "Immediate reversible translated UI with preserved state"],
    ])
    add_callout(document, "Ownership rule.", "Each stage validates the contract it receives and emits a narrower explicit contract. Later stages must not reach backward into a UI form or provider response to reconstruct missing state.")

    add_engineering_chapter(document, 4)
    add_paragraphs(document, [
        "DAT.Core.Types is the vocabulary root. TProjectProfile describes target identity/framework/files/targets. TLocaleProfile describes locale facts. TTranslationEntry carries stable identity, source/target text, status, origin, context, ownership, checksums, and review facts. TTranslationCatalog owns project/locale metadata and its entries.",
        "Status is workflow state, not provenance. Origin records how target text arrived; Reviewed/Approved record human decisions. A machine result must not become Approved merely because structural validation passes.",
    ])
    add_table(document, ["Catalog event", "Required behavior", "Forbidden behavior"], [
        ["Same key, same source", "Preserve translation, origin, review, approval, and applicable decisions.", "Retranslate or reset state without explicit operator action."],
        ["Same key, changed source", "Retain prior target for reference and mark source-changed/attention-required.", "Treat old target as current approved wording."],
        ["New key", "Add with context, ownership, source snapshot, and unresolved status.", "Borrow another key's approval automatically."],
        ["Missing key", "Retain as obsolete for audit/recovery.", "Delete silently."],
        ["Equivalent duplicate", "Collapse occurrence while retaining provenance/count.", "Create conflicting duplicate stable keys."],
        ["Semantic recovery", "Recreate known application-owned dynamic contract deterministically.", "Scan arbitrary runtime values as translatable source."],
    ])

    add_engineering_chapter(document, 5)
    add_paragraphs(document, [
        r"TTranslationWorkspace roots project state at %LOCALAPPDATA%\DelphiAppTranslationStudio\Workspaces\<ApplicationId>. Development, Languages, Glossaries, and Deployment separate editable data, deployed-format packs, terminology, and destination settings.",
        "TAtomicTextFile writes a candidate, validates it, preserves a prior known-good version, and replaces the active file atomically. Readers can recover .previous, quarantine corrupt input, and avoid converting a process interruption into an empty or half-written catalog.",
    ])
    add_table(document, ["Artifact", "Writer", "Validation before commit", "Recovery behavior"], [
        ["Development catalog", "TCatalogJson", "Schema/domain validation and parse round trip.", "Previous catalog retained; corrupt input quarantined."],
        ["Runtime pack", "TRuntimePackBuilder", "Required metadata, checksums, locale, entries, layout contract.", "Invalid pack never replaces the prior admitted copy."],
        ["Provider settings", "TProviderSettings", "Known provider/plan, bounded timeout/batch, no secret field.", "Defaults and previous valid settings remain available."],
        ["Language preference", "TLanguagePreference", "Admitted locale code.", "Invalid preference falls back to source/system policy."],
        ["Glossary", "TProjectGlossary", "Term/source/target and match-rule validation.", "Last valid glossary can be restored."],
        ["Layout overrides", "TLayoutOverrides", "Application/locale/form/control identity and numeric bounds.", "Invalid overrides are ignored/quarantined rather than moving controls unpredictably."],
    ])

    add_engineering_chapter(document, 6)
    add_paragraphs(document, [
        "TProjectDetector establishes the source closure before scanning. TProjectScanner orchestrates form and Pascal scanners, rules, context, domain profile, quality analysis, progress/cancellation, source snapshots, and the final result.",
        "TTextFormScanner reads text DFM/FMX structure and eligible string properties. TPascalResourceStringScanner reads resourcestring declarations and approved runtime assignments. TScanRuleSet is the universal policy boundary; it must not contain named-project or version-specific exceptions.",
    ])
    add_table(document, ["Count", "Meaning", "Why it is separate"], [
        ["Raw observations", "Direct form/Pascal/resource observations.", "Shows scanner input volume before merge behavior."],
        ["Recovered semantic contracts", "Explicit dynamic/application-owned strings added by stable rules.", "Dynamic coverage is auditable rather than hidden in code."],
        ["Equivalent duplicates collapsed", "Repeated equivalent ownership represented once.", "Prevents inflated translation counts without losing occurrence evidence."],
        ["Unique catalog entries", "Final stable-key records after recovery/collapse.", "This is the canonical Wizard/Maintenance headline."],
    ])
    add_callout(document, "Scan snapshot invariant.", "Final processing must reject a source/form save that occurred after the accepted scan. Exporting a stale snapshot can omit new text even when every later stage succeeds.")

    add_engineering_chapter(document, 7)
    add_paragraphs(document, [
        "Context and terminology are layered because provider text alone is insufficient for short interface strings. TScanContextAnalyzer records UI role, semantic concept, description, confidence, and surrounding evidence. TDomainProfiler builds project vocabulary and senses. TProjectGlossary and TSharedDictionary carry approved local and shared terms. TTranslationMemory reuses reviewed prior segments with provenance.",
        "Resolution order must preserve the most specific trusted evidence. Project glossary outranks shared/general terminology; reviewed contextual memory outranks an unreviewed machine draft; protected tokens and placeholders are never rewritten as terminology.",
    ])
    add_table(document, ["Layer", "Strength", "Allowed effect"], [
        ["Protected token/placeholder", "Absolute structural constraint", "Preserve exactly and validate restoration."],
        ["Project glossary", "Application-specific approved terminology", "Supply or repair the target term without touching protected Reviewed/Approved exceptions."],
        ["Contextual translation memory", "Trusted prior segment in compatible context", "Suggest or reuse according to explicit confidence/status rules."],
        ["Shared dictionary/calendar terms", "Vetted cross-project UI vocabulary", "Repair common UI/calendar wording where locale support exists."],
        ["Provider context", "Advisory description supplied with DeepL-capable requests", "Improve disambiguation but never override structural validation."],
        ["Provider draft", "Unreviewed external result", "Populate eligible unresolved target text as machine origin."],
    ])

    add_engineering_chapter(document, 8)
    add_paragraphs(document, [
        "TProviderSettings stores only provider, DeepL plan, remember policy, timeout, and batch size. TProviderCredentialStore stores the selected secret in Windows Credential Manager or session memory. TTranslationProviderClient builds and posts bounded HTTPS requests and maps provider-specific errors into ETranslationProviderError.",
        "TContextBatching groups compatible contexts and limits request size. TPlaceholderProtection replaces immutable tokens before network transmission and requires exact recovery. TProviderRetry retries only transient failures with bounded backoff. TProviderLanguageCodes prevents sending a canonical locale code that a provider does not accept.",
    ])
    add_table(document, ["Failure class", "Examples", "Engineering response"], [
        ["Permanent authentication/configuration", "401, 403, wrong DeepL plan, disabled Google API.", "Fail promptly with actionable account/endpoint guidance; do not loop retries."],
        ["Quota/rate", "429.", "Honor bounded retry policy and surface quota context."],
        ["Transient service/network", "5xx, temporary transport failure.", "Retry with bounded backoff and cancellation."],
        ["Timeout", "Request exceeds configured seconds.", "Cancel/abort the operation and return UI control; never leave the Wizard uncloseable."],
        ["Cancellation", "Operator stops final processing or closes within allowed boundary.", "Raise controlled cancellation, preserve completed atomic artifacts, and report STOPPED safely."],
        ["Structural response error", "Wrong response count, malformed JSON, missing placeholders.", "Reject the batch; never write partial unvalidated translations as complete."],
    ])

    add_engineering_chapter(document, 9)
    add_paragraphs(document, [
        "TLocalizationReviewer combines source geometry, code-geometry ownership, translated measurements, text roles, and language behavior into findings and proposals. ITextMeasurer isolates metric acquisition; FMX and GDI implementations provide framework-appropriate measurements.",
        "Layout changes are stored as per-language rules and applied only after baseline restoration. A proposal is not application state until explicitly accepted. Code-positioned controls and complex layout relationships are excluded from automatic movement.",
    ])
    add_table(document, ["Fitting order", "Engineering rationale", "Stop condition"], [
        ["Natural space wrapping", "Preserves words, typography, and shared column geometry.", "Text fits without clipping or harming neighbors."],
        ["Intelligent hyphenation", "Breaks a long unspaced word at language-aware points.", "Readable result fits and remains linguistically acceptable."],
        ["Modest font reduction", "Uses remaining control capacity without expanding shared layout.", "Do not cross the readable minimum or create inconsistent hierarchy."],
        ["Conservative size/layout proposal", "Adds locale-specific room when ownership and neighbors permit.", "Reject if it moves code-owned controls or introduces collision."],
        ["Developer redesign", "Handles complex collision, grid, graph, or container constraints.", "Required when automated proposals cannot preserve the design."],
    ])

    add_engineering_chapter(document, 10)
    add_paragraphs(document, [
        "TCatalogValidator separates blocking runtime safety from linguistic state. Required metadata, duplicate keys, source changes, placeholder/accelerator integrity, runtime wiring, and eligible target text are validated; Reviewed and Approved remain explicit human states.",
        "TRuntimePackBuilder exports only after structural validation. It creates a compact pack with identity, framework, version/schema, source/target locale, checksum linkage, texts, templates, source maps needed for reversibility, and accepted layout rules.",
    ])
    add_bullets(document, [
        "An error blocks export.",
        "A warning requires developer judgment and may describe linguistic or runtime-readiness risk.",
        "Information explains a condition that is safe but relevant.",
        "Canonical source and every target pack in one deployment must share compatible application identity and catalog/source contract.",
    ])

    add_engineering_chapter(document, 11)
    add_paragraphs(document, [
        "TRuntimeLanguagePack parses one JSON file and exposes text, templates, source maps, locale, and layout rules. TTranslationRuntime discovers files, validates pack headers/checksums/framework/application identity, admits compatible packs, loads source and selected target, updates format settings, and serves translation lookups.",
        "TLanguagePreference remembers an admitted locale. Startup loads the source pack as the safety anchor, then attempts the stored or system language according to policy. A stale preference cannot force an invalid pack into the process.",
    ])
    add_table(document, ["Admission check", "Reason"], [
        ["ApplicationId", "Prevents one product's pack from appearing in another."],
        ["Framework", "Prevents incompatible layout/runtime assumptions between VCL and FMX."],
        ["Schema/version", "Prevents a runtime from misreading a newer/older contract."],
        ["Locale/source locale", "Prevents mislabeled files and invalid fallback relationships."],
        ["Checksum/source compatibility", "Prevents a translated pack from pairing with a different source catalog."],
        ["Required data/layout structure", "Prevents incomplete pack activation and partial translation."],
    ])

    add_engineering_chapter(document, 12)
    add_paragraphs(document, [
        "TDATVCLLanguageManager connects TTranslationRuntime to VCL application lifecycle. Idle inspection, modal boundaries, active-form change, notifications, and explicit ApplyToForm cover existing and later forms without replacing application event handlers destructively.",
        "TVCLTranslationApplicator snapshots source text, geometry, alignment, direction, and writable state; restores the baseline; applies stable-key text/templates/layout; handles menus, dialogs, status panels, lists/grids, wrapping, and RTL; then restores focus, selection, and other live state.",
    ])
    add_callout(document, "VCL handle rule.", "Do not force unnecessary handle recreation or overwrite window state merely to change language. Native controls and Windows messages can regenerate captions, so VCL template/caption interception reapplies only the owned text contract.")

    add_engineering_chapter(document, 13)
    add_paragraphs(document, [
        "TDATFMXLanguageManager connects TTranslationRuntime to FMX form lifecycle and styled-control behavior. It handles open/later forms, tabs, scroll content bounds, browsers, grids, menus, dynamic refresh timers, and lifecycle-safe application without relying on one form-specific event.",
        "TFMXTranslationApplicator snapshots/restores the designer baseline, applies text and layout to the FMX object tree, preserves state/order, and coordinates direction. Template rewrite is essential because FMX styles can recreate visible text after a property was translated.",
    ])
    add_callout(document, "FMX styled-control rule.", "A language switch is not complete merely because Text changed once. Styled controls, tab activation, browser document load, and scroll-layout recalculation can repaint later; lifecycle contracts must reapply only when the owned text is regenerated.")

    add_engineering_chapter(document, 14)
    add_paragraphs(document, [
        "DAT.Components.Core exposes published properties/events and the framework-neutral manager contract. DAT.Components.VCL/FMX implement framework lifecycle and applicator bridges. The LanguageSelector units provide connected designer combo boxes. Registration units expose those components only to RAD Studio design packages.",
        "The design packages are Win32 IDE packages because RAD Studio is a Win32 design host even when the target application also builds Win64. Runtime packages are separated so a deployed application never links IDE registration code.",
    ])
    add_table(document, ["Package", "Role", "Contains / requires"], [
        ["DATLanguageManagerCoreRuntime", "Framework-neutral runtime package.", "Core runtime/manager/component contract used by VCL and FMX runtime packages."],
        ["DATLanguageManagerVCLRuntime", "VCL runtime package.", "VCL applicator, VCL manager, VCL selector, and framework-specific helpers."],
        ["DATLanguageManagerFMXRuntime", "FMX runtime package.", "FMX applicator, FMX manager, FMX selector, and framework-specific helpers."],
        ["DATLanguageManagerVCLDesign", "VCL IDE design package.", "Requires VCL runtime package and contains VCL Register."],
        ["DATLanguageManagerFMXDesign", "FMX IDE design package.", "Requires FMX runtime package and contains FMX Register."],
    ])

    add_engineering_chapter(document, 15)
    add_paragraphs(document, [
        "ComponentSource is a generated source-closure distribution for one target framework. The package generator computes the exact shared and framework-specific set and copies all of it into the kit. The target compiler must see one internally consistent version of that set.",
        r"The optional target-local dependencies folder is <Target>\dependencies\DelphiAppTranslation\source. It is not created automatically because the Studio cannot assume how a target project vendors third-party source or what repository policy permits.",
    ])
    add_table(document, ["Layer", "Units", "Why inseparable"], [
        ["Shared persistence/diagnostics", "DAT.Core.AtomicFile; DAT.Core.Diagnostics", "Preference/layout/runtime parsing depends on the same atomic/error contract."],
        ["Shared runtime", "DAT.Runtime.LanguagePack; Preference; Manager; LayoutOverrides; SplashTranslation; TemplateRewrite", "Pack model, active runtime, restoration, early startup, and later template repair share interfaces."],
        ["Shared component API", "DAT.Components.Core", "VCL/FMX managers and application code compile against its published properties/events/functions."],
        ["VCL layer", "DAT.Runtime.VCL; SplashTranslation.VCL; TemplateRewrite.VCL; DAT.Components.VCL; VCL.LanguageSelector", "All VCL lifecycle, applicator, splash, selector, and rewrite contracts must match."],
        ["FMX layer", "DAT.Runtime.FMX; SplashTranslation.FMX; TemplateRewrite.FMX; DAT.Components.FMX; FMX.LanguageSelector", "All FMX lifecycle, styles, selector, and rewrite contracts must match."],
    ])
    add_callout(document, "Partial-copy failure mode.", "Copying only the units named by the first compiler errors can compile against older interfaces elsewhere on Search path, creating ambiguous unit resolution or runtime behavior. Refresh the complete generated set together.")

    add_engineering_chapter(document, 16)
    add_paragraphs(document, [
        "TComponentIntegrationPackageGenerator produces the recommended non-mutating kit. TTargetBuildDeployer runs selected builds and deploys packs. TIntegrationPackageGenerator and TDelphiIntegrationSourceEditor retain an advanced explicitly authorized source-edit path with preview/backup/restore, but it is not the Wizard default.",
        "Wizard final processing passes ComponentSource to its own build environment without persisting a DPROJ edit. Manual IDE builds require the developer to set Search path in Project Options. Deployment verifies the intended executable/application directory before copying packs.",
    ])
    add_bullets(document, [
        "Search path must preserve existing values and $(DCC_UnitSearchPath).",
        "All configurations/all platforms is the normal setting when one vendored dependency folder serves every target.",
        "Pack deployment is independent of executable compilation; a build can succeed while packs are stale or absent, so both results must be checked.",
        "Unavailable optional removable/network destinations warn and skip; they must not fail an otherwise valid local build.",
        "No integration code should write outside the explicitly selected Studio export, workspace, target dependency, or authorized application destinations.",
    ])

    add_engineering_chapter(document, 17)
    add_paragraphs(document, [
        "TfrmTranslationStudio is the landing and seven-page Maintenance state owner. TfrmSetupWizard is an eight-step validation state machine whose downstream state is invalidated when project/language/provider inputs change. TfrmLocalizationReview is modal within final processing and returns saved decisions to a continuation. DAT.Studio.Translation owns the Studio's own interface-language runtime.",
        "The forms are designer-authored FMX resources. Runtime code should update state, enablement, text, and data; it should not rebuild the form hierarchy that belongs in the FMX designer. Default buttons, Enter/Esc behavior, hidden irrelevant buttons, and asynchronous provider state are explicit UI contracts.",
    ])
    add_table(document, ["Form", "Primary state", "Critical transition"], [
        ["MainForm", "Opened project, scan result, catalog, selected maintenance page, provider/settings, validation/export/integration status.", "Landing to Wizard or Maintenance; language refresh must rebuild current visible content immediately."],
        ["SetupWizard", "Current/highest step, project/profile, destinations, locales, provider, scan/catalog, authorization, final-processing continuation.", "Begin Final Processing disables navigation; Localization Review temporarily suspends and resumes the pipeline."],
        ["LocalizationReview", "Findings, glossary, suggestions, proposals, decisions, output package.", "Close saves accepted decisions and returns to Wizard validation/export."],
    ])

    add_engineering_chapter(document, 18)
    add_paragraphs(document, [
        "Text direction is pack locale metadata interpreted by the framework applicator. Direction and alignment are applied at the smallest correct paragraph/control ownership boundary; the whole interface is not blindly mirrored when the application contains protected LTR identifiers, paths, APIs, or numeric content.",
        "Dynamic text uses stable semantic keys or source-map translation through DATTranslateDynamicText. HTML uses DATTranslateHtmlText and must rebuild the current page on language change. A language switch restores source/dynamic baseline before applying the new target, preventing Arabic text from surviving into Greek or another language.",
    ])
    add_table(document, ["Content class", "Translation contract", "Refresh contract"], [
        ["Designer property", "Stable form.component.property key.", "Applicator updates open form and lifecycle handles later forms."],
        ["Resourcestring", "Stable unit.symbol key plus explicit application use.", "Application code reads translated value at the intended message point."],
        ["Dynamic status/message", "Semantic key/template or source-map contract.", "Recompute immediately on language change without altering live data."],
        ["HTML/report", "Canonical HTML/text templates and semantic keys; table grid owns header/body alignment.", "Rebuild current document synchronously on switch; no previous-language DOM/text may remain."],
        ["Identifier/path/API token", "Protected LTR/nontranslatable token.", "Preserve spelling and isolate direction inside RTL prose."],
        ["Layout rule", "Locale/form/component property decision.", "Restore designer baseline, then apply selected locale rule only."],
    ])

    add_engineering_chapter(document, 19)
    add_paragraphs(document, [
        "TDATDiagnostics and structured exceptions provide severity, operation, path, provider status, stable key, and recovery context. UI code must convert them into actionable status/progress text without exposing secrets. Background provider/build work must always return the form to a closable and internally consistent state.",
        "Performance-sensitive language switching must use already admitted in-memory packs. It should not rescan the project, call a provider, reread every JSON file repeatedly, or rebuild hidden pages that have no active content. Only open/later managed objects and explicit dynamic/HTML refresh consumers participate.",
    ])
    add_table(document, ["Operation", "Bound", "Cancellation / recovery"], [
        ["Provider request", "Configured timeout (default 30 seconds) and bounded batch (default 40 strings).", "Cancel check and retry policy; no endless wait."],
        ["Build/deploy process", "Explicit process timeout and termination wait.", "Controlled process stop, captured output, destination-specific failure."],
        ["Scan", "Progress by stage/file and source snapshot.", "EProjectScanCancelled; prior catalog remains valid."],
        ["Atomic save", "One candidate validation and replace operation.", "Prior valid file and corruption quarantine."],
        ["Language switch", "Admitted pack lookup plus managed open objects.", "Reentrancy/main-thread guards and error behavior event/exception policy."],
    ])

    add_engineering_chapter(document, 20)
    add_bullets(document, [
        "Secrets: Credential Manager/session memory only; never serialize or log API-key text.",
        "Network: development Studio only; HTTPS endpoints determined by selected provider/DeepL plan; response structure and placeholder count validated.",
        "Filesystem: normalize and validate intended paths; atomic writes; no broad recursive deletion; deployment only to verified application folders.",
        "Source: scanner parses text; it does not execute target Pascal, forms, scripts, or arbitrary JSON.",
        "Runtime: admit only compatible packs; reject malformed/wrong-app/wrong-framework/wrong-version/checksum mismatches before selection.",
        "Logs/reports: actionable identifiers and paths are acceptable; secrets and proprietary unneeded text are not.",
        "Distribution: exclude Local AppData workspaces, credentials, transient test folders, compiled output when source-only, and proprietary catalogs unless explicitly licensed.",
        "Dependencies: generated ComponentSource is source code and must be versioned/licensed/reviewed like any vendored library.",
    ])

    add_engineering_chapter(document, 21)
    add_paragraphs(document, [
        "The test suite is intentionally many small Delphi programs rather than one opaque test executable. Each program names and isolates a contract: scanning, placeholders, batching, retries, memory, validation, pack layout, VCL/FMX runtime, lifecycle, direction, wrapping, splash, templates, design streaming, and Studio form creation.",
        "tools\\verify_all.ps1 is the release-level orchestration entry. Contract scripts compile and run groups using the RAD Studio environment. check_shipped_units_complete.ps1 prevents package/kit/source-distribution omissions; check_build_paths_agree.ps1 detects conflicting outputs; check_source_encoding.ps1 detects unsupported encoding drift.",
    ])
    add_table(document, ["Gate", "Required evidence"], [
        ["Source hygiene", "Encoding/path checks, no stale lock or unintended generated files, Git diff limited to approved scope."],
        ["Compile", "Studio, package set, tools, and affected tests compile for intended platform/configuration."],
        ["Foundation", "Types, JSON, workspace, atomic persistence, locale, placeholders, retry, terminology, memory, validation."],
        ["Scanner", "VCL/FMX form discovery, Pascal/resource contracts, counts, context, quality, source snapshot."],
        ["Runtime", "VCL/FMX load/apply/switch-back/restart, state preservation, lifecycle, templates, splash, wrapping, RTL."],
        ["Design-time", "VCL/FMX form streaming and package registration with no missing-class or published-property error."],
        ["Studio", "Main/Wizard/Review form creation, UI transitions, provider timeout/cancel, scan consistency, self-localization."],
        ["Pilot matrix", "Real target project, all supported languages, HTML/dynamic text, LTR/RTL, actual output folders."],
        ["Documentation", "DOCX source-driven accuracy, PDF generation, page rendering, visual inspection, accessibility/structure audits."],
    ])

    document.add_page_break()
    add_engineering_chapter(document, 22)
    add_paragraphs(document, [
        "This chapter is the production PAS inventory. Each entry names its file, purpose, reason for separation, public surface, direct DAT dependencies, direct production consumers, named test consumers, and change risk. Dependencies are parsed from the current source; descriptions state the architectural responsibility rather than merely repeating the filename.",
    ])
    grouped: dict[str, list[PascalUnit]] = defaultdict(list)
    for unit in units:
        grouped[unit.subsystem].append(unit)
    subsystem_order = ("core", "scan", "provider", "review", "validation", "runtime", "components", "design", "integration", "studio")
    for subsystem in subsystem_order:
        document.add_heading(f"23.{subsystem_order.index(subsystem) + 1} {subsystem.title()} units", level=2)
        add_paragraphs(document, [f"These units own {SUBSYSTEM_PURPOSE[subsystem]}. The entries are ordered by unit name."])
        for unit in grouped[subsystem]:
            add_unit_entry(document, unit)

    document.add_page_break()
    add_engineering_chapter(document, 23)
    add_paragraphs(document, [
        "Every Delphi test DPR is listed below. Test-support PAS/DFM/FMX units and formal contract fixtures follow. These files are not runtime dependencies; they supply isolated evidence and intentionally artificial application behavior.",
    ])
    document.add_heading("24.1 Test executable programs", level=2)
    test_paths = sorted((ROOT / "tools" / "tests").glob("*.dpr"))
    for path in test_paths:
        add_test_program_entry(document, path)
    document.add_heading("24.2 Test-support Pascal units", level=2)
    support_paths = sorted((ROOT / "tools" / "tests").rglob("*.pas"))
    for path in support_paths:
        document.add_heading(path.name, level=3)
        add_support_pascal_entry(document, path)
    document.add_heading("24.3 Formal scan/layout Pascal fixtures", level=2)
    for path in sorted((ROOT / "contracts").rglob("*.pas")):
        document.add_heading(path.name, level=3)
        add_support_pascal_entry(document, path)

    document.add_page_break()
    add_engineering_chapter(document, 24)
    document.add_heading("25.1 Application, command-line, and package entry points", level=2)
    add_important_file(document, "DelphiAppTranslationStudio.dpr", "FMX application entry point that initializes the Studio translation runtime and creates the main Studio form.", "Defines process startup ordering and the first form/language initialization boundary.")
    add_important_file(document, "DelphiAppTranslationStudio.dproj", "RAD Studio project metadata, unit search paths, platforms/configurations, resources, outputs, and main form identity.", "A build can compile different source or write to a different output when DPROJ configuration drifts, even if DPR is unchanged.")
    add_important_file(document, "tools\\DATBatch.dpr", "Command-line/batch entry point for supported noninteractive translation operations.", "Provides automation without coupling scripts directly to Studio form classes.")
    for path in sorted((ROOT / "packages").rglob("*.dpk")):
        add_package_entry(document, path)
        dproj = path.with_suffix(".dproj")
        if dproj.exists():
            add_key_value(document, "Companion DPROJ", str(dproj.relative_to(ROOT)) + " - platform/configuration/output metadata for this package.")

    add_engineering_chapter(document, 25)
    document.add_heading("26.1 JSON schemas", level=2)
    for path in sorted((ROOT / "source" / "schemas").glob("*.json")):
        description, properties, required = schema_summary(path)
        document.add_heading(path.name, level=3)
        add_key_value(document, "File", str(path.relative_to(ROOT)))
        add_key_value(document, "What it defines", description)
        add_key_value(document, "Top-level properties", ", ".join(properties) or "Schema composition does not expose direct top-level properties.")
        add_key_value(document, "Required top-level properties", ", ".join(required) or "Requirements are expressed in nested/composed schema clauses.")
        add_key_value(document, "Change rule", "A schema change requires compatible reader/writer changes, version handling, fixtures, round-trip tests, and release notes; never change emitted structure alone.")
    document.add_heading("26.2 Production FMX form resources", level=2)
    form_descriptions = {
        "DAT.Studio.MainForm.fmx": "Designer-authored landing screen and seven Maintenance pages, navigation rail, dialogs, status, and control defaults.",
        "DAT.Studio.SetupWizard.fmx": "Designer-authored eight-step Wizard, step rail, page controls, keyboard defaults, authorization, and final-processing layout.",
        "DAT.Studio.LocalizationReview.fmx": "Designer-authored findings/glossary/suggestions/layout-decision review window.",
    }
    for name, description in form_descriptions.items():
        path = ROOT / "source" / "studio" / name
        add_important_file(document, str(path.relative_to(ROOT)), description, "It is the editable IDE layout authority paired with the same-base-name PAS class. Runtime code must not replace this designer ownership without explicit approval.")
    document.add_heading("26.3 Test form resources", level=2)
    for path in sorted((ROOT / "tools" / "tests").rglob("*.dfm")) + sorted((ROOT / "tools" / "tests").rglob("*.fmx")):
        add_important_file(document, str(path.relative_to(ROOT)), "Designer resource streamed by a lifecycle or design test.", "Proves published properties, inheritance, object names, lifecycle behavior, and form-resource compatibility.")

    add_engineering_chapter(document, 26)
    document.add_heading("27.1 Build, verification, distribution, and guide tools", level=2)
    tools = [
        ("tools\\build_packages.ps1", "Builds the five runtime/design package projects in controlled platform/configuration order.", "Package installation and ComponentSource correctness depend on reproducible package outputs."),
        ("tools\\verify_all.ps1", "Runs the broad engineering verification suite.", "Primary pre-release gate across foundation, scanner, runtime, components, Studio, and artifacts."),
        ("tools\\check_shipped_units_complete.ps1", "Compares required runtime/component units with shipped packages/kits/distributions.", "Prevents a compile/runtime failure caused by an omitted dependency."),
        ("tools\\check_build_paths_agree.ps1", "Checks project, script, and documentation output/search-path agreement.", "Prevents testing one binary while shipping another."),
        ("tools\\check_source_encoding.ps1", "Audits source encoding and known text hazards.", "Scanner/provider/UI correctness depends on stable Unicode source."),
        ("tools\\build_source_distribution.ps1", "Creates the source-only distribution with approved files and structure.", "Distribution must be complete while excluding secrets, transient workspaces, and unintended binaries."),
        ("tools\\run_form_scan_contracts.ps1", "Builds/runs formal DFM/FMX scan contracts.", "Protects canonical scan keys, counts, inclusion/exclusion, and context behavior."),
        ("tools\\run_layout_contracts.ps1", "Builds/runs layout and direction contracts.", "Protects designer ownership, wrapping, fitting, RTL, and code-positioned exclusions."),
        ("tools\\run_pascal_scan_contracts.ps1", "Builds/runs Pascal resourcestring/runtime-assignment contracts.", "Protects text extraction without executing target source or translating arbitrary literals."),
        ("tools\\Install-DATLanguageManagerComponents.ps1", "Builds/locates and guides installation of design components.", "Keeps Tool Palette setup repeatable without copying arbitrary BPLs into RAD Studio folders."),
        ("tools\\render_guides_pdf.py", "Converts documentation DOCX files to companion PDFs through LibreOffice.", "Keeps printable output tied to editable source documents."),
        ("tools\\build_user_guide.py", "Builds the detailed source-driven User Guide and its static printable TOC.", "Documents actual application behavior and provides deterministic regeneration."),
        ("tools\\build_wizard_engineering_guides.py", "Builds this Wizard Guide and Engineering Guide from current source inventories and curated architecture descriptions.", "Keeps exhaustive file/dependency reference synchronized with repository content."),
        ("tools\\finalize_guides.ps1", "Performs the maintained guide finalization workflow.", "Ensures DOCX/PDF artifacts and verification steps remain repeatable."),
    ]
    for path, what, why in tools:
        add_important_file(document, path, what, why)
    document.add_heading("27.2 Important engineering reference documents", level=2)
    references = [
        ("docs\\guides\\Total Stabilization Contracts.md", "Canonical cross-cutting runtime, layout, language-switch, and universal-fix contracts."),
        ("docs\\guides\\Translation Studio Completed Test Report and Remediation Plan.md", "Evidence, resolved defects, residual risks, and verification results."),
        ("docs\\guides\\TDATLanguageManager Deep-Dive and Implementation Plan.md", "Manager architecture and staged implementation rationale."),
        ("docs\\STAGES-1-3-REVIEW-AND-SCANNING-ENGINEERING-NOTES.md", "Detailed scan/context/review engineering notes."),
        ("docs\\STAGES-4-10-RUNTIME-LAYOUT-ENGINEERING-NOTES.md", "Detailed runtime/layout/direction/stabilization engineering notes."),
        ("docs\\guides\\Release Checklist.md", "Human release verification checklist complementing automated tests."),
    ]
    for path, what in references:
        add_important_file(document, path, what, "Read before changing the named subsystem; these records explain contracts that a source signature alone cannot express.")

    add_engineering_chapter(document, 27)
    add_paragraphs(document, [
        "Every production unit entry in Chapter 23 includes direct DAT dependencies and direct production consumers computed from current source. A dependency means the unit imports another DAT unit; a consumer means another production unit imports it. The list is direct, not the complete transitive closure.",
        "Leaf/foundation units may have no DAT dependencies but many consumers. Adapter/entry units may have many dependencies and no production consumer because a DPR/DPK or application code selects them. Test consumers show named DPRs that import the unit directly; indirect coverage is broader.",
    ])
    add_table(document, ["Graph signal", "Likely meaning", "Required engineering response"], [
        ["Many consumers", "Shared API or foundational contract.", "Treat signature/behavior changes as cross-subsystem and run full verification."],
        ["Framework-neutral unit imports VCL/FMX", "Layer violation.", "Move framework behavior into the applicator/component adapter."],
        ["Runtime imports provider/studio", "Offline trust-boundary violation.", "Remove network/UI dependency from shipped runtime."],
        ["Design unit in runtime package", "IDE-only code can leak into deployment.", "Correct DPK containment/requirements."],
        ["ComponentSource missing a direct/transitive dependency", "Generated kit cannot compile consistently.", "Fix closure computation and shipped-unit check; do not hand-copy one file."],
        ["No named tests for high-risk unit", "Coverage may be only indirect.", "Add focused contract test when behavior cannot be proven by existing subsystem tests."],
    ])

    add_engineering_chapter(document, 28)
    add_steps(document, [
        "Define the contract and affected ownership layer before editing. Read the relevant reference document and unit entries in Chapter 23.",
        "Check Git status and remote currency; preserve unrelated work. Make a pre-change backup for material changes.",
        "Trace direct dependencies, consumers, package membership, ComponentSource closure, schemas, and tests. Do not edit only the first file named by a compiler error.",
        "Make the smallest universal change. Do not add named-project/version exceptions, runtime UI construction, or helper layers without explicit approval.",
        "Update schema/version or compatibility handling when serialized contracts change. Preserve atomic persistence and older valid recovery files.",
        "Compile affected foundation/runtime/package/Studio targets and run focused tests first, then the complete verification suite.",
        "Run real VCL and FMX pilot applications across LTR/RTL, language-to-language, switch-back, restart, dynamic/HTML, layout, and supported platforms.",
        "Update User, Wizard, Engineering, help, release, and test documentation as applicable. Render and visually inspect DOCX/PDF outputs.",
        "Review Git diff, stage only intended files, commit clearly, and push configured public/private remotes without rewriting history.",
    ])

    add_engineering_chapter(document, 29)
    add_bullets(document, [
        "[ ] Product invariants and trust boundaries remain true.",
        "[ ] All production PAS units compile in their intended packages/projects.",
        "[ ] Five package projects build; Win32 Release design BPLs stream components in RAD Studio.",
        "[ ] Complete ComponentSource closure contains matching shared and framework units.",
        "[ ] Development catalogs, settings, preferences, glossaries, layout overrides, and packs pass atomic/recovery tests.",
        "[ ] Canonical scanner totals agree between Wizard and Maintenance for identical inputs.",
        "[ ] Provider timeout, cancellation, retries, placeholders, language codes, quota/auth errors, and credential isolation pass.",
        "[ ] Localization Review findings, glossary, suggestions, layout decisions, HTML package, and resume continuation pass.",
        "[ ] Validation blocks structural errors without pretending to approve language quality.",
        "[ ] Pack admission rejects wrong application/framework/version/checksum/source contract.",
        "[ ] VCL and FMX preserve state and designer baseline through every language transition.",
        "[ ] LTR, RTL, protected LTR tokens, space wrapping, intelligent hyphenation, readable font fitting, and neighbor alignment pass.",
        "[ ] English/source switch-back repaints current forms/HTML immediately and previous-language content is absent.",
        "[ ] Startup preference, missing/stale pack fallback, splash, templates, later forms, dialogs, menus, grids, reports, and dynamic text pass.",
        "[ ] Build/deployment output and actual running executable folder contain the identical intended pack set.",
        "[ ] Source distribution excludes secrets/transient work and includes required source, packages, schemas, tools, guides, and licenses.",
        "[ ] Documentation is source-current, rendered, visually inspected, accessible, committed, and pushed.",
    ])

    document.add_page_break()
    add_engineering_chapter(document, 30)
    add_table(document, ["Purpose", "Path / identifier"], [
        ["Main project", "DelphiAppTranslationStudio.dproj / DelphiAppTranslationStudio.dpr"],
        ["RAD environment", r"C:\Program Files (x86)\Embarcadero\Studio\37.0\bin\rsvars.bat"],
        ["Win32 compiler", r"C:\Program Files (x86)\Embarcadero\Studio\37.0\bin\dcc32.exe"],
        ["Production source", r"source\core, scan, provider, review, validation, runtime, components, design, integration, studio"],
        ["Schemas", r"source\schemas"],
        ["Packages", r"packages\runtime and packages\design"],
        ["Tests", r"tools\tests"],
        ["Release verification", r"tools\verify_all.ps1"],
        ["Studio settings", r"%LOCALAPPDATA%\DelphiAppTranslationStudio"],
        ["Project workspace", r"%LOCALAPPDATA%\DelphiAppTranslationStudio\Workspaces\<ApplicationId>"],
        ["Generated kit", r"export\component-integration\<ApplicationId>"],
        ["Review package", r"export\localization-review\<ApplicationId>\<locale>"],
        ["Vendored target dependencies", r"<Target>\dependencies\DelphiAppTranslation\source"],
        ["Target packs", r"<EXE>\Localization\Languages"],
        ["DeepL credential target", "DelphiAppTranslationStudio/Providers/DeepL"],
        ["Google credential target", "DelphiAppTranslationStudio/Providers/Google Cloud Translation"],
    ])

    add_license_appendix(document, ENGINEERING_HEADINGS[31])

    mark_first_rows_as_accessibility_headers(document)
    path = GUIDES_DIR / "Delphi App Translation Studio Engineering Guide.docx"
    finish_document(document, path)
    return path


def extract_toc_pages(pdf_path: Path, headings: list[str]) -> dict[str, int]:
    from pypdf import PdfReader

    reader = PdfReader(str(pdf_path))
    texts = [(page.extract_text() or "") for page in reader.pages]
    body_pages: dict[str, int] = {}
    physical_pages: dict[str, int] = {}
    for heading in headings:
        matches = [index for index, text in enumerate(texts) if heading in text]
        if not matches:
            continue
        physical_pages[heading] = matches[-1]
    if not physical_pages:
        return body_pages
    first_heading = headings[0]
    body_start = physical_pages.get(first_heading, min(physical_pages.values()))
    for heading, physical in physical_pages.items():
        body_pages[heading] = physical - body_start + 1
    return body_pages


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--toc-json", type=Path)
    parser.add_argument("--discover-toc", nargs=2, metavar=("WIZARD_PDF", "ENGINEERING_PDF"))
    parser.add_argument("--output-toc-json", type=Path, default=TOC_DEFAULTS)
    args = parser.parse_args()

    if args.discover_toc:
        wizard_pdf, engineering_pdf = (Path(item) for item in args.discover_toc)
        data = {
            "wizard": extract_toc_pages(wizard_pdf, WIZARD_HEADINGS),
            "engineering": extract_toc_pages(engineering_pdf, ENGINEERING_HEADINGS),
        }
        args.output_toc_json.parent.mkdir(parents=True, exist_ok=True)
        args.output_toc_json.write_text(json.dumps(data, indent=2), encoding="utf-8")
        print(args.output_toc_json)
        return

    pages = load_toc_pages(args.toc_json)
    wizard = build_wizard_guide(pages.get("wizard", {}))
    engineering = build_engineering_guide(pages.get("engineering", {}))
    print(wizard)
    print(engineering)


if __name__ == "__main__":
    main()
