from __future__ import annotations

import base64
import html
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader


PROJECT_ROOT = Path(__file__).resolve().parents[1]
GUIDE_DIRECTORY = PROJECT_ROOT / "docs" / "guides"
PDF_DIRECTORY = PROJECT_ROOT / "docs" / "pdf"
TEMP_DIRECTORY = PROJECT_ROOT / "tmp" / "pdfs"
PRINT_SCRIPT = PROJECT_ROOT / "tools" / "print_guide_pdf.js"
ICON = PROJECT_ROOT / "images and icons" / "DelphiAppTranslationStudio-Icon-Master-v2_150.png"
GUIDES = (
    "Delphi App Translation Studio User Guide",
    "Delphi App Translation Studio Setup Wizard Guide",
    "Delphi App Translation Studio Engineering Guide",
)


def iter_blocks(document: Document):
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def normalize(value: str) -> str:
    value = value.translate(str.maketrans({
        "\ufb00": "ff", "\ufb01": "fi", "\ufb02": "fl",
        "\ufb03": "ffi", "\ufb04": "ffl", "\ufb05": "ft",
        "\ufb06": "st",
    }))
    return re.sub(r"\s+", " ", value).strip()


def image_data() -> str:
    if not ICON.exists():
        return ""
    encoded = base64.b64encode(ICON.read_bytes()).decode("ascii")
    return f'<img class="cover-logo" src="data:image/png;base64,{encoded}" alt="Delphi App Translation Studio icon">'


def table_html(table: Table) -> str:
    rows = [[normalize(cell.text) for cell in row.cells] for row in table.rows]
    if len(rows) == 1 and len(rows[0]) == 1:
        text = html.escape(rows[0][0])
        first, separator, rest = text.partition(". ")
        if separator:
            text = f"<strong>{first}.</strong> {rest}"
        return f'<aside class="callout">{text}</aside>'
    if not rows:
        return ""
    header = "".join(f"<th>{html.escape(value)}</th>" for value in rows[0])
    body = "".join(
        "<tr>" + "".join(f"<td>{html.escape(value)}</td>" for value in row) + "</tr>"
        for row in rows[1:]
    )
    return f"<table><thead><tr>{header}</tr></thead><tbody>{body}</tbody></table>"


def build_html(docx_path: Path, page_map: dict[str, int] | None = None) -> tuple[str, list[str]]:
    document = Document(docx_path)
    headings = [
        normalize(p.text)
        for p in document.paragraphs
        if p.style.name in ("Heading 1", "Heading 2", "Heading 3") and normalize(p.text)
    ]
    title = next((normalize(p.text) for p in document.paragraphs if p.style.name == "Title"), docx_path.stem)
    subtitle = next((normalize(p.text) for p in document.paragraphs if p.style.name == "Subtitle"), "")
    metadata = next((normalize(p.text) for p in document.paragraphs if "Last changed:" in p.text), "Version 1.0")

    toc_rows = []
    for heading in headings:
        level = 1
        for paragraph in document.paragraphs:
            if normalize(paragraph.text) == heading and paragraph.style.name.startswith("Heading"):
                level = int(paragraph.style.name[-1])
                break
        number = "" if page_map is None else str(page_map.get(heading, ""))
        toc_rows.append(
            f'<div class="toc-row toc-{level}"><span>{html.escape(heading)}</span>'
            f'<span class="toc-dots"></span><span>{number}</span></div>'
        )

    body_parts: list[str] = []
    started = False
    list_kind: str | None = None

    def close_list() -> None:
        nonlocal list_kind
        if list_kind:
            body_parts.append(f"</{list_kind}>")
            list_kind = None

    for block in iter_blocks(document):
        if isinstance(block, Paragraph):
            style = block.style.name
            text = normalize(block.text)
            if style == "Heading 1":
                started = True
            if not started or not text:
                continue
            if style in ("List Number", "List Bullet"):
                wanted = "ol" if style == "List Number" else "ul"
                if list_kind != wanted:
                    close_list()
                    list_kind = wanted
                    body_parts.append(f"<{wanted}>")
                body_parts.append(f"<li>{html.escape(text)}</li>")
                continue
            close_list()
            if style.startswith("Heading "):
                level = int(style[-1])
                body_parts.append(f"<h{level}>{html.escape(text)}</h{level}>")
            elif style == "Code Block":
                body_parts.append(f"<pre>{html.escape(text)}</pre>")
            else:
                body_parts.append(f"<p>{html.escape(text)}</p>")
        elif started:
            close_list()
            body_parts.append(table_html(block))
    close_list()

    css = """
      @page { size: Letter; }
      * { box-sizing: border-box; }
      body { margin: 0; color: #26384a; font-family: Calibri, Arial, sans-serif; font-size: 10.7pt; line-height: 1.34; }
      .cover { min-height: 8.75in; display: flex; flex-direction: column; align-items: center; justify-content: center; text-align: center; page-break-after: always; }
      .cover-logo { width: 1.45in; height: 1.45in; object-fit: contain; margin-bottom: 0.28in; }
      .cover h1 { margin: 0; color: #234c80; font-size: 29pt; line-height: 1.08; }
      .cover h2 { margin: 0.14in 0 0.25in; color: #5d7693; font-size: 15pt; font-weight: 400; max-width: 6in; }
      .accent { width: 6.25in; height: 0.09in; display: flex; margin: 0.1in 0 0.35in; }
      .accent span:first-child { background: #234c80; flex: 1; } .accent span:last-child { background: #f28a1b; flex: 1; }
      .meta { white-space: pre-line; color: #5d7693; font-size: 10.5pt; }
      .toc { min-height: 8.75in; page-break-after: always; }
      .toc h1 { color: #2e74b5; font-size: 18pt; margin: 0 0 0.28in; }
      .toc-row { display: flex; align-items: baseline; margin: 0.055in 0; color: #26384a; }
      .toc-2 { padding-left: 0.24in; font-size: 9.7pt; } .toc-3 { padding-left: 0.48in; font-size: 9.2pt; }
      .toc-dots { flex: 1; border-bottom: 1px dotted #9aabba; margin: 0 0.08in 0.04in; }
      main h1 { color: #2e74b5; font-size: 16pt; margin: 0.24in 0 0.1in; break-after: avoid; }
      main h2 { color: #2e74b5; font-size: 13pt; margin: 0.18in 0 0.08in; break-after: avoid; }
      main h3 { color: #1f4d78; font-size: 12pt; margin: 0.14in 0 0.06in; break-after: avoid; }
      p { margin: 0 0 0.085in; orphans: 3; widows: 3; }
      ol, ul { margin: 0.03in 0 0.1in 0.26in; padding-left: 0.25in; }
      li { margin: 0 0 0.055in; padding-left: 0.03in; }
      table { width: 100%; border-collapse: collapse; margin: 0.1in 0 0.16in; font-size: 9.2pt; }
      thead { display: table-header-group; } tr { break-inside: avoid; }
      th { background: #234c80; color: white; padding: 7px 8px; text-align: left; vertical-align: middle; }
      td { border: 1px solid #c8d6e5; padding: 7px 8px; text-align: left; vertical-align: top; }
      tbody tr:nth-child(even) td { background: #f4f7fb; }
      .callout { background: #eaf3ff; border-left: 5px solid #f28a1b; color: #163a63; padding: 0.13in 0.16in; margin: 0.1in 0 0.16in; break-inside: avoid; }
      .callout strong { color: #234c80; }
      pre { background: #f3f5f7; border: 1px solid #d7e0e8; padding: 0.12in; font: 8.5pt/1.35 Consolas, monospace; white-space: pre-wrap; break-inside: avoid; }
    """
    output = f"""<!doctype html><html><head><meta charset="utf-8"><style>{css}</style></head><body>
      <section class="cover">{image_data()}<h1>{html.escape(title)}</h1><h2>{html.escape(subtitle)}</h2>
      <div class="accent"><span></span><span></span></div><div class="meta">{html.escape(metadata)}</div></section>
      <section class="toc"><h1>Table of Contents</h1>{''.join(toc_rows)}</section>
      <main>{''.join(body_parts)}</main></body></html>"""
    return output, headings


def page_map_from_pdf(pdf_path: Path, headings: list[str]) -> dict[str, int]:
    pages = [normalize(page.extract_text() or "") for page in PdfReader(pdf_path).pages]
    result: dict[str, int] = {}
    for heading in headings:
        needle = normalize(heading)
        # Page 1 is the cover and page 2 is the generated TOC, which contains
        # every heading. Search only the body or every entry would map to 2.
        for index, page_text in enumerate(pages[2:], start=3):
            if needle in page_text:
                result[heading] = index
                break
    return result


def print_pdf(html_path: Path, pdf_path: Path, title: str) -> None:
    bundled_dependencies = Path(sys.executable).resolve().parents[1]
    bundled_node = bundled_dependencies / "node" / "bin" / "node.exe"
    node = bundled_node if bundled_node.exists() else Path(shutil.which("node") or "")
    if not node.exists():
        raise RuntimeError("Node.js was not found for Playwright PDF generation.")
    environment = os.environ.copy()
    bundled_node_modules = bundled_dependencies / "node" / "node_modules"
    if bundled_node_modules.exists():
        environment["NODE_PATH"] = str(bundled_node_modules)
    subprocess.run(
        [str(node), str(PRINT_SCRIPT), str(html_path), str(pdf_path), title],
        check=True,
        env=environment,
    )


def render_guide(name: str) -> None:
    docx_path = GUIDE_DIRECTORY / f"{name}.docx"
    final_pdf = PDF_DIRECTORY / f"{name}.pdf"
    html_path = TEMP_DIRECTORY / f"{name}.html"
    draft_pdf = TEMP_DIRECTORY / f"{name}.draft.pdf"
    intermediate_pdf = TEMP_DIRECTORY / f"{name}.intermediate.pdf"
    first_html, headings = build_html(docx_path)
    html_path.write_text(first_html, encoding="utf-8")
    print_pdf(html_path, draft_pdf, name)
    mapped_html, _ = build_html(docx_path, page_map_from_pdf(draft_pdf, headings))
    html_path.write_text(mapped_html, encoding="utf-8")
    print_pdf(html_path, intermediate_pdf, name)
    final_html, _ = build_html(
        docx_path, page_map_from_pdf(intermediate_pdf, headings)
    )
    html_path.write_text(final_html, encoding="utf-8")
    print_pdf(html_path, final_pdf, name)
    print(final_pdf)


def main() -> None:
    PDF_DIRECTORY.mkdir(parents=True, exist_ok=True)
    TEMP_DIRECTORY.mkdir(parents=True, exist_ok=True)
    for name in GUIDES:
        render_guide(name)


if __name__ == "__main__":
    main()
