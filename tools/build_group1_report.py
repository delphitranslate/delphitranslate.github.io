from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "guides" / "Translation Studio Group 1 Runtime Lifecycle and State Ownership Report.md"
OUTPUT = ROOT / "docs" / "guides" / "Translation Studio Group 1 Runtime Lifecycle and State Ownership Report.docx"

BLUE = "1F4E78"
ORANGE = "F28C18"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MUTED = "5B6573"
CONTENT_WIDTH_DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start),
                        ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_pr.append(repeat_header)


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_inline_markup(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(BLUE)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, "17365D"),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))


def add_page_furniture(section):
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("DELPHI APP TRANSLATION STUDIO  |  GROUP 1")
    set_font(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    add_field(footer, " PAGE ")


def create_numbering(doc, ordered):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId")))
                    for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId")))
               for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "\u2022")
    level.append(lvl_text)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_id_node))
    p_pr.append(num_pr)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(88)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("DELPHI APP TRANSLATION STUDIO")
    set_font(run, size=12, bold=True, color=ORANGE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Group 1 Runtime Lifecycle\nand State Ownership")
    set_font(run, size=28, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("Implementation and Validation Report")
    set_font(run, size=15, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Status")
    set_font(run, size=10, bold=True, color=BLUE)
    run = p.add_run("  Complete - automated validation passed")
    set_font(run, size=10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Last changed")
    set_font(run, size=10, bold=True, color=BLUE)
    run = p.add_run("  August 9, 2026")
    set_font(run, size=10.5)

    p = doc.add_paragraph()
    run = p.add_run("Safety boundary")
    set_font(run, size=10, bold=True, color=BLUE)
    run = p.add_run("  No target-application source was changed")
    set_font(run, size=10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(58)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Static text is localized automatically. Live state remains owned by the application.")
    set_font(run, size=13, italic=True, color=BLUE)
    doc.add_page_break()


def add_table(doc, rows):
    column_count = len(rows[0])
    if column_count == 2:
        widths = [2700, 6660]
    elif column_count == 3:
        if rows[0][0].strip() == "Runtime text role":
            widths = [1700, 3400, 4260]
        else:
            widths = [4200, 2580, 2580]
    else:
        widths = [CONTENT_WIDTH_DXA // column_count] * column_count
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_markup(paragraph, value.strip())
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build():
    lines = SOURCE.read_text(encoding="utf-8-sig").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_page_furniture(section)
    add_title_page(doc)

    index = 0
    in_code = False
    code_lines = []
    ordered_num_id = None
    bullet_num_id = None
    while index < len(lines):
        line = lines[index]
        if line.startswith("# Delphi App Translation Studio") or \
                line.startswith("## Group 1 Runtime") or \
                line.startswith("**Last changed:") or \
                line.startswith("**Status:") or \
                line.startswith("**Scope:") or \
                line.startswith("**Target-source policy:"):
            index += 1
            continue
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                for code_index, code_line in enumerate(code_lines):
                    if code_index:
                        p.add_run("\n")
                    run = p.add_run(code_line)
                    run.font.name = "Consolas"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
                    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor.from_string("17365D")
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and \
                re.match(r"^\|[\s:|-]+\|$", lines[index + 1]):
            table_lines = [line]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            rows = [[cell.strip() for cell in row.strip("|").split("|")]
                    for row in table_lines]
            add_table(doc, rows)
            continue
        if line.startswith("## "):
            ordered_num_id = None
            bullet_num_id = None
            if line == "## 10. Exit decision":
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            add_inline_markup(p, line[3:])
        elif re.match(r"^\d+\. ", line):
            if ordered_num_id is None:
                ordered_num_id = create_numbering(doc, True)
            bullet_num_id = None
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            apply_numbering(p, ordered_num_id)
            add_inline_markup(p, re.sub(r"^\d+\. ", "", line))
        elif line.lstrip().startswith("- "):
            if line.startswith("- "):
                ordered_num_id = None
            if bullet_num_id is None:
                bullet_num_id = create_numbering(doc, False)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            apply_numbering(p, bullet_num_id)
            add_inline_markup(p, line.lstrip()[2:])
        elif line.strip():
            ordered_num_id = None
            bullet_num_id = None
            p = doc.add_paragraph()
            add_inline_markup(p, line)
        index += 1

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    doc.core_properties.title = "Translation Studio Group 1 Runtime Lifecycle and State Ownership Report"
    doc.core_properties.subject = "Group 1 implementation and validation evidence"
    doc.core_properties.author = "Delphi App Translation Studio Project"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
