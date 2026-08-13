from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
GUIDES_DIR = PROJECT_ROOT / "docs" / "guides"
ICON = (
    PROJECT_ROOT
    / "images and icons"
    / "DelphiAppTranslationStudio-Icon-Master-v2_150.png"
)
LAST_CHANGED = "August 12, 2026"
BLUE = "234C80"
BRIGHT_BLUE = "1974DF"
ORANGE = "F28A1B"
PALE_BLUE = "EAF3FF"
INK = "163A63"
GRAY = "5D7693"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)

def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    if sum(widths) != 9360:
        raise ValueError("Table widths must total 9360 DXA.")
    table.autofit = False
    properties = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblCellMar"):
        existing = properties.find(qn(tag))
        if existing is not None:
            properties.remove(existing)

    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), "9360")
    table_width.set(qn("w:type"), "dxa")
    properties.append(table_width)
    table_indent = OxmlElement("w:tblInd")
    table_indent.set(qn("w:w"), str(indent))
    table_indent.set(qn("w:type"), "dxa")
    properties.append(table_indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    properties.append(layout)

    margins = OxmlElement("w:tblCellMar")
    for side, value in (
        ("top", 80), ("bottom", 80), ("start", 120), ("end", 120)
    ):
        margin = OxmlElement(f"w:{side}")
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    properties.append(margins)

    grid = table._tbl.tblGrid
    for column in list(grid):
        grid.remove(column)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "The table of contents is refreshed during document finalization."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction_text, separate, placeholder, end])


def set_page_number_format(section, fmt: str, start: int | None = None) -> None:
    properties = section._sectPr
    page_number = properties.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        properties.append(page_number)
    page_number.set(qn("w:fmt"), fmt)
    if start is not None:
        page_number.set(qn("w:start"), str(start))


def unlink_headers(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False


def configure_page(section) -> None:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_header_footer(section, title: str, numbered: bool) -> None:
    unlink_headers(section)
    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if numbered:
        footer.add_run("Page ")
        add_field(footer, "PAGE")
    for run in footer.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)


def setup_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("26384A")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in (
        ("Title", 29, BLUE),
        ("Subtitle", 15, GRAY),
        ("TOC Heading", 16, "2E74B5"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.keep_with_next = True
        if name in ("Heading 1", "TOC Heading"):
            style.paragraph_format.space_before = Pt(18)
            style.paragraph_format.space_after = Pt(10)
        elif name == "Heading 2":
            style.paragraph_format.space_before = Pt(14)
            style.paragraph_format.space_after = Pt(7)
        elif name == "Heading 3":
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(5)
        else:
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(8)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Cascadia Mono"
    code.font.size = Pt(8)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(5)

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Aptos"
    callout.font.size = Pt(9)
    callout.font.color.rgb = RGBColor.from_string(INK)
    callout.paragraph_format.left_indent = Inches(0.2)
    callout.paragraph_format.right_indent = Inches(0.2)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(6)


def add_cover(document: Document, title: str, subtitle: str) -> None:
    section = document.sections[0]
    configure_page(section)
    add_header_footer(section, "", False)

    document.add_paragraph("")
    document.add_paragraph("")
    logo = document.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ICON.exists():
        logo.add_run().add_picture(str(ICON), width=Inches(1.5))

    heading = document.add_paragraph(style="Title")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run(title)
    subheading = document.add_paragraph(style="Subtitle")
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheading.add_run(subtitle)

    accent = document.add_table(rows=1, cols=2)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    accent.columns[0].width = Inches(3.25)
    accent.columns[1].width = Inches(3.25)
    set_table_geometry(accent, [4680, 4680], indent=0)
    set_cell_shading(accent.cell(0, 0), BLUE)
    set_cell_shading(accent.cell(0, 1), ORANGE)
    for cell in accent.rows[0].cells:
        cell.height = Inches(0.08)

    document.add_paragraph("")
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Version 1.0\nLast changed: {LAST_CHANGED}\n"
        "Windows • Delphi VCL and FireMonkey • Win32 and Win64"
    )
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(GRAY)


def add_toc_section(document: Document, guide_title: str,
                    contents: list[tuple[str, int]] | None = None) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(section)
    set_page_number_format(section, "lowerRoman", 1)
    add_header_footer(section, guide_title, True)
    document.add_paragraph("Table of Contents", style="TOC Heading")
    paragraph = document.add_paragraph()
    if contents:
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        paragraph.add_run()._r.extend([begin, instruction, separate])
        for index, (heading, page_number) in enumerate(contents):
            if index:
                paragraph.add_run().add_break()
            entry = paragraph.add_run(heading)
            entry.font.name = "Aptos"
            entry.font.size = Pt(9)
            entry.font.color.rgb = RGBColor.from_string(INK)
            dots = max(4, 72 - len(heading))
            leader = paragraph.add_run(
                " " + "." * dots + " " + str(page_number))
            leader.font.name = "Aptos"
            leader.font.size = Pt(9)
            leader.font.color.rgb = RGBColor.from_string(GRAY)
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        paragraph.add_run()._r.append(end)
    else:
        add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')
    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(content_section)
    set_page_number_format(content_section, "decimal", 1)
    add_header_footer(content_section, guide_title, True)


def add_paragraphs(document: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        document.add_paragraph(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_steps(document: Document, items: list[str]) -> None:
    numbering = document.part.numbering_part.element
    style_num_id = int(
        document.styles["List Number"].element.pPr.numPr.numId.val
    )
    base_number = numbering.xpath(
        f'./w:num[@w:numId="{style_num_id}"]'
    )[0]
    abstract_id = base_number.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [
        int(element.get(qn("w:numId"))) for element in numbering.findall(qn("w:num"))
    ]
    list_num_id = max(existing_ids, default=0) + 1
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(list_num_id))
    abstract_number = OxmlElement("w:abstractNumId")
    abstract_number.set(qn("w:val"), abstract_id)
    number.append(abstract_number)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    number.append(level_override)
    numbering.append(number)

    for item in items:
        paragraph = document.add_paragraph(item, style="List Number")
        number_properties = paragraph._p.get_or_add_pPr().get_or_add_numPr()
        number_properties.get_or_add_ilvl().val = 0
        number_properties.get_or_add_numId().val = list_num_id


def add_callout(document: Document, title: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360])
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    row_properties.append(cannot_split)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.style = "Callout"
    lead = paragraph.add_run(title + " ")
    lead.bold = True
    lead.font.color.rgb = RGBColor.from_string(BLUE)
    paragraph.add_run(text)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_by_count = {
        2: [2700, 6660],
        3: [1800, 3900, 3660],
        4: [1500, 2500, 2680, 2680],
    }
    widths = widths_by_count.get(
        len(headers), [9360 // len(headers)] * len(headers)
    )
    widths[-1] += 9360 - sum(widths)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        row_properties = table.rows[-1]._tr.get_or_add_trPr()
        cannot_split = OxmlElement("w:cantSplit")
        row_properties.append(cannot_split)
        for index, value in enumerate(row):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_geometry(table, widths)


def finish_document(document: Document, path: Path) -> None:
    document.core_properties.title = path.stem
    document.core_properties.subject = "Delphi App Translation Studio documentation"
    document.core_properties.author = "TMartinDub"
    document.core_properties.keywords = (
        "Delphi, VCL, FireMonkey, FMX, localization, translation"
    )
    document.settings.update_fields_on_open = True
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def build_user_guide() -> Path:
    document = Document()
    setup_styles(document)
    title = "Delphi App Translation Studio - User Guide"
    add_cover(
        document,
        "User Guide",
        "Delphi App Translation Studio",
    )
    add_toc_section(document, title)

    document.add_heading("1. About This Guide", level=1)
    add_paragraphs(
        document,
        [
            "This guide is the complete product reference for Delphi App Translation Studio. It is written for a Delphi developer who may be new to localization, language packs, design-time components, or translation-provider APIs. It explains what each part does, why it is needed, what files are created, and what result to expect before it gives operating steps.",
            "For the shortest first-time path, use the separate Delphi App Translation Studio Setup Wizard Guide. The Wizard Guide leads one project from preparation through translation, component setup, deployment, and final verification. This User Guide remains the authoritative reference for the Studio's full seven-page workflow, catalog editing, providers, validation, export, integration, troubleshooting, security, and file locations.",
            "Delphi App Translation Studio is an offline-first Windows developer tool for adding localization to existing Delphi VCL and FireMonkey applications. It scans designer-authored forms and Delphi resourcestring declarations, stores the results in an editable development catalog, automatically translates unresolved text through Google Cloud Translation or DeepL, validates translation safety, and exports compact JSON language packs for the target application.",
            "The Studio itself is built with FireMonkey, but it works with both VCL and FMX target projects. The supported build targets are Windows Win32 and Win64. macOS, iOS, Android, Linux, C++Builder, and runtime cloud translation are outside the product scope.",
            "When the Studio starts, its landing screen offers two deliberate paths: Run Setup Wizard for a new localization setup, or Open Maintenance Studio for an existing catalog, language pack, validation, export, or integration task. This keeps first-time users out of a page whose purpose is not yet clear.",
        ],
    )
    add_callout(
        document,
        "Offline by design.",
        "Only the developer's Studio computer needs Internet access while creating translations. The finished target application reads local JSON language packs and never needs Internet access, a provider account, or an API key.",
    )

    document.add_heading("1.1 The Parts of the Localization System", level=2)
    add_table(
        document,
        ["Part", "What it is", "Why it exists"],
        [
            ["Translation Studio", "The developer tool used while preparing a project.", "Scans text, calls Google or DeepL, validates results, and builds deployable files."],
            ["Development catalog", "A detailed JSON working file under Localization\\Development.", "Preserves source text, translations, review state, origin, runtime classification, and locale settings between scans."],
            ["Runtime language pack", "A compact JSON file under Localization\\Languages, such as es-ES.json.", "Supplies translated text to the finished application without Internet access or an API key."],
            ["TDAT language manager component", "One nonvisual VCL or FMX component placed on the application's primary form.", "Loads packs, remembers the selected language, translates open forms, and supervises forms opened later."],
            ["Language selector", "A required user-facing language choice, normally the visual DAT combo-box linked to the manager.", "Shows installed languages and switches immediately. A connected Language menu may be used instead."],
            ["Component integration kit", "A generated, project-specific folder under the Studio export tree.", "Provides component source, packs, deployment script, manifest, and exact integration instructions without rewriting the target project."],
        ],
    )
    add_callout(
        document,
        "What 'manager' means in this guide.",
        "Manager is shorthand for TDATVCLLanguageManager or TDATFMXLanguageManager. It is a Delphi component installed on the DAT Localization Tool Palette page and placed once on the target application's primary form. It is not a separate program, Windows service, or online account.",
    )

    document.add_heading("1.2 What the Studio Changes - and Does Not Change", level=2)
    add_bullets(
        document,
        [
            "Scanning is read-only and does not alter the selected project.",
            "Saving creates project-local Localization\\Development catalog files.",
            "Export creates project-local Localization\\Languages runtime packs.",
            "Component Integration generates a setup kit under the Studio export folder and does not write to the selected target project.",
            "The developer places one TDAT language manager component on the primary form in Delphi. Delphi then makes its normal, visible form-resource and unit-reference changes, exactly as it does for any other design-time component.",
            "Automatic Source Integration remains an explicit advanced fallback with preview, verified backup, apply, and restore.",
        ],
    )
    add_callout(
        document,
        "Protect the original project.",
        "Use a disposable test copy for the first evaluation. A Git repository is helpful but is not required. If Git is unavailable, copy the complete project to a separately named backup folder, confirm that the copy opens and builds, and run the Studio only against the test copy.",
    )

    document.add_heading("2. Before You Begin", level=1)
    add_paragraphs(
        document,
        [
            "The published source is currently built and release-tested with RAD Studio 13 Florence. The code is intended for modern Delphi VCL and FMX Windows projects; an older RAD Studio release may require source or package adjustments and should be treated as unverified until it passes the included build and smoke tests.",
            "No Git installation is required to use the Studio. Git is one convenient safety check for developers who already use it; a verified folder copy or other normal backup system is equally acceptable.",
            "Build products are placed under bin\\<Platform>\\<Configuration>. Compiler units are placed under dcu. The project deliberately does not target mobile or macOS platforms. After final Wizard processing, the Finish page can build and deploy selected Win32/Win64 Debug/Release combinations automatically; leave the option clear when you want to build manually.",
        ],
    )
    add_table(
        document,
        ["Prepare", "Reason"],
        [
            ["A clean test copy of the target project", "Keeps the original application available if the evaluation is abandoned."],
            ["A successful Win32 or Win64 build", "Confirms the target project is healthy before localization work begins."],
            ["Text-based DFM or FMX resources", "The scanner reads designer text resources; binary legacy DFM files must first be converted through Delphi."],
            ["A Google Cloud Translation or DeepL API key", "Automatic translation runs on the developer's computer. The deployed application remains offline."],
            ["RAD Studio access for one manual component step", "Delphi itself installs the design package and places the manager component on the primary form."],
        ],
    )

    document.add_heading("2.1 Choose the Setup Path", level=2)
    add_table(
        document,
        ["Path", "Choose it when", "Where to begin"],
        [
            ["Setup Wizard - recommended", "This is the first project, or the developer wants the Studio to lead and verify the sequence.", "Select Run Setup Wizard on the landing screen and follow the separate Setup Wizard Guide."],
            ["Maintenance Studio", "The project is already integrated, or the developer needs direct catalog, validation, export, provider, or integration maintenance.", "Select Open Maintenance Studio on the landing screen, then choose the required page."],
            ["Manual Studio workflow", "The developer understands catalogs and wants direct access to individual Project, Scan, Translate, Validation, Export, Integration, and Provider Settings pages.", "Open Maintenance Studio and continue with Chapter 3 of this User Guide."],
            ["Automatic Source Integration - advanced", "The component approach is unsuitable and the developer explicitly accepts reviewed source edits.", "Read Chapter 11.3 and use only a protected project copy."],
        ],
    )

    document.add_heading("2.2 Build or Start the Studio", level=2)
    add_steps(
        document,
        [
            "To use an existing build, run DelphiAppTranslationStudio.exe from bin\\Win32\\Release or bin\\Win64\\Release. Debug builds are intended for diagnosis.",
            "To compile from source, open DelphiAppTranslationStudio.dproj in RAD Studio, select Win32 or Win64 and Debug or Release, then choose Project > Build DelphiAppTranslationStudio.",
            "Confirm that the build writes the executable under bin\\<Platform>\\<Configuration> and compiler units under dcu\\<Platform>\\<Configuration>.",
            "Start the executable and confirm that Delphi App Translation Studio opens maximized on its landing screen. The landing screen offers Run Setup Wizard for first-time use and Open Maintenance Studio for the manual workflow; the maintenance workflow rail is hidden until that choice is made.",
        ],
    )

    document.add_heading("2.3 Install the Localization Components", level=2)
    add_paragraphs(
        document,
        [
            "The Studio and the target application have different jobs. The Studio prepares the translation files. The target application needs a TDAT language manager component so it can load and apply those files at runtime. Component installation is normally performed once per RAD Studio installation, not once per translated project.",
            "The design package must be installed through RAD Studio's Component > Install Packages dialog. Do not use Component > Install Component, and do not try to install a .dpk source file. The correct file is the compiled Win32 Release design-time .bpl selected by the Studio.",
        ],
    )
    add_steps(
        document,
        [
            "If this is the first localization project, complete the Setup Wizard through its Delphi Component step. In the manual workflow, open the project and complete Scan, Translate, Validation, and Export, then open Integration and leave Component Integration selected.",
            "Choose Build Integration Plan. The Studio enables Show Design BPL for the detected VCL or FMX framework.",
            "Choose Show Design BPL. File Explorer selects DATLanguageManagerVCLDesign.bpl or DATLanguageManagerFMXDesign.bpl in the Studio's stable bin\\packages\\Win32\\Release folder.",
            "Start RAD Studio without opening the target form. Choose Component > Install Packages, choose Add, select the exact design .bpl shown by the Studio, and choose Open.",
            "Confirm DAT Language Manager VCL design-time package or DAT Language Manager FireMonkey design-time package is listed and checked, then choose OK. Open a form and confirm DAT Localization appears in the Tool Palette.",
            "Never choose Component > Install Component and never select a .dpk. A .dpk is package source; Delphi installs the compiled Win32 design .bpl through Install Packages.",
            "Return to the Studio and choose Generate Component Kit. The clean kit contains ComponentSource, JSON packs, deployment support, a manifest, and README instructions; it deliberately contains no BPLs or installer scripts.",
            "Normally linked applications compile the generated kit's ComponentSource units into the executable. The target does not need DAT runtime BPLs beside it.",
        ],
    )
    add_callout(
        document,
        "Delphi-controlled installation.",
        "The Studio does not copy BPLs into Delphi folders, write the BDS registry, or install packages automatically. RAD Studio registers the stable package from the Studio build tree through its normal Install Packages dialog. Generated per-project kits remain replaceable and are never package-registration locations.",
    )

    document.add_heading("3. The End-to-End Workflow", level=1)
    add_table(
        document,
        ["Step", "Purpose", "Primary output"],
        [
            ["1 Project", "Open and identify a Delphi project.", "Project profile"],
            ["2 Scan", "Extract designated designer text and resourcestrings.", "Scan result"],
            ["3 Translate", "Select a language and translate unresolved text automatically.", "Saved machine-translation catalog"],
            ["4 Validation", "Check completeness and structural safety.", "Issue list"],
            ["5 Export", "Create the compact offline pack.", "Runtime JSON"],
            ["6 Integration", "Generate a non-mutating component setup kit; optionally use advanced automatic integration.", "Component kit or advanced preview"],
            ["7 Provider Settings", "Configure and test Google Cloud Translation or DeepL.", "Secure provider configuration"],
        ],
    )

    document.add_heading("4. Open and Scan a Project", level=1)
    add_steps(
        document,
        [
            "Choose Project and select Open Delphi Project.",
            "Open a .dproj when available; a .dpr is also accepted.",
            "Review the detected framework, platforms, form count, and source count.",
            "Choose Scan. The Studio reads text DFM/FMX resources and resourcestring declarations.",
            "Review the entry count, elapsed milliseconds, breakdown, and result list.",
        ],
    )
    add_callout(
        document,
        "Binary forms.",
        "The scanner expects text DFM/FMX resources. If a legacy DFM is binary, use Delphi's normal form conversion facilities before scanning and keep a source-control copy.",
    )
    document.add_heading("4.1 Text that is collected", level=2)
    add_bullets(
        document,
        [
            "Captions, Text values, prompts, hints, headings, and designated list content.",
            "Memo and string-list content when it represents user-visible text.",
            "Delphi resourcestring symbols with stable unit-and-symbol keys.",
            "Multiline DFM/FMX string expressions and supported runtime UI assignments in Pascal, including Text, Caption, Header, Hint, TextPrompt, and Format templates.",
            "Nested utility or conversion projects are excluded when a separate DPROJ or DPR identifies them as a different application.",
            "Locale metadata such as date, time, number, and currency formatting is entered on the Languages page rather than guessed from UI text.",
        ],
    )

    document.add_heading("5. Create and Edit a Language Catalog", level=1)
    add_steps(
        document,
        [
            "Choose Translate after scanning.",
            "Select the source language, normally English (United States) [en-US].",
            "Select the target language from the built-in language list. The Studio records its locale code and native name.",
            "Review the automatically selected left-to-right or right-to-left direction.",
            "Review or edit the date, time, decimal, thousands, and currency fields. Blank fields are initially populated from Delphi TFormatSettings for the target locale.",
            "Choose Save.",
            "Click the blue catalog path to select the saved JSON file in File Explorer.",
        ],
    )
    add_paragraphs(
        document,
        [
            "The catalog is saved under Localization\\Development using the project name and target locale. Each target language has its own development catalog. Existing translations are preserved on later scans when the stable key and source text remain unchanged. Changed source text is flagged for review, and removed entries become obsolete rather than disappearing silently.",
        ],
    )

    document.add_heading("5.1 Translation statuses", level=2)
    add_table(
        document,
        ["Status", "Meaning"],
        [
            ["Needs translation", "No usable target text is present."],
            ["AI draft", "Legacy provenance retained when opening a catalog created by an earlier Studio build."],
            ["Machine translated", "DeepL or Google produced a draft that requires review."],
            ["Imported", "CSV supplied target text that requires review."],
            ["Edited", "A person changed or entered the target text."],
            ["Reviewed", "A person reviewed the target text."],
            ["Approved", "The entry is release-approved."],
            ["Source changed", "The source changed after a translation existed."],
            ["Excluded", "The entry is intentionally omitted."],
            ["Obsolete", "The source entry no longer exists in the latest scan."],
            ["Error", "The entry needs correction after a failed operation or check."],
        ],
    )

    document.add_heading("5.2 Automatic Google or DeepL workflow", level=2)
    add_steps(
        document,
        [
            "Open Provider Settings and select Google Cloud Translation or DeepL.",
            "Paste the provider API key into the masked field, choose secure Windows Credential Manager storage or session-only use, and select Replace / Save Key.",
            "Choose Test Connection. Resolve any authentication, billing, restriction, firewall, or quota error before translating a project.",
            "Return to Translate, select the target language, and choose Translate Automatically.",
            "Confirm the displayed provider and unresolved-entry count. The Studio sends only eligible unresolved source strings; Reviewed, Approved, Excluded, and Obsolete entries are not replaced.",
            "The Studio first reuses approved contextual translation memory and vetted UI terminology where available. Remaining entries are sent in bounded batches. DeepL receives inferred UI context; Google Basic v2 does not support a context field, so short or ambiguous Google results are flagged for focused review. The catalog is saved automatically.",
        ],
    )
    add_callout(
        document,
        "No extra software installation.",
        "Automatic translation is built into the Studio. It requires only a supported provider account and API key; no command-line AI agent, Python package, browser extension, or separate translation utility is required.",
    )

    document.add_heading("5.3 Safety, context, and focused review", level=2)
    add_bullets(
        document,
        [
            "Designer-property entries are applied automatically by the VCL or FMX runtime adapter.",
            "Pascal resourcestring entries require an explicit TranslateText(Key, Fallback) call at the intended code location. Mark Manual wiring confirmed only after adding and reviewing that call.",
            "Translation origin is independent of Draft, Reviewed, and Approved status. A provider result is never silently approved.",
            "Only Needs Translation, Source Changed, and Error entries are eligible. Existing Machine-translated, Imported, Edited, Reviewed, Approved, Excluded, and Obsolete work is protected unless the developer changes its status deliberately.",
            "Each scanned entry records form, component, class, property, UI role, semantic concept, contextual description, and context confidence. This distinguishes meanings such as media Play, game Play, command Schedule, the noun Schedule, weekday abbreviations, Save, Close, and media-playback timing.",
            "Vetted UI terminology is applied before a provider call where a supported target-language term is known, and it repairs unreviewed provider drafts when a definitive application term is available. Reviewed and Approved entries remain protected. Reviewed and Approved entries also form contextual translation memory for matching later entries.",
            "DeepL accepts the Studio's contextual description. Google Cloud Translation Basic v2 does not accept context or glossaries; the Studio uses local context, terminology, memory, and focused warnings around Google.",
            "Validation focuses on placeholders, accelerators, inconsistent repeated terms, source changes, and runtime wiring.",
            "Exact-source translations from other stable keys are suggestions only. The developer must explicitly accept one, and the accepted target remains Edited rather than inheriting approval.",
        ],
    )

    document.add_heading("6. Provider Settings and Secure Key Storage", level=1)
    add_paragraphs(
        document,
        [
            "Provider Settings supports DeepL API Free, DeepL API Pro, and Google Cloud Translation Basic v2. Provider accounts, billing, quotas, prices, and supported languages are controlled by the provider and may change.",
            "A remembered key is stored as a Windows Generic Credential for the current Windows user. The key is not written to the Delphi project, JSON catalogs, exported packs, source distribution, or Git repository. Session-only keys are cleared when the Studio closes.",
        ],
    )
    add_table(
        document,
        ["Control", "Behavior"],
        [
            ["Provider", "Select DeepL or Google Cloud Translation."],
            ["DeepL API plan", "Select API Free or API Pro so the Studio uses the matching endpoint."],
            ["API key", "Masked field for a new or replacement key."],
            ["Remember securely", "Stores the key as a Windows Generic Credential."],
            ["Session only", "Clear Remember before saving; the key is held only until shutdown."],
            ["Replace / Save Key", "Records the entered key using the selected storage choice."],
            ["Test Connection", "Sends a small English-to-Italian request."],
            ["Remove Key", "Deletes stored and session copies for the selected provider."],
            ["Timeout / batch", "Controls 5-300 second requests and 1-50 strings per request."],
        ],
    )

    document.add_heading("7. Obtain and Enter a DeepL API Key", level=1)
    add_callout(
        document,
        "Use a DeepL API account.",
        "A normal DeepL Translator subscription is not automatically a DeepL API plan. Confirm that the account includes API Free or API Pro.",
    )
    add_steps(
        document,
        [
            "Open the official DeepL developer site at https://developers.deepl.com/ and review current API account requirements.",
            "Create or sign in to a DeepL account and subscribe to DeepL API Free or DeepL API Pro as appropriate.",
            "Open the account's API Keys tab. Create a separate key for this Studio when the account interface permits multiple keys.",
            "Copy the key once and keep it out of source code, JSON catalogs, issue trackers, email, and screenshots.",
            "In the Studio, choose Provider Settings and select DeepL.",
            "Select API Free or API Pro to match the account. Free uses api-free.deepl.com; Pro uses api.deepl.com.",
            "Paste the key into the masked field.",
            "Leave Remember securely checked for Windows Credential Manager, or clear it for Use for This Session Only.",
            "Choose Replace / Save Key, then Test Connection.",
            "If the test fails, verify the plan selection, key status, account quota/billing, firewall, and target-language availability.",
        ],
    )
    add_paragraphs(
        document,
        [
            "DeepL key and authentication reference: https://developers.deepl.com/docs/getting-started/auth",
            "DeepL quickstart: https://developers.deepl.com/docs/getting-started/quickstart",
            "DeepL multiple-key guidance: https://developers.deepl.com/docs/multiple-api-keys",
        ],
    )

    document.add_heading("8. Obtain and Enter a Google API Key", level=1)
    add_callout(
        document,
        "Google Basic v2 only.",
        "The Studio uses Cloud Translation Basic v2 because it supports API-key authentication. Cloud Translation Advanced v3 uses different authentication and is not implemented.",
    )
    add_steps(
        document,
        [
            "Open Google Cloud Console at https://console.cloud.google.com/ and sign in with the account that will own billing and credentials.",
            "Use the project selector to create a dedicated project, such as Delphi App Translation Studio, or select an existing project whose billing and key lifecycle you control.",
            "Open Billing for that project and link an active billing account. Cloud Translation setup requires billing to be enabled even when usage might fall within a current no-charge allowance.",
            "Open APIs & Services, then Library. Search for Cloud Translation API, open it, and choose Enable. Confirm that the selected project is still the intended project.",
            "Open APIs & Services, then Credentials. Choose Create credentials, then API key. Create a standard API key; do not bind it to a service account for this Studio.",
            "Name the key clearly, for example Delphi App Translation Studio - <computer or developer>. Google requires at least one API restriction when a key is created in the console.",
            "Under API restrictions, choose Restrict key and select only Cloud Translation API (service translate.googleapis.com), then save. Do not leave the key usable with every Google API.",
            "Application restrictions are recommended by Google, but the available types are websites, server IP addresses, Android, and iOS. A native Windows desktop application has no matching signed-application restriction. Use an IP-address restriction only when the developer computer exits through a stable public IP that you control; do not select website, Android, or iOS restrictions for the Studio.",
            "Copy the displayed key string immediately and keep it out of source code, JSON, screenshots, email, issue trackers, and chat. The key ID or display name is not the usable key string.",
            "Review Cloud Translation pricing, quotas, and billing budget alerts before a large project. API keys associate requests with the project for quota and billing.",
            "In the Studio, choose Provider Settings and select Google Cloud Translation. The DeepL plan list becomes disabled because it does not apply.",
            "Paste the key into the masked field. Leave Remember securely checked to store it as a Windows Generic Credential, or clear the check box for session-only use.",
            "Choose Replace / Save Key, then Test Connection. A successful test translates one small English phrase to Italian.",
            "If the test returns HTTP 403, verify the selected project, billing, API enablement, and API/application restrictions. HTTP 429 normally indicates a quota or rate limit. After changing a Google restriction, allow a few minutes for propagation and test again.",
        ],
    )
    add_paragraphs(
        document,
        [
            "Google Cloud Translation setup: https://docs.cloud.google.com/translate/docs/setup",
            "Google API-key creation and management: https://docs.cloud.google.com/docs/authentication/api-keys",
            "Google API-key restrictions: https://docs.cloud.google.com/api-keys/docs/add-restrictions-api-keys",
            "Cloud Translation authentication: https://docs.cloud.google.com/translate/docs/authentication",
        ],
    )

    document.add_heading("9. Direct Provider Translation", level=1)
    add_steps(
        document,
        [
            "Open or create the target-language catalog.",
            "Confirm the source and target language codes.",
            "Configure and test a provider under Provider Settings.",
            "Choose Translate Automatically on the Translate page.",
            "Read the confirmation message showing how many unresolved strings will be sent to the provider. Cross-key suggestions are never accepted automatically.",
            "Choose Yes to begin. Existing complete reviewed or approved work is preserved.",
            "Wait for all batches. Transient HTTP 429 and provider server failures receive bounded retries.",
            "Review every machine-translated entry for meaning, placeholders, accelerators, product terminology, tone, and available control space.",
            "The Studio saves the catalog automatically. Run Validation after reviewing the results.",
            "For focused work, Review and Approve the selected entry. For a large catalog, Review All and Approve All provide separate, confirmed catalog-wide decisions and save once after each operation.",
        ],
    )
    add_callout(
        document,
        "Cost control.",
        "The confirmation count is not a price quote. Provider billing can depend on characters and account terms. Check the provider dashboard and quotas before a large request.",
    )

    document.add_heading("10. Validate and Export", level=1)
    add_paragraphs(
        document,
        [
            "Structural validation checks catalog metadata, required translations, duplicate keys, source changes, Delphi indexed/sequential placeholders, accelerator keys, and excluded/obsolete conditions. Errors block export. Manual resourcestring wiring is reported as a runtime-readiness warning. Linguistic Reviewed and Approved states are reported separately and are not inferred from structural validation.",
        ],
    )
    add_steps(
        document,
        [
            "Choose Validation and Run Validation.",
            "Read the summary: errors block export, warnings request review, and information messages require no action. Double-click an entry-specific issue to open that catalog entry on Translate.",
            "Repeat until no errors remain.",
            "Choose Export and Export Runtime Pack.",
            "Record the output path under Localization\\Languages.",
        ],
    )

    document.add_heading("11. Integrate a Target Application", level=1)
    add_paragraphs(
        document,
        [
            "Component Integration is the recommended path. It generates validated JSON packs, an English source pack, the applicable component/runtime source, a machine-readable manifest, deployment script, and exact setup instructions. The Studio writes these files only under export\\component-integration and never opens the selected target project for writing.",
            "The target application remains completely offline. Provider code and API keys are never added to it. One manager on the primary form supervises the application; ordinary forms do not need individual components.",
        ],
    )
    add_steps(
        document,
        [
            "Select Integration and leave Integration method set to Component Integration (Recommended).",
            "Choose Build Integration Plan. Confirm the detected framework and translated-pack count.",
            "Choose Generate Component Kit. Select README.txt and the generated manifest/files in the read-only inspection panes.",
            "If DAT Localization is not already in the Tool Palette, choose Show Design BPL. It selects the stable package under the Studio's bin\\packages\\Win32\\Release folder. In RAD Studio choose Component > Install Packages > Add, select that exact design BPL, confirm the package is listed and checked, and choose OK.",
            "Open the target application's primary form in the Delphi Form Designer and place one TDATVCLLanguageManager or TDATFMXLanguageManager from DAT Localization.",
            "Set ApplicationId exactly to the detected Delphi project name. Leave LanguagesFolder as Localization\\Languages and set SourceLanguage to the catalog source locale, normally en-US.",
            "Place TDATVCLLanguageComboBox or TDATFMXLanguageComboBox and set its LanguageManager property to the manager. A connected Language menu is the supported alternative; some visible selector is required.",
            "The Setup Wizard adds the kit's ComponentSource folder to the DPROJ Base Search Path, which is inherited by Debug/Release and Win32/Win64.",
            "The Setup Wizard adds a post-build command that deploys the kit's Localization folder to the active DCC_ExeOutput directory with PowerShell ExecutionPolicy Bypass.",
            "Build and run the target. The saved language is applied during startup. Choosing another locale immediately retranslates open forms and saves the preference.",
        ],
    )

    document.add_heading("11.1 Component properties and behavior", level=2)
    add_table(
        document,
        ["Property or control", "Purpose"],
        [
            ["ApplicationId", "Must match runtime-pack applicationId; normally the Delphi project name."],
            ["LanguagesFolder", "Pack folder relative to the executable unless an absolute path is supplied."],
            ["SourceLanguage", "Source locale used when no translated pack is selected."],
            ["AutoLoadPreferred", "Loads the saved per-user locale during manager initialization."],
            ["AutoTranslateOwner / AutoTranslateNewForms", "Applies the active pack to the primary and later forms."],
            ["ReapplyOpenForms", "Immediately retranslates open forms after a language change."],
            ["PreserveControlState", "Protects writable user data, focus, selections, and list selection while text changes."],
            ["FormIdentityMappings", "Optional FormClass=ScannerRoot mappings for renamed or inherited forms."],
            ["Language combo box", "Populates from validated packs and calls the linked manager without taking over OnChange."],
        ],
    )

    document.add_heading("11.2 Runtime files and preferences", level=2)
    add_bullets(
        document,
        [
            "Deploy JSON packs beside the target executable under Localization\\Languages.",
            "Deploy-LanguagePacks.ps1 and manual copying support custom executable output layouts.",
            "The selected language is stored in %LOCALAPPDATA%\\<ApplicationId>\\language.ini.",
            "The executable folder can therefore remain read-only, as it normally is under Program Files.",
            "FMX uses additive before-show/release messages. VCL uses additive application idle/modal notifications and direct owner application.",
            "Selecting or reselecting a language applies the selected pack to every open form immediately. English is a real generated pack, so returning from another language restores all scanned source strings without restarting.",
        ],
    )

    document.add_heading("11.3 Automatic Source Integration (Advanced)", level=2)
    add_paragraphs(
        document,
        [
            "Use the advanced mode only when the component path is unsuitable and after committing or otherwise protecting the target project. It retains the previous generated-unit, menu-resource, preview, verified backup, atomic apply, build/deploy, restore, and Complete Reset workflow.",
        ],
    )
    add_steps(
        document,
        [
            "Change Integration method to Automatic Source Integration (Advanced).",
            "Enter the designer language-menu component name and build the plan.",
            "Generate the exact preview and inspect any file that needs closer review.",
            "Authorize once, optionally select build/deploy, and choose Apply.",
            "Use Restore for the session backup or Complete Reset for a recorded pre-integration baseline.",
        ],
    )

    document.add_heading("12. Translate the Studio Itself", level=1)
    add_steps(
        document,
        [
            "Run the English Studio and select DelphiAppTranslationStudio.dproj.",
            "Scan it and create the desired language catalog.",
            "Translate, review, validate, and export the pack.",
            "Copy the exported pack to the Studio project or deployed executable's Localization\\Languages folder.",
            "Add the locale to the designer-authored Studio Language menu when a selectable menu item is required.",
            "Close the running Studio, rebuild when the menu changed, and start the executable.",
            "Select the new language. Open forms change immediately and the preference is retained for the next launch.",
        ],
    )
    add_paragraphs(
        document,
        [
            "The running executable is never overwritten. The JSON pack and preference are separate from the executable, and the rebuild supplies the updated menu on the next run.",
        ],
    )

    document.add_heading("13. Troubleshooting", level=1)
    add_table(
        document,
        ["Symptom", "Likely action"],
        [
            ["Project framework unknown", "Open the .dproj, confirm VCL/FMX units and project metadata, then rescan."],
            ["No scan entries", "Confirm text DFM/FMX resources and designated user-visible properties."],
            ["HTTP 401/403", "Replace the key; verify provider plan, API enablement, restrictions, and billing."],
            ["HTTP 429", "Quota or rate limit reached; wait, reduce batch size, or review provider quota."],
            ["DeepL test fails only on one plan", "Select API Free or API Pro to match the account endpoint."],
            ["Google test fails", "Confirm Basic v2 Cloud Translation API is enabled and key API restriction permits it."],
            ["Export blocked", "Run Validation and correct every error."],
            ["Language selector is empty", "Confirm valid nonempty packs, matching applicationId, and the deployed Localization\\Languages folder; call RefreshLanguages after late deployment."],
            ["Component missing from Tool Palette", "Use Show Design BPL, then Delphi Component > Install Packages > Add. Select the exact self-contained Win32 design BPL."],
            ["Preference cannot be found", "Check %LOCALAPPDATA%\\<ApplicationId>\\language.ini, not the executable folder."],
            ["Target remains English", "Confirm manager ApplicationId, JSON applicationId/locale, deployment folder, preference, and manager initialization errors."],
        ],
    )

    document.add_heading("14. Security and Privacy", level=1)
    add_bullets(
        document,
        [
            "Only confirmed source strings are sent to the selected provider.",
            "Do not send secrets, personal data, customer data, or regulated information as UI text.",
            "Remembered keys are Windows Generic Credentials; session keys are not persisted.",
            "Removing a key affects only the selected provider.",
            "Provider settings JSON contains no API key.",
            "Catalogs, runtime packs, deployment packages, logs, backups, and Git should contain no key.",
            "Rotate a key immediately in the provider console if it may have been exposed, then replace it in the Studio.",
        ],
    )

    document.add_heading("15. File and Folder Reference", level=1)
    add_table(
        document,
        ["Location", "Contents"],
        [
            ["Localization\\Development", "Editable full development catalogs."],
            ["Localization\\Languages", "Compact offline runtime JSON packs."],
            ["export", "Generated previews and temporary product output."],
            ["export\\component-integration", "Recommended non-mutating component setup kits."],
            ["packages\\runtime / packages\\design", "Core/framework runtime packages and VCL/FMX IDE packages."],
            ["docs\\guides / docs\\pdf", "Editable guides and companion PDFs."],
            ["%LOCALAPPDATA%\\DelphiAppTranslationStudio", "Studio settings and language preference, but no remembered key."],
            ["Windows Credential Manager", "Remembered DeepL/Google Generic Credentials."],
            ["%LOCALAPPDATA%\\<ApplicationId>", "Integrated application's selected-language preference."],
        ],
    )

    document.add_heading("16. Provider Reference and Change Notice", level=1)
    add_paragraphs(
        document,
        [
            "Provider setup screens and commercial terms can change after this guide is published. Before creating a production key, compare these steps with the official pages listed in Chapters 7 and 8. Prefer a dedicated, restricted key and review the provider dashboard after the first bulk run.",
            "This guide documents the implemented Studio as of August 10, 2026.",
        ],
    )

    path = GUIDES_DIR / "Delphi App Translation Studio User Guide.docx"
    finish_document(document, path)
    return path


def build_setup_wizard_guide() -> Path:
    document = Document()
    setup_styles(document)
    title = "Delphi App Translation Studio - Setup Wizard Guide"
    add_cover(
        document,
        "Setup Wizard Guide",
        "A First-Time User's Path from Delphi Project to Offline Translation",
    )
    add_toc_section(document, title, [
        ("1. Purpose and Expected Result", 1),
        ("2. Complete the RAD Studio Preparation Before Starting", 1),
        ("3. Start and Navigate the Setup Wizard", 3),
        ("4. Complete the Nine Wizard Steps", 4),
        ("5. Verify the RAD Studio Component Setup", 9),
        ("6. Build and Deploy the Language Packs", 10),
        ("7. Verify the Translated Application", 10),
        ("8. Files Created by the Wizard", 11),
        ("9. Update, Resume, or Add Another Language", 12),
        ("10. Troubleshooting", 13),
        ("11. Where to Learn More", 15),
    ])

    document.add_heading("1. Purpose and Expected Result", level=1)
    add_paragraphs(
        document,
        [
            "The Setup Wizard is the recommended way to localize a Delphi application for the first time. It presents one decision at a time, verifies prerequisites before allowing the next step, and records the files and commands needed to finish the integration.",
            "At completion, the developer has a saved development catalog, compact JSON runtime language packs, a component integration kit, a required safety backup, and automatic build deployment. Pascal, DFM, FMX, and DPR files remain untouched. The Wizard adds one clearly marked, backed-up block to the DPROJ for the ComponentSource Search Path and post-build language-pack deployment. The Finish page also offers an optional build-and-deploy action so the selected Win32/Win64 Debug/Release outputs can be produced without leaving the Wizard.",
            "Before opening the Wizard, RAD Studio must install the design package through its normal Install Packages dialog. The developer then places and configures the language manager and visible selector, adds any supporting labels or other localization UI, and saves those changes. The Wizard consequently scans the completed interface once instead of requiring a scan, a component-related UI change, and a second scan.",
        ],
    )
    add_callout(
        document,
        "End-user applications remain offline.",
        "Google Cloud Translation or DeepL is contacted only by the Translation Studio on the developer's computer. The finished application reads local JSON files and requires no Internet connection, provider account, or API key.",
    )

    document.add_heading("2. Complete the RAD Studio Preparation Before Starting", level=1)
    add_table(
        document,
        ["Requirement", "What to prepare", "Why it matters"],
        [
            ["Target project", "A .dproj file is preferred; .dpr is accepted.", "The Wizard identifies VCL or FMX, supported Windows platforms, and form resources."],
            ["Safe working copy", "Use a clearly named test copy and confirm it builds before localization.", "The original application remains available even if the evaluation is abandoned."],
            ["Backup method", "Git is optional. A verified folder copy, archive, or normal backup system is sufficient.", "The Wizard does not assume every Delphi developer uses Git."],
            ["Translation provider", "A working Google Cloud Translation Basic v2 or DeepL API key.", "Automatic translation cannot run without provider authentication."],
            ["RAD Studio", "The current release is verified with RAD Studio 13 Florence.", "Install the matching design package and finish the component-related UI before opening the Wizard."],
            ["Healthy target build", "Build at least one intended Windows configuration successfully.", "A localization test should not begin with unrelated compiler or form-streaming failures."],
        ],
    )
    add_steps(
        document,
        [
            "Close the target application if it is running.",
            "Create the pristine or otherwise protected copy, then create a separate test copy for the Wizard.",
            "Open the untouched test copy in RAD Studio, build it, and run it briefly to confirm that it is healthy before localization changes begin.",
        ],
    )

    document.add_heading("2.1 Install the Localization Components", level=2)
    add_paragraphs(
        document,
        [
            "Complete this normally one-time RAD Studio installation before opening the Wizard. If DAT Localization is already visible in the Tool Palette for the target framework, do not reinstall it; continue with Section 2.2.",
        ],
    )
    add_steps(
        document,
        [
            "Close the target project while changing installed packages.",
            "In RAD Studio choose Component > Install Packages. Do not choose Component > Install Component.",
            "Select Add. From the Translation Studio installation, open bin\\packages\\Win32\\Release and select DATLanguageManagerFMXDesign.bpl for a FireMonkey project or DATLanguageManagerVCLDesign.bpl for a VCL project.",
            "Confirm that the matching DAT Language Manager design-time package is listed and checked, then select OK.",
            "Open a matching FMX or VCL form and confirm that DAT Localization appears in the Tool Palette with the language manager and language combo-box components.",
        ],
    )
    add_callout(
        document,
        "Install the compiled design package.",
        "Install the stable Win32 Release design-time BPL. Do not select a .dpk source file, do not copy BPLs into Delphi folders, and do not install a package from a generated per-project kit.",
    )

    document.add_heading("2.2 Complete and Save the Localization UI", level=2)
    add_paragraphs(
        document,
        [
            "Finish every interface change caused by adding language selection before the Wizard performs its scan. Text added here, including a Language: label, is then included in the first scan and does not require an immediate rescan.",
        ],
    )
    add_steps(
        document,
        [
            "Open the target test project and its primary form in the Form Designer.",
            "Place one TDATFMXLanguageManager for an FMX project or one TDATVCLLanguageManager for a VCL project. One manager supervises the application; do not place one on every form.",
            "Place the matching TDAT language combo box. In Object Inspector, set its LanguageManager property to the manager component; do not leave this property blank.",
            "Set ApplicationId to the project filename without .dproj. For example, MyApplication.dproj uses MyApplication. Leave LanguagesFolder as Localization\\Languages and set SourceLanguage to the application's source locale, such as en-US.",
            "Add any visible supporting interface, such as a Language: label or a designer-authored language menu. Complete its wording, placement, size, font, color, alignment, and related layout changes now.",
            "Choose File > Save All. The Wizard reads saved DFM/FMX and Pascal files; it cannot scan an unsaved Form Designer buffer.",
            "Close the target project before opening the Translation Studio. Do not worry if the newly added component units cannot compile yet: final Wizard processing configures the generated ComponentSource Search Path before the localized build.",
            "Have the provider API key available. Do not paste it into source code, JSON files, email, screenshots, or issue reports.",
        ],
    )
    add_callout(
        document,
        "Why these steps come first.",
        "The first Wizard scan now sees the manager, selector, Language: label, and every related saved UI change. Ordinarily, only one scan is needed for initial setup. Any later source-text or Form Designer change must still be saved and rescanned before export.",
    )

    document.add_heading("3. Start and Navigate the Setup Wizard", level=1)
    add_steps(
        document,
        [
            "Run DelphiAppTranslationStudio.exe from the selected Win32 or Win64 Debug or Release folder.",
            "On the landing screen, select Run Setup Wizard. The separate Open Maintenance Studio path begins the advanced manual workflow and is not required for this guide.",
            "Read the Welcome page, especially the orange safety notice, then select Next.",
        ],
    )
    add_table(
        document,
        ["Navigation control", "Behavior"],
        [
            ["Next", "Validates the current step and moves forward only when required information is present."],
            ["Back", "Returns to the previous step before final processing."],
            ["Left step rail", "Selects a step already reached. Future steps remain unavailable, preventing required work from being skipped."],
            ["Cancel", "Closes the Wizard before final processing and returns to the Studio landing screen. The target project remains unchanged; a deliberately saved provider credential may remain in Windows Credential Manager."],
            ["Begin Final Processing", "Starts the controlled translation/export sequence. Back, Cancel, the left rail, and window closing are disabled until the sequence completes or stops safely."],
            ["Finish", "Becomes available only after final processing succeeds."],
        ],
    )
    add_callout(
        document,
        "Cancel boundary.",
        "Cancel is safe through the Review and Authorize page, before Begin Final Processing is selected. Cancel returns to the landing screen, where the developer may start the Wizard again or select Maintenance Studio. During final processing the Wizard cannot be closed. If a failure occurs after the controlled DPROJ update, the transaction restores the prior project configuration automatically. Pascal and form source are not rewritten.",
    )

    document.add_heading("4. Complete the Nine Wizard Steps", level=1)

    document.add_heading("4.1 Step 1 - Welcome", level=2)
    add_paragraphs(
        document,
        [
            "The Welcome page summarizes the complete job and the safety boundary. No project has been selected and no target file is being written at this point.",
        ],
    )
    add_steps(document, ["Read the safety notice.", "Select Next."])
    add_callout(document, "Expected result.", "The Delphi Project step opens and Step 1 remains selectable in the left rail.")

    document.add_heading("4.2 Step 2 - Delphi Project", level=2)
    add_paragraphs(
        document,
        [
            "This step identifies the application that owns the text. Browse does not compile, execute, or modify the project. A .dproj file usually supplies the best framework and platform metadata; use .dpr only when no .dproj is available.",
        ],
    )
    add_steps(
        document,
        [
            "Select Browse.",
            "Navigate to the test copy and select its .dproj file.",
            "Confirm the displayed project name, VCL or FireMonkey framework, Windows targets, and form-resource count.",
            "Confirm the detected Application ID. It is the project filename without .dproj; for example, MyApplication.dproj produces MyApplication. Use Copy ID if the value will be needed in RAD Studio.",
            "If the project is wrong, select Browse again before continuing.",
            "Select Next.",
        ],
    )
    add_callout(document, "Expected result.", "The footer reports that the project was identified and no target file was changed. The exact Application ID is displayed explicitly rather than inferred from the folder name.")

    document.add_heading("4.3 Step 3 - Deployment Destinations", level=2)
    add_paragraphs(
        document,
        [
            "The Wizard detects normal Win32 and Win64 build-output folders automatically. This page is for additional installed, portable, network, or USB folders that contain—or will contain—the application executable. Entries are staged until authorized final processing, so Cancel still leaves the project and Studio-owned deployment settings unchanged.",
        ],
    )
    add_steps(
        document,
        [
            "Leave the list empty when the application will run only from normal Delphi build-output folders.",
            "For a separate application copy, select Add Application Folder and choose the full folder, including its drive letter. Select the folder containing the executable, not the Localization or Languages subfolder.",
            "Repeat Add Application Folder for every installed, portable, network, or USB destination that should receive language packs automatically.",
            "Use Remove Selected to remove an obsolete or incorrect destination before final processing.",
            "Read the summary, then select Next.",
        ],
    )
    add_callout(document, "Expected result.", "Final processing deploys immediately to every configured destination that is currently available. The choices are saved as Studio-owned JSON and included in the component kit. Future Delphi builds retry them automatically; an unplugged or unavailable drive is skipped with a warning and does not fail the build.")

    document.add_heading("4.4 Step 4 - Languages", level=2)
    add_paragraphs(
        document,
        [
            "Source language means the language currently written in the forms and source strings. Target language means the language the application user will select. The Wizard uses the locale code to name the JSON pack and to obtain date, time, number, currency, and text-direction defaults.",
        ],
    )
    add_steps(
        document,
        [
            "Leave Source language as English (United States) [en-US] when the application was authored in U.S. English.",
            "Choose the target language from the list.",
            "Confirm the native language name. This is the human-readable name displayed to the application's users.",
            "Read the locale summary, then select Next.",
        ],
    )
    add_callout(document, "Expected result.", "A locale such as es-ES is selected. Choosing a different target language invalidates downstream scan/catalog state so the wrong language cannot be exported accidentally.")

    document.add_heading("4.5 Step 5 - Translation Service", level=2)
    add_paragraphs(
        document,
        [
            "The Wizard supports Google Cloud Translation Basic v2 and DeepL. A remembered key is stored as a Windows Generic Credential for the current Windows user. A session-only key is held in memory until the Studio closes. Keys are never written into catalogs, runtime packs, component kits, or target source files.",
        ],
    )
    add_steps(
        document,
        [
            "Select Google Cloud Translation or DeepL. For DeepL, also select API Free or API Pro to match the account.",
            "If a key is already stored for this provider, the status reports that it is available. Otherwise paste a new key into the masked API key field.",
            "Leave Remember securely on this computer selected for Windows Credential Manager storage, or clear it for this Studio session only.",
            "Select Save / Replace Key when a new key was entered.",
            "Select Test Connection and wait for Connection test passed. Do not continue after an authentication, billing, quota, restriction, or network failure.",
            "Select Next.",
        ],
    )
    add_callout(document, "Expected result.", "The provider connection succeeds using one small test translation. No target application text has been sent yet.")

    document.add_heading("4.6 Step 6 - Scan Project", level=2)
    add_paragraphs(
        document,
        [
            "Scanning reads text DFM/FMX resources and Delphi resourcestring declarations. It creates stable keys such as form.component.property and unit.symbol. If a development catalog already exists for the selected project and language, matching translations are preserved; only new, changed, or unresolved entries require provider work.",
            "The scan count is not the same as the translation count. Dynamic values, identifiers, excluded entries, obsolete entries, and already complete translations may be scanned but deliberately withheld from automatic translation.",
        ],
    )
    add_steps(
        document,
        [
            "Select Scan Project.",
            "Review the summary for total, new, changed, unchanged, and obsolete entries.",
            "Scroll through representative rows and confirm that keys and source text belong to the selected project.",
            "If the count is zero, stop and verify that the forms are text resources and that the correct project was selected.",
            "Select Next.",
        ],
    )
    add_callout(document, "Expected result.", "The footer reports that the development catalog is ready. The scan itself remains read-only with respect to the target project.")

    document.add_heading("4.7 Step 7 - Delphi Component", level=2)
    add_paragraphs(
        document,
        [
            "This page restates the component requirements completed before the Wizard was opened and provides the exact project-specific values for verification. No additional component or label should normally be required now.",
        ],
    )
    add_steps(
        document,
        [
            "Read every project-specific instruction shown on the page and compare it with the component setup already saved in RAD Studio.",
            "Confirm that the displayed ApplicationId matches the value assigned to the manager. The project filename without .dproj is the controlling value.",
            "Use Show Design BPL only if the DAT Localization components were not installed during Section 2.1 or the package location must be verified. File Explorer selects the package; the Wizard does not install it.",
            "Select I understand this manual RAD Studio step to confirm that the pre-Wizard component setup has been completed and will be verified after final processing.",
            "Select Next.",
        ],
    )
    add_callout(document, "Expected result.", "The Wizard records the acknowledgement. It does not modify the package registry or the target form; those designer-owned changes were already completed and saved before the scan.")

    document.add_heading("4.8 Step 8 - Review and Authorize", level=2)
    add_paragraphs(
        document,
        [
            "This is the final cancellation point. The summary identifies the exact project, framework, target language, provider, scan count, unresolved translation count, component integration method, and backup choice. Read the project path character by character, especially when pristine and test folders have similar names.",
        ],
    )
    add_steps(
        document,
        [
            "Confirm the project path and application name.",
            "Confirm the target language and provider.",
            "Compare Scanned entries with Unresolved entries to translate. A lower unresolved count is normal when entries are protected, excluded, obsolete, or already translated.",
            "Close the target project in RAD Studio, then select Required: the target project is closed in RAD Studio. This prevents the IDE from overwriting the controlled DPROJ update with an older in-memory copy.",
            "Confirm that the required timestamped ZIP safety backup is selected. It cannot be disabled because final processing adds controlled properties inside Delphi's native Base compiler group.",
            "Select I reviewed these choices and authorize final processing.",
            "Select Begin Final Processing only when every item is correct.",
        ],
    )
    add_callout(document, "Expected result.", "The Wizard moves to Process and Finish, disables cancellation and navigation, and begins the logged operation sequence.")

    document.add_heading("4.9 Step 9 - Process and Finish", level=2)
    add_paragraphs(
        document,
        [
            "Final processing runs in a protected single pass: required ZIP backup, automatic translation of eligible unresolved entries, initial catalog save, automatic Localization Review, application of saved terminology/layout decisions, final validation, runtime-pack export, component-kit generation, deployment to existing output folders and every available configured application destination, backed-up transactional DPROJ configuration, and completion-report creation.",
            "Machine translations are recorded with Google or DeepL provenance. Existing Reviewed, Approved, Excluded, and Obsolete entries are not overwritten. Blocking validation errors stop the sequence; warnings are recorded but do not necessarily prevent export.",
        ],
    )
    add_steps(
        document,
        [
            "Watch the timestamped progress log. Do not terminate the Studio or Windows while processing is active.",
            "When processing completes, select the blue component-kit path or Open Kit Folder.",
            "Read the repeat/troubleshooting explanation. Normally, neither deployment button must be clicked: final processing has already deployed detected outputs and configured application destinations.",
            "Build each target configuration you actually intend to ship or test. A Delphi build produces one selected platform/configuration at a time unless you use Project > Build All Projects with the required configurations enabled. The Wizard's post-build step automatically deploys the current JSON packs to that build's executable Localization\\Languages folder.",
            "Use Redeploy Build Outputs only to repeat deployment after an unusual manual change. Use Deploy New App Folder only for a new or temporary folder that was not entered on Step 3. The Wizard verifies the EXE before copying anything.",
            "Use Copy Fallback Commands only for troubleshooting or a deliberately manual deployment.",
            "In the Build the application now? card, optionally select Build now, choose Win32 only, Win64 only, or Win32 and Win64, choose Debug only, Release only, or Debug and Release, and select Build and Deploy Selected Targets. The progress message reports each target and deployment result. Leave Build now clear if you want to finish without compiling.",
            "If a selected build fails, correct the Delphi project or environment, run the Wizard again, and repeat the build choice. A failed build does not invalidate the already-created catalog, pack, backup, or kit.",
            "Select Finish after reviewing the completion results. The Wizard closes and returns to the Studio; continue with the verification steps in Chapter 5.",
        ],
    )
    add_callout(document, "Expected result.", "The progress log ends successfully, the component-kit folder exists, Wizard-Completion-Report.txt records the exact Application ID, Search Path, backups, and deployments, and the DPROJ contains one marked Wizard block.")

    document.add_heading("5. Verify the RAD Studio Component Setup", level=1)
    add_paragraphs(
        document,
        [
            "The package and component-related UI were completed before the Wizard so the first scan included all visible text. After final processing, reopen the target project and verify the saved component properties and the Search Path added by the Wizard. Do not place duplicate components.",
        ],
    )
    document.add_heading("5.1 Verify the Design Package", level=2)
    add_steps(
        document,
        [
            "Start RAD Studio and confirm that the matching DAT Language Manager design-time package remains enabled under Component > Install Packages.",
            "Open the target project and its primary form. Confirm that DAT Localization remains visible in the Tool Palette and that Delphi loads the saved form without a missing-class error.",
            "If the package is unexpectedly absent, close the project and repeat Section 2.1 using the stable Win32 Release design BPL. Do not ignore a missing component class while opening the form.",
        ],
    )
    add_callout(
        document,
        "No reinstall is normally needed.",
        "A successfully installed DAT package remains registered for the current RAD Studio installation. Repeat package installation only when it is missing or disabled.",
    )

    document.add_heading("5.2 Verify the Saved Components", level=2)
    add_steps(
        document,
        [
            "Confirm that the primary form contains exactly one DAT language manager and the intended visible selector.",
            "Confirm that ApplicationId exactly matches the Wizard's detected value, LanguagesFolder is Localization\\Languages unless intentionally customized, and SourceLanguage matches the authored source locale.",
            "Confirm that the selector's LanguageManager property references the manager and that every supporting label, menu item, or other language-selection UI is present and correctly positioned.",
            "If any visible text or Form Designer layout must be changed now, choose File > Save All and run the Wizard scan again before relying on the runtime pack. A post-scan UI change cannot be added safely to the existing export without rescanning.",
            "If verification requires no UI or source-text change, continue directly to Search Path verification and the localized build.",
        ],
    )

    document.add_heading("5.3 Verify the Automatic Search Path", level=2)
    add_paragraphs(
        document,
        [
            "The Wizard adds the generated kit's ComponentSource directory under the DPROJ Base configuration. Base inheritance makes the same path available to Debug and Release on both Win32 and Win64. No four-way manual entry is normally required.",
            "To verify or repair the setting manually in RAD Studio, open Project > Options. Select Building > Delphi Compiler. At the top of the page select All configurations and All platforms, then locate Search path. The value must contain the exact ComponentSource path recorded in Wizard-Completion-Report.txt. Do not delete existing paths or the inherited $(DCC_UnitSearchPath) value.",
        ],
    )
    add_callout(
        document,
        "Manual fallback screen reference.",
        "RAD Studio path: Project > Options > Building > Delphi Compiler > Search path. Select All configurations and All platforms before editing. The Wizard normally performs this operation automatically; this reference exists for verification and recovery.",
    )

    document.add_heading("6. Build and Deploy the Language Packs", level=1)
    add_paragraphs(
        document,
        [
            "The JSON packs must be beside each executable under Localization\\Languages. The Wizard configures a post-build event that uses Delphi's active DCC_ExeOutput value, so Debug, Release, Win32, Win64, and intentional custom output paths are handled by the build that produced the executable. If the Finish-page build action was used, this deployment has already occurred for each successful selected target.",
        ],
    )
    add_steps(
        document,
        [
            "Build each target platform/configuration that will be tested or released.",
            "Confirm each successful build reports language-pack deployment in its build output. The command runs PowerShell with -NoProfile -ExecutionPolicy Bypass.",
            "For each executable, verify that Localization\\Languages contains the English source pack and the translated target pack.",
            "If an already-built output needs refreshing, select Redeploy Build Outputs. Copy Fallback Commands remains available for troubleshooting.",
            "For a portable or USB copy, enter its full application folder on Step 3 so final processing and future builds deploy it automatically. Deploy New App Folder remains available for a new or temporary destination. The LanguagesFolder component property remains the relative value Localization\\Languages, while the physical destination includes the drive.",
        ],
    )
    add_callout(
        document,
        "PowerShell execution policy.",
        "Generated commands launch C:\\Windows\\System32\\WindowsPowerShell\\v1.0\\powershell.exe with -NoProfile -ExecutionPolicy Bypass. This bypass applies only to that launched process and prevents the deployment script from failing under the common local script-policy restriction.",
    )

    document.add_heading("7. Verify the Translated Application", level=1)
    add_steps(
        document,
        [
            "Run the target executable from one deployed build folder.",
            "Confirm that the language selector lists English and the translated language.",
            "Choose the translated language. Open forms should update immediately without restarting the application.",
            "Open secondary forms and pages. Confirm captions, labels, prompts, column headings, and other scanned static text use the selected language.",
            "Confirm that live data, identifiers, dates, numeric values, selections, focused controls, and application connection state remain intact.",
            "Close and restart the application. Confirm that the selected language is remembered.",
            "Switch back to English and confirm that the English source pack restores the source text.",
            "Repeat the test for every built Win32 and Win64 configuration intended for release.",
        ],
    )
    add_callout(
        document,
        "Translation length is application-specific.",
        "A correct translation can be substantially longer than its English source—20 to 50 percent growth is common for short interface phrases, and individual terms can grow more. The Review Center proposes conservative language-specific Width, Height, WordWrap, AutoSize, and grid-column rules and shows current and proposed geometry in its Visual Review. Accepted rules are stored in the JSON pack, applied only for that language, and reversed when another language is selected. Complex collisions, graph geometry, and substantial redesigns still require developer review in the target Form Designer.",
    )

    document.add_heading("8. Files Created by the Wizard", level=1)
    add_table(
        document,
        ["Location or file", "Purpose"],
        [
            ["<Target>\\Localization\\Development\\<Project>.<locale>.translation-project.json", "Detailed working catalog retained for later rescans, updates, and review."],
            ["<Target>\\Localization\\Languages\\<locale>.json", "Compact runtime pack for deployment beside the application."],
            ["Documents\\Delphi App Translation Backups\\<Project>\\<timestamp>.zip", "Required pre-processing safety backup."],
            ["<Target>.dproj marked Wizard block", "Inherits ComponentSource through all configurations/platforms and runs pack deployment after builds."],
            ["<Studio>\\export\\component-integration\\<Project>", "Generated component integration kit."],
            ["ComponentSource", "Framework-neutral and VCL/FMX component/runtime units needed by the target."],
            ["Deploy-LanguagePacks.ps1", "Copies the kit's Localization folder to a specified executable directory."],
            ["component-integration.json", "Machine-readable integration manifest."],
            ["README.txt", "Project-specific manual integration instructions."],
            ["Wizard-Completion-Report.txt", "Final record of project, language, catalog, pack, backup, kit, manual step, and commands."],
        ],
    )

    document.add_heading("9. Update, Resume, or Add Another Language", level=1)
    document.add_heading("9.1 Update an Existing Translation After UI or Text Changes", level=2)
    add_paragraphs(
        document,
        [
            "Changing an existing application is normal. New labels, revised captions, renamed controls, added forms, and changed resourcestrings must be saved and rescanned before they can enter the JSON catalog. The Wizard reuses the existing catalog and preserves matching work; it does not start the language over.",
            "For the fewest provider calls and the clearest review, gather related interface changes into one practical batch. A single urgent correction can still be processed by itself.",
        ],
    )
    add_steps(
        document,
        [
            "Make the intended UI or source-text changes in Delphi, then choose File > Save All. The scanner reads saved DFM/FMX and Pascal files, not unsaved Form Designer buffers.",
            "Run the current Translation Studio and start the Setup Wizard. Select the same .dproj, source language, target language, and provider used for the existing catalog.",
            "Run Scan Project again. Read the new, changed, unchanged, and obsolete counts. A renamed component receives a new stable key; its former key remains Obsolete for traceability.",
            "Continue to final processing. Automatic translation sends only eligible new, changed, or otherwise unresolved entries. Matching Reviewed and Approved translations are preserved.",
            "If any target DFM/FMX or Pascal file is saved after the scan, final processing stops instead of exporting stale results. Return to Scan Project and scan again.",
            "The already installed design package, manager, selector, and configured search path normally remain in place. Review the Delphi Component step and acknowledge the existing setup; do not add duplicate components.",
            "Let final processing validate the catalog, export fresh JSON, regenerate the kit, and deploy to known build outputs and every available configured application destination.",
            "Rebuild the target because its UI/source changed. Test every affected form in the source and target languages, including text wrapping, alignment, control state, and language persistence.",
        ],
    )
    add_callout(
        document,
        "Repeatable update sequence.",
        "Batch changes -> Save All -> Scan -> Translate unresolved -> Validate -> Export -> Deploy -> Rebuild -> Test.",
    )

    document.add_heading("9.2 What the Wizard Repeats Automatically", level=2)
    add_bullets(
        document,
        [
            "Completed steps may be revisited from the left rail while the Wizard remains open and final processing is not running.",
            "Selecting a different project clears downstream scan/catalog state. Run the scan again.",
            "Selecting a different target language creates or opens that language's own development catalog and also requires a new scan in the Wizard session.",
            "A later scan preserves matching translations, marks changed source text for attention, adds new entries, and retains removed entries as obsolete.",
            "Automatic translation sends only eligible unresolved entries. Existing reviewed and approved work is preserved.",
            "Run the Wizard again when adding another language. Deploy all generated packs to each executable folder and refresh or restart the target application as required by its integration.",
            "The Wizard can rescan, merge, translate unresolved entries, validate, export, regenerate the component kit, configure supported project metadata, and deploy packs. It cannot save an open Delphi editor buffer, decide whether a translation is linguistically ideal, or redesign controls for longer text.",
        ],
    )

    document.add_heading("10. Troubleshooting", level=1)
    add_table(
        document,
        ["Problem", "What it means", "Action"],
        [
            ["Project browse or identification fails", "An old Wizard build may be running, or the selected file lacks recognizable VCL/FMX metadata.", "Use the current executable, select the .dproj rather than .dpr, and confirm the project opens normally in RAD Studio."],
            ["Connection test fails", "Provider authentication, billing, restriction, quota, plan, or network setup is incomplete.", "Correct the provider account before scanning or final processing."],
            ["Scan returns zero", "The wrong project was selected or the forms are not readable text resources.", "Verify the project path and convert binary DFM resources through Delphi."],
            ["Unresolved count is lower than scan count", "Some entries are dynamic, excluded, obsolete, protected, or already translated.", "This is normally correct; review classifications in the full Studio when needed."],
            ["Package or DAT Localization is unavailable", "The design package is not built, installed, or enabled in this RAD Studio installation.", "Build the package set if needed, then use Component > Install Packages > Add with the exact Win32 Release design BPL."],
            ["Packs are absent after a build", "The marked DPROJ settings are missing, the kit moved, the destination was unavailable, or the post-build command failed.", "Inspect the build output and configured destination, verify the kit path, then use Redeploy Build Outputs, Deploy New App Folder, or the documented manual fallback."],
            ["Selector is empty", "Packs are missing, invalid, or have a different applicationId.", "Check Localization\\Languages beside the running executable and verify ApplicationId."],
            ["Some text remains English", "The text may be dynamic, excluded, added after the scan, or not wired as a resourcestring.", "Rescan and translate new entries; use the full Studio to inspect runtime classifications and manual wiring warnings."],
            ["A newly added label is missing from the scan", "The form was not saved, the wrong project/catalog is open, or the new property is not an eligible text property.", "Choose File > Save All in Delphi, confirm the same .dproj and target language, then rescan and inspect the entry classification."],
            ["Final processing says source was saved after the scan", "The scan snapshot is stale and exporting it could omit a recent UI or source-text change.", "Choose Save All in Delphi, return to Scan Project, scan again, and then continue. Do not bypass this safeguard."],
            ["Translated text clips or overlaps", "The target layout was sized only for the shorter source language.", "In the Delphi Form Designer, enable appropriate wrapping/AutoSize and adjust layout containers, margins, or control dimensions; then rebuild and visually retest every supported language."],
        ],
    )

    document.add_heading("11. Where to Learn More", level=1)
    add_bullets(
        document,
        [
            "User Guide: full manual workflow, provider setup, catalog review, validation, export, integration, troubleshooting, security, and folder reference.",
            "Engineering Guide: architecture, JSON schema, scanners, provider clients, runtime managers, packages, lifecycle, security, and validation strategy.",
            "Generated kit README.txt: exact instructions for the selected target project and framework.",
            "Wizard-Completion-Report.txt: exact paths and commands produced by the completed Wizard run.",
        ],
    )

    path = GUIDES_DIR / "Delphi App Translation Studio Setup Wizard Guide.docx"
    finish_document(document, path)
    return path


def build_engineering_guide() -> Path:
    document = Document()
    setup_styles(document)
    title = "Delphi App Translation Studio - Engineering Guide"
    add_cover(
        document,
        "Engineering Guide",
        "Architecture, Formats, Integration, Security, and Validation",
    )
    add_toc_section(document, title)

    document.add_heading("1. Product Scope and Invariants", level=1)
    add_paragraphs(
        document,
        [
            "Delphi App Translation Studio is a Windows-only, open-source localization workspace written in Delphi FireMonkey. It targets Delphi VCL and FMX applications compiled for Win32 and Win64. It is not a runtime translation service and does not require target machines to have Internet access.",
            "The central invariant is separation: scan and catalog operations do not modify target source; provider access exists only in the Studio; recommended Component Integration writes only to the Studio export tree; advanced automatic integration remains explicit, previewed, transactional, backed up, and reversible; target applications consume compact offline JSON only.",
        ],
    )
    add_bullets(
        document,
        [
            "All user interface controls are persisted in DAT.Studio.MainForm.fmx and remain editable in the RAD Studio designer.",
            "No UI is constructed at runtime.",
            "Stable keys, not source text alone, identify translated properties.",
            "Original designer text remains the runtime fallback.",
            "API keys never enter project artifacts.",
            "One designer-persisted manager supervises the application; ordinary forms require no component.",
        ],
    )

    document.add_heading("2. Repository and Build Layout", level=1)
    add_table(
        document,
        ["Path", "Responsibility"],
        [
            ["source\\core", "Catalog types, JSON persistence, project detection, workspace paths, runtime pack generation."],
            ["source\\scan", "VCL/FMX form text, Pascal resourcestrings, extraction rules, diagnostics, incremental merge."],
            ["source\\validation", "Catalog safety and completeness checks."],
            ["source\\provider", "Provider types/settings, Credential Manager, DeepL/Google HTTPS client."],
            ["source\\runtime", "Pack discovery/loading, preference, manager, VCL and FMX applicators."],
            ["source\\components", "Framework-neutral manager core, VCL/FMX lifecycle adapters, and supplied language selectors."],
            ["source\\design", "VCL and FMX Tool Palette registration units."],
            ["packages\\runtime / packages\\design", "Core and framework runtime BPLs plus Win32 IDE design packages."],
            ["source\\integration", "Non-mutating component kits plus advanced planning, resource/source changes, transactions, and reset."],
            ["source\\studio", "Designer-authored FMX UI and Studio self-localization."],
            ["source\\schemas", "Development and runtime JSON Schemas."],
            ["tools\\tests", "Foundation, runtime, integration, form-streaming, launch, and self-localization smoke tests."],
            ["docs / help / samples", "Product documentation, help source, and safe fixtures."],
        ],
    )
    document.add_heading("2.1 Toolchain", level=2)
    document.add_paragraph(
        'call "C:\\Program Files (x86)\\Embarcadero\\Studio\\37.0\\bin\\rsvars.bat"\n'
        "msbuild DelphiAppTranslationStudio.dproj /t:Build /p:Config=Debug /p:Platform=Win32\n"
        "msbuild DelphiAppTranslationStudio.dproj /t:Build /p:Config=Release /p:Platform=Win64",
        style="Code Block",
    )
    add_paragraphs(
        document,
        [
            "The DPROJ identifies the main form as FMX, includes the standard project resource, uses source\\provider in its unit search path, and routes EXE/DCU output to bin and dcu by platform/configuration.",
        ],
    )

    document.add_heading("3. System Architecture", level=1)
    add_table(
        document,
        ["Boundary", "Input", "Output"],
        [
            ["Detection and scan", ".dproj/.dpr, text DFM/FMX, Pascal source", "TProjectProfile and TProjectScanResult"],
            ["Catalog", "Scan result plus existing development JSON", "Merged TTranslationCatalog"],
            ["Provider", "Confirmed source strings plus in-memory credential", "Machine-translated entry values"],
            ["Validation", "Development catalog", "Errors, warnings, information"],
            ["Pack builder", "Valid catalog", "Compact runtime JSON"],
            ["Component integration", "Project profile, scanner roots, packs", "Non-mutating component kit under Studio export"],
            ["Advanced integration", "Project profile, packs, designated menu", "Preview package and transactional change set"],
            ["Runtime", "Local JSON and per-user preference", "Translated existing controls and locale settings"],
        ],
    )

    document.add_heading("4. Core Data Model and JSON", level=1)
    add_paragraphs(
        document,
        [
            "DAT.Core.Types defines target framework/platform flags, locale metadata, translation status, catalog entries, and catalog ownership. Development JSON retains source text, translation, stable key, checksum, status, source location, component/property metadata, and locale data. Runtime JSON deliberately omits development-only detail.",
        ],
    )
    document.add_heading("4.1 Stable keys", level=2)
    add_bullets(
        document,
        [
            "Form properties use form.component.property, for example frmMain.btnSave.Text.",
            "resourcestring entries use the Pascal unit and symbol.",
            "Keys remain stable when wording changes, allowing source-changed detection.",
            "Fallback source text remains in the compiled DFM/FMX or code.",
        ],
    )
    document.add_heading("4.2 Development catalog", level=2)
    document.add_paragraph(
        '{\n'
        '  "schemaVersion": 5,\n'
        '  "applicationId": "SampleApp",\n'
        '  "sourceLanguage": "en-US",\n'
        '  "locale": {"languageCode": "it-IT", "nativeLanguageName": "Italiano"},\n'
        '  "entries": [{"key": "frmMain.btnSave.Text", "sourceText": "Save",\n'
        '               "translatedText": "Salva", "status": "machineTranslated",\n'
        '               "translationOrigin": "google",\n'
        '               "runtimeApplication": "automatic"}]\n'
        "}",
        style="Code Block",
    )
    document.add_heading("4.3 Runtime pack", level=2)
    document.add_paragraph(
        '{\n'
        '  "schemaVersion": 2,\n'
        '  "applicationId": "SampleApp",\n'
        '  "languageCode": "it-IT",\n'
        '  "language": {"code": "it-IT", "nativeName": "Italiano", "direction": "ltr"},\n'
        '  "sourceCatalogChecksum": "...",\n'
        '  "strings": {"frmMain.btnSave.Text": "Salva"},\n'
        '  "sources": {"frmMain.btnSave.Text": "Save"},\n'
        '  "sourceStrings": {"Save": "Salva"},\n'
        '  "sourceTemplates": {}\n'
        "}",
        style="Code Block",
    )
    add_paragraphs(
        document,
        [
            "The schema files are source\\schemas\\development-project.schema.json and runtime-language-pack.schema.json. Changes require schema-version handling, round-trip fixtures, and compatibility notes.",
        ],
    )

    document.add_heading("5. Project Detection and Scanning", level=1)
    add_paragraphs(
        document,
        [
            "DAT.Core.ProjectDetection reads project metadata, resolves the DPR, detects VCL or FMX through project/form evidence, records Win32/Win64 support, and enumerates forms and Pascal sources. DAT.Scan.Project coordinates form and Pascal scanners while measuring elapsed milliseconds. Generated folders and nested directories that contain another DPROJ or DPR are excluded so a conversion utility is not mistaken for part of the selected application.",
            "DAT.Scan.FormText parses text DFM/FMX without instantiating target forms, including Delphi multiline string expressions. DAT.Scan.PascalResources extracts resourcestring declarations plus supported runtime UI assignments and Format templates. It rejects SQL and HTML payloads that merely use Text properties. DAT.Scan.TextCodec preserves Delphi text encodings and escaped string syntax. DAT.Scan.Rules centralizes designated property decisions.",
        ],
    )
    document.add_heading("5.1 Incremental merge", level=2)
    add_steps(
        document,
        [
            "Index the existing catalog by stable key.",
            "Add unseen keys as needs-translation.",
            "Preserve translation and status when source text is unchanged.",
            "Preserve existing translation but mark source-changed when source text differs.",
            "Mark catalog-only entries obsolete unless excluded.",
            "Report new, changed, unchanged, and obsolete counts.",
        ],
    )

    document.add_heading("6. Studio UI and Workflow State", level=1)
    add_paragraphs(
        document,
        [
            "DAT.Studio.MainForm.fmx contains the complete orange-and-blue interface. The form persists WindowState=wsMaximized, uses client-aligned workflow cards, anchors resizable work areas to every relevant edge, and keeps the status card at the bottom. The workflow selection rectangle moves among Project, Scan, Translate, Validation, Export, Integration, and Provider Settings. Language and provider choices, Translate Automatically, key-management actions, and all other controls are persisted designer objects. DAT.Studio.MainForm.pas contains event and state logic only; it does not construct controls.",
            "Each workflow TLabel explicitly persists HitTest=True. FireMonkey labels default to HitTest=False, which would otherwise pass mouse events through the visible label even when an OnClick event is assigned. The setting remains editable in the Object Inspector.",
            "The form owns the project profile, scan result, catalog, validation result, integration change set, provider settings, and per-provider session-key strings. Destructors release owned objects. Catalog updates invalidate validation and export state. The designer-authored Integration method combo defaults to Component Integration; mutation controls are hidden in that mode and restored only for the explicitly selected advanced mode.",
        ],
    )
    add_callout(
        document,
        "FMX validation requirement.",
        "A successful compile is insufficient. Persisted FMX property-type errors surface only while streaming. StudioFormSmokeTests directly constructs the form, and launch tests verify the real title in every configuration.",
    )

    document.add_heading("6.1 Setup Wizard Architecture", level=2)
    add_paragraphs(
        document,
        [
            "DAT.Studio.SetupWizard is a separate, designer-authored FMX modal form opened by Start Setup Wizard. It is borderless, fixed-size, centered, and deliberately omits menu and system controls. Its nine hidden TabControl pages, including the Deployment Destinations page, left navigation rail, footer controls, labels, dialogs, lists, check boxes, memos, and layout geometry are all persisted in DAT.Studio.SetupWizard.fmx and remain editable in the Form Designer. A designer-authored modal backdrop dims and blocks the Studio while the Wizard is open so the two blue interfaces remain visually distinct.",
            "The Wizard owns isolated project-profile, scan-result, catalog, provider-key, backup, and component-kit state. Steps already reached may be revisited; future steps remain disabled. Changing the project or target language invalidates downstream scan/catalog state. Before final processing, the target project is read only except for an explicitly saved provider credential outside the project. During final processing, navigation and closing are disabled.",
            "Final processing requires confirmation that the target project is closed, creates a timestamped ZIP, translates eligible unresolved entries, saves the development catalog, validates, exports the runtime JSON pack, generates the component kit, deploys packs to existing outputs, and transactionally inserts marked Search Path and PostBuildEvent properties inside Delphi's native Base compiler property group. Base inheritance covers Debug/Release and Win32/Win64. The Wizard never edits target Pascal, form, or DPR files and never writes RAD Studio's package registry.",
        ],
    )
    add_callout(
        document,
        "Wizard regression requirement.",
        "StudioFormSmokeTests must instantiate the Wizard and verify every primary designer component, including dlgOpenProject. The project-selection access violation of August 10, 2026 was caused by a declared TOpenDialog omitted from the FMX resource; the explicit presence assertion prevents that class of streaming omission from returning.",
    )

    document.add_heading("7. Provider Settings and Secret Storage", level=1)
    add_paragraphs(
        document,
        [
            "DAT.Provider.Settings persists only provider, DeepL plan, remember choice, timeout, and batch size in %LOCALAPPDATA%\\DelphiAppTranslationStudio\\provider-settings.json. Bounds are normalized to 5-300 seconds and 1-50 strings.",
            "DAT.Provider.CredentialStore uses CredWriteW, CredReadW, CredDeleteW, and CredFree with CRED_TYPE_GENERIC and local-machine persistence. Credential targets are provider-specific and begin TMartinDub/DelphiAppTranslationStudio/. Secret bytes are cleared after conversion where practical.",
        ],
    )
    add_table(
        document,
        ["Choice", "Persistence", "Replacement behavior"],
        [
            ["Remember checked", "Windows Credential Manager", "Writes/replaces selected provider credential and clears session copy."],
            ["Remember cleared", "Process memory only", "Deletes selected provider stored credential and retains entered key until shutdown."],
            ["Remove Key", "None", "Deletes stored and session copies for selected provider."],
        ],
    )
    add_bullets(
        document,
        [
            "The masked edit never displays a stored credential.",
            "The effective key resolves new field text, then session memory, then Credential Manager.",
            "No error message contains the key, request headers, or provider response body.",
            "Credential operations occur only when needed; the form can start without network access.",
        ],
    )

    document.add_heading("8. Provider HTTP Clients", level=1)
    add_paragraphs(
        document,
        [
            "DAT.Provider.Client uses Delphi THTTPClient. It validates a nonblank key, normalizes settings, builds provider-specific JSON, sends HTTPS POST requests, parses provider response arrays, verifies result counts, and returns translations in source order.",
        ],
    )
    add_table(
        document,
        ["Provider", "Endpoint", "Authentication", "Payload"],
        [
            ["DeepL API Free", "https://api-free.deepl.com/v2/translate", "Authorization: DeepL-Auth-Key", "text array, source_lang, target_lang, context"],
            ["DeepL API Pro", "https://api.deepl.com/v2/translate", "Authorization: DeepL-Auth-Key", "text array, source_lang, target_lang, context"],
            ["Google Basic v2", "https://translation.googleapis.com/language/translate/v2", "X-Goog-Api-Key", "q array, source, target, format=text"],
        ],
    )
    document.add_heading("8.1 Reliability and cancellation", level=2)
    add_bullets(
        document,
        [
            "Batch size is capped at 50 and Google language tags are reduced to base language where appropriate.",
            "DeepL source codes are normalized to two-letter uppercase; targets retain supported regional forms.",
            "DeepL receives a batch context assembled from each entry's inferred UI role and semantic concept. DeepL documents context as unbilled supporting text.",
            "Google Basic v2 has no context or glossary request field. Context is used locally for terminology, memory matching, and ambiguity warnings but is not sent to Google.",
            "HTTP 429 and 5xx responses receive three bounded attempts with short backoff.",
            "Other non-2xx statuses fail immediately with provider, status, and corrective categories.",
            "A cancellation callback is checked between batches.",
            "Test Connection translates one harmless English phrase to Italian.",
        ],
    )
    document.add_heading("8.2 Official protocol references", level=2)
    add_bullets(
        document,
        [
            "DeepL authentication: https://developers.deepl.com/docs/getting-started/auth",
            "DeepL quickstart: https://developers.deepl.com/docs/getting-started/quickstart",
            "Google Basic translate method: https://docs.cloud.google.com/translate/docs/reference/rest/v2/translate",
            "Google Cloud Translation authentication: https://docs.cloud.google.com/translate/docs/authentication",
            "Google API-key best practices: https://docs.cloud.google.com/docs/authentication/api-keys-best-practices",
        ],
    )

    document.add_heading("9. Automatic Provider Translation and Review", level=1)
    add_paragraphs(
        document,
        [
            "Translate Automatically is the canonical machine-translation path. It validates the selected catalog language, loads the selected provider settings and effective key, counts eligible unresolved entries, and shows the provider and count before any network request.",
            "Every scan entry receives ContextKind, ContextDescription, SemanticConcept, and ContextConfidence metadata. Before a network request, the Studio reuses a Reviewed or Approved translation only when source text and semantic context match, then applies supported vetted UI terminology. The same authoritative terminology repairs unreviewed Google or DeepL drafts when a definitive application term is known, including commands, media playback, schedule terms, weekday abbreviations, and runtime uptime templates. Remaining strings are sent through DAT.Provider.Client in bounded batches. Reviewed, Approved, Edited, Excluded, and Obsolete entries are not overwritten.",
            "DeepL receives contextual descriptions. Google Basic results are tagged provider-basic; short terms with unknown meaning receive a review note. DeepL contextual results are tagged contextual-provider. Terminology and translation-memory results retain their own provenance.",
            "If no key is available, the workflow opens Provider Settings and gives a specific corrective message. HTTP or parsing failures leave the catalog entries unchanged for the failed operation and report a redacted summary in the status card.",
            "CSV interchange and manual editing remain optional alternatives for translation-company collaboration or specialized review, not prerequisites for automatic translation.",
            "The Studio may automatically reuse a Reviewed or Approved translation when both source text and semantic context match. It does not reuse approval state: the new entry remains Machine translated for review. Suggestions remain available for deliberate acceptance of other exact-source candidates.",
        ],
    )
    add_bullets(
        document,
        [
            "Mark Reviewed requires nonblank translated text.",
            "Approve requires the entry to have reached Reviewed first.",
            "Review All changes every nonblank active draft that is not already Reviewed, Approved, Excluded, or Obsolete after one count-bearing confirmation. Approve All changes only Reviewed entries after a second confirmation. Each bulk operation invalidates prior validation and saves the canonical JSON catalog once.",
            "Translation origin is independent of linguistic status.",
            "Inconsistent repeated terms and structural defects form the focused exception queue; a valid machine-translated draft is not itself an error.",
            "Structural validity, linguistic status, automatic runtime coverage, and manual wiring readiness are separate measures.",
            "A Pascal resourcestring remains manual wiring until the developer confirms the generated TranslateText call site.",
        ],
    )

    document.add_heading("10. Validation and Runtime Pack Export", level=1)
    add_paragraphs(
        document,
        [
            "DAT.Validation.Catalog validates metadata, missing target text, duplicates, changed source, Delphi indexed and sequential Format arguments, accelerators, inconsistent repeated-source terminology within the same semantic context, ambiguous context, and status conditions. Export is blocked by errors. Common runtime placeholders such as -- are informational when intentionally unchanged. The UI explains severity and maps a double-clicked entry issue back to the Translate list. Manual resourcestring wiring is reported separately from structural errors. DAT.Core.RuntimePack serializes only runtime-required metadata and strings and records a source-catalog checksum.",
            "Locale data is retained in the runtime pack so DAT.Runtime.Manager can expose a pack-specific TFormatSettings without globally mutating the developer's source code.",
        ],
    )

    document.add_heading("11. Runtime Engine", level=1)
    add_paragraphs(
        document,
        [
            "DAT.Runtime.LanguagePack loads and discovers JSON packs, checks application identity, rejects empty/invalid packs, canonicalizes native language names, de-duplicates exact locale codes, and suppresses a generic locale when a regional variant exists. Runtime schema 2 retains keyed strings and templates while adding source text, source-string, and source-template indexes; schema 1 packs remain readable. DAT.Runtime.Preference reads/writes the selected locale. DAT.Runtime.Manager owns the active pack, discovery, preference, translation lookup, and locale format settings. The source language loads its generated JSON pack when present and falls back to designer text only for older deployments without that pack.",
            "DAT.Runtime.VCL and DAT.Runtime.FMX traverse existing component trees and supported collection properties. They apply values by stable key to already designer-created controls. The FMX adapter also translates anonymous runtime-created components, grid cell text, and supported browser-generated text through the source indexes. Missing keys retain source text. Accepted JSON layout rules may adjust only the recorded safe properties for the active language; the adapters do not create controls or rewrite Delphi form source.",
        ],
    )
    document.add_heading("11.1 Component manager core", level=2)
    add_paragraphs(
        document,
        [
            "TDATCustomLanguageManager owns TTranslationRuntime, immutable post-initialization configuration, stable form-identity mappings, generation tracking, deterministic removal, main-thread and reentrancy guards, exclusion policy, diagnostics, and lifecycle events. SelectLanguage advances the generation, applies the active pack to open forms, saves the preference, and raises additive notifications.",
            "TDATFMXLanguageManager subscribes additively to TFormBeforeShownMessage and TFormReleasedMessage. It translates after streaming but before OnShow/first paint and never replaces a global handler. Its designer-visible AutoRefreshDynamicText and DynamicRefreshInterval properties reapply the active pack to visible forms so timers and application code cannot permanently overwrite translated status text, menu captions, grid headers, or uptime displays. TDATVCLLanguageManager owns a private TApplicationEvents, discovers visible forms on throttled idle, and inspects modal forms before display. VCL has no public additive before-show notification for every dynamic modeless form; such a form can paint once in the source language unless the application calls ApplyToForm before Show.",
            "PreserveControlState protects writable edit/memo data, focus, selection ranges, and list/combo ItemIndex. Collection replacement temporarily suppresses and then restores OnChange. Read-only instructional memo content remains translatable.",
        ],
    )
    document.add_heading("11.2 Required language-selection UI", level=2)
    add_paragraphs(
        document,
        [
            "TDATVCLLanguageComboBox and TDATFMXLanguageComboBox are the supplied designer-owned language selectors. Their typed LanguageManager property streams in DFM/FMX resources. At runtime AutoPopulate obtains validated descriptors, ShowLanguageCode controls display formatting, and selection calls the manager while preserving the inherited OnChange notification. RefreshLanguages supports packs deployed after startup. A connected Language menu is permitted instead, but the localized application must expose some visible language-selection UI.",
        ],
    )
    document.add_heading("11.3 Deployment paths", level=2)
    add_bullets(
        document,
        [
            "Packs: <ExecutableDirectory>\\Localization\\Languages\\<locale>.json.",
            "Preference: %LOCALAPPDATA%\\<ApplicationId>\\language.ini.",
            "Studio preference: %LOCALAPPDATA%\\DelphiAppTranslationStudio\\language.ini.",
            "Developer catalogs remain in the target source tree under Localization\\Development.",
        ],
    )

    document.add_heading("12. Component Packages and Integration Modes", level=1)
    add_paragraphs(
        document,
        [
            "The package graph contains DATLanguageManagerCoreRuntime, framework-specific VCL/FMX runtime packages, and separate VCL/FMX design packages. Each design package contains its DAT runtime/component units directly and therefore imports no custom DAT runtime BPL; this eliminates the IDE loader failure caused by missing dependent modules. Design packages register only their applicable manager and selector on DAT Localization. Runtime packages build for Win32 and Win64; RAD Studio consumes Win32 design BPLs.",
            "The Studio's Show Design BPL action selects the matching, verified, self-contained package under bin\\packages\\Win32\\Release. The developer installs that stable file with RAD Studio's Component > Install Packages > Add command. The Studio does not copy BPLs to public Delphi folders, write Known Packages registry values, close/restart RAD Studio, or expose automatic package installation. Delphi therefore remains the sole owner of design-package registration. Generated per-project kits contain no BPLs and are regenerated from a clean output directory, so they can be replaced without invalidating Delphi's registered package path.",
            "DAT.Integration.ComponentPackage implements the recommended non-mutating mode. It emits applicable runtime/component units, validated translated packs, an English source pack, component-integration.json, scanner form roots, README instructions, and a deployment script exclusively below export\\component-integration. SHA-256 fixture tests prove target DPROJ and DFM/FMX hashes remain unchanged.",
            "The advanced fallback retains DAT.Integration.Plan, DAT.Integration.Package, DAT.Integration.Engine, MenuResource, DelphiSource, Transaction, BuildDeploy, and Reset. It generates the application unit and exact changes in memory, requires explicit authorization, verifies a pre-change backup, writes atomically, rolls back failures, and supports restore/complete reset.",
        ],
    )
    document.add_heading("12.1 Component startup contract", level=2)
    add_bullets(
        document,
        [
            "The primary form streams one manager through ordinary Delphi designer mechanics.",
            "Loaded initializes the runtime, loads the preferred/source language, and applies the owner.",
            "FMX before-show messages cover later normal, inherited, and popup forms.",
            "VCL idle/modal discovery covers ordinary forms; strict modeless pre-display callers use ApplyToForm.",
            "SelectLanguage applies the pack to currently open forms immediately and persists the locale.",
            "The generated en-US pack makes switching back to source text deterministic.",
            "Manager and lifecycle subscriptions are released deterministically.",
        ],
    )
    add_callout(
        document,
        "Original code preservation.",
        "Recommended Component Integration performs zero automatic writes to the target. The developer's normal act of placing the manager causes Delphi to persist one component field/resource and applicable uses unit. Original DFM/FMX text remains the source-language fallback.",
    )

    document.add_heading("13. Self-Localization", level=1)
    add_paragraphs(
        document,
        [
            "DAT.Studio.Translation resolves the project/deployed pack directory, initializes a Studio runtime before form creation, applies the active pack in FormCreate, and maps persisted language-menu names to locale codes. Package streaming fixtures run in isolated bin\\tests\\packages directories so their local JSON cannot shadow the Studio's project/deployed packs.",
        ],
    )

    document.add_heading("14. Error Handling and Diagnostics", level=1)
    add_bullets(
        document,
        [
            "Project, scan, catalog, validation, export, integration, credential, and provider errors are caught at Studio event boundaries and summarized in the status card.",
            "Provider exceptions carry an optional HTTP status but not response bodies.",
            "Form scanners return file/line diagnostics without instantiating target UI.",
            "Integration previews show a complete line-numbered exact diff for every affected file before mutation.",
            "Transaction errors trigger rollback and retain backup evidence.",
            "Provider keys must never be added to troubleshooting logs.",
        ],
    )

    document.add_heading("15. Test Strategy and Verified Matrix", level=1)
    add_table(
        document,
        ["Test", "Coverage"],
        [
            ["FoundationSmokeTests", "Detection, scan-to-catalog, schema/provenance round-trip, provider contracts, review/approval, validation, runtime pack, preference, advanced integration, component-kit contents, and SHA-256 non-mutation proof."],
            ["Language manager suites", "Core guards/generations, FMX before-show lifecycle, VCL discovery/modal boundary, stable identities, instant switching, state preservation, and deterministic cleanup on Win32/Win64."],
            ["Package/selector suites", "Debug and Release runtime/design packages, DFM/FMX streaming, typed manager references, pack discovery, and selector language propagation on Win32/Win64."],
            ["VCLRuntimeSmokeTests", "VCL controls, menu items, locale, generated unit."],
            ["FMXRuntimeSmokeTests", "All representative FMX form properties, menu items, locale, generated unit."],
            ["StudioFormSmokeTests", "Direct FMX stream/create, maximized client-aligned pages, provider-only automatic-translation controls, exact-review controls, linguistic action wiring, and Provider Settings activation."],
            ["RunRuntimeSmokeTests.ps1", "Both compilers, scanner/catalog/provider contracts, disposable integrated VCL/FMX builds, deployed Italian pack, and required Italian launch title."],
            ["RunStudioLaunchSmokeTests.ps1", "Debug/Release Win32/Win64 real main-window title."],
            ["RunStudioSelfLocalizationSmokeTest.ps1", "Italian Studio title in all four configurations with state restoration."],
        ],
    )
    add_paragraphs(
        document,
        [
            "On August 9, 2026, the complete release harness passed uninterrupted. Debug and Release component packages, Win32/Win64 manager suites, scanner/catalog/provider contracts, component non-mutation, runtime and advanced integration all passed. The Studio built in Debug and Release for both architectures, streamed its FMX form directly, launched normally, and self-localized to Italian in all four configurations. Disposable real-application pilots also built and opened translated first forms on Win32/Win64 for Website Analytics (FMX) and Courier Herald Reader (VCL).",
        ],
    )
    document.add_heading("15.1 Optional live provider testing", level=2)
    add_paragraphs(
        document,
        [
            "Automated fixtures must not contain live keys. Live provider acceptance uses an owner-supplied restricted key through Provider Settings, confirms Test Connection, translates a small disposable catalog, checks Machine translated status and provider provenance, then removes or rotates the test key. Provider availability is external state and cannot be made deterministic, but live provider acceptance is required before a public release that claims current Google or DeepL compatibility.",
        ],
    )

    document.add_heading("16. Security Review", level=1)
    add_table(
        document,
        ["Threat", "Control"],
        [
            ["Key committed to Git", "No key field in settings JSON/catalogs; Credential Manager or memory only; release secret scan."],
            ["Key leaked through URL/log", "Authentication headers, redacted errors, no body/header logging."],
            ["Wrong target changed", "Recommended component kit has no target-write path; advanced mode uses profile display, preview, and explicit Apply."],
            ["Unrecoverable source mutation", "Verified G-drive project backup by workflow policy plus transaction backup/manifest/restore."],
            ["Installed-folder write failure", "Per-user language preference; packs are read-only deployment assets."],
            ["Machine mistranslation shipped", "Machine-translated status, validation, required human review."],
            ["Pack for wrong application", "applicationId validation in pack discovery/loading."],
        ],
    )
    add_paragraphs(
        document,
        [
            "Windows Credential Manager follows Microsoft guidance for application credentials (CredWrite/CredRead). It protects persistence under the signed-in Windows context; it does not make a key safe from malware running as that user. Provider-side restrictions, quotas, rotation, and least privilege remain necessary.",
        ],
    )

    document.add_heading("17. Release and Contribution Workflow", level=1)
    add_steps(
        document,
        [
            "Read global and repository directives.",
            "Confirm the exact requested scope and make a pre-change backup when changes are approved.",
            "Keep UI edits in the FMX designer resource and source event logic separate.",
            "Add deterministic tests without credentials.",
            "Run Win32/Win64 Debug/Release builds elevated.",
            "Run runtime, form-streaming, launch, and self-localization suites.",
            "Update phase notes, help, User Guide, Engineering Guide, real TOCs, PDFs, and visual QA.",
            "Check for a stale zero-byte .git\\index.lock with no Git process.",
            "Review status, stage only intended files, commit clearly, and push every configured remote.",
        ],
    )

    document.add_heading("18. Known Boundaries and Future-Compatible Design", level=1)
    add_bullets(
        document,
        [
            "Only Windows Win32/Win64 Delphi applications are supported.",
            "Google Advanced v3, OAuth/service-account workflows, and other providers are not implemented.",
            "Provider operations are explicit Studio actions; target runtime remains offline.",
            "No automatic control resizing/reflow is performed.",
            "Binary DFM conversion is outside the scanner.",
            "Automatic runtime application covers scanned designer properties. Application-specific strings still require the documented TranslateText wiring.",
            "VCL dynamic modeless forms may paint once in source language before idle discovery; call ApplyToForm before Show when a no-flicker first display is mandatory.",
            "Provider account terms and language support are external and must be rechecked for each release.",
        ],
    )

    document.add_heading("19. Documentation Generation", level=1)
    add_paragraphs(
        document,
        [
            "The editable guides are generated from actual source and engineering notes with python-docx. Each DOCX contains a real TOC field, a dedicated TOC section, and a new-page content section. LibreOffice performs the approved headless PDF conversion with an isolated user profile. Microsoft Word is not used. Every final PDF is rendered to page images and inspected before release.",
        ],
    )

    document.add_heading("20. Source References", level=1)
    add_bullets(
        document,
        [
            "Repository source, forms, project metadata, schemas, and smoke tests as of August 10, 2026.",
            "Microsoft credential handling: https://learn.microsoft.com/en-us/windows/win32/secbp/handling-passwords",
            "Windows CREDENTIAL structure: https://learn.microsoft.com/en-us/windows/win32/api/wincred/ns-wincred-credentialw",
            "DeepL developer documentation: https://developers.deepl.com/docs/getting-started/auth",
            "Google Cloud Translation documentation: https://docs.cloud.google.com/translate/docs/authentication",
            "Google API-key guidance: https://docs.cloud.google.com/docs/authentication/api-keys-best-practices",
        ],
    )

    path = GUIDES_DIR / "Delphi App Translation Studio Engineering Guide.docx"
    finish_document(document, path)
    return path


if __name__ == "__main__":
    for generated_path in (
        build_user_guide(),
        build_setup_wizard_guide(),
        build_engineering_guide(),
    ):
        print(generated_path)
