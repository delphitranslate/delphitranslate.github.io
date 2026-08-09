from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_guides import (
    BLUE,
    GRAY,
    INK,
    ORANGE,
    PALE_BLUE,
    add_bullets,
    add_callout,
    add_field,
    add_header_footer,
    add_paragraphs,
    add_steps,
    configure_page,
    finish_document,
    set_cell_shading,
    set_page_number_format,
    set_table_geometry,
    setup_styles,
)


PROJECT_ROOT = Path(__file__).resolve().parents[1]
GUIDES_DIR = PROJECT_ROOT / "docs" / "guides"
ICON = (
    PROJECT_ROOT
    / "images and icons"
    / "DelphiAppTranslationStudio-Icon-Master-v2_150.png"
)
DOCX_PATH = GUIDES_DIR / "Delphi App Translation Studio Complete Test Guide.docx"
LAST_CHANGED = "August 9, 2026"


def set_run_font(
    run,
    size: float = 11,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
    name: str = "Calibri",
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_cover(document: Document) -> None:
    section = document.sections[0]
    configure_page(section)
    add_header_footer(section, "", False)
    document.add_paragraph("")
    document.add_paragraph("")
    logo = document.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ICON.exists():
        logo.add_run().add_picture(str(ICON), width=Inches(1.45))
    heading = document.add_paragraph(style="Title")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run("Complete Test Guide")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Delphi App Translation Studio")
    accent = document.add_table(rows=1, cols=2)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(accent, [4680, 4680], indent=0)
    set_cell_shading(accent.cell(0, 0), BLUE)
    set_cell_shading(accent.cell(0, 1), ORANGE)
    for cell in accent.rows[0].cells:
        cell.height = Inches(0.08)
    document.add_paragraph("")
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Version 1.0\n"
        f"Last changed: {LAST_CHANGED}\n"
        "User acceptance, component integration, regression, and release testing\n"
        "Windows - Delphi VCL and FireMonkey - Win32 and Win64"
    )
    set_run_font(run, 11, GRAY)
    purpose = document.add_paragraph()
    purpose.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = purpose.add_run(
        "A printable, start-to-finish procedure based on the current source, "
        "designer forms, build metadata, package graph, and verified test harness."
    )
    set_run_font(run, 10, GRAY, italic=True)


def add_toc(document: Document) -> None:
    toc_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(toc_section)
    set_page_number_format(toc_section, "lowerRoman", 1)
    add_header_footer(toc_section, "Complete Test Guide", True)
    document.add_paragraph("Table of Contents", style="TOC Heading")
    paragraph = document.add_paragraph()
    add_field(paragraph, 'TOC \\o "1-2" \\h \\z \\u')
    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(content_section)
    set_page_number_format(content_section, "decimal", 1)
    add_header_footer(content_section, "Complete Test Guide", True)


def add_code(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph(style="Code Block")
        paragraph.add_run(line)


def add_matrix(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    font_size: float = 9.2,
) -> None:
    if sum(widths) != 9360:
        raise ValueError("Table widths must total 9360 DXA")
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, font_size, "FFFFFF", bold=True)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        cannot_split = OxmlElement("w:cantSplit")
        table.rows[-1]._tr.get_or_add_trPr().append(cannot_split)
        for index, value in enumerate(row):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                set_cell_shading(cells[index], "F6F9FC")
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(run, font_size, INK)
    set_table_geometry(table, widths)
    document.add_paragraph("")


def add_test_case(
    document: Document,
    case_id: str,
    title: str,
    objective: str,
    steps: list[str],
    expected: list[str],
) -> None:
    document.add_heading(f"{case_id} {title}", level=2)
    add_callout(document, "Objective.", objective)
    document.add_paragraph("Procedure", style="Heading 3")
    add_steps(document, steps)
    document.add_paragraph("Expected result", style="Heading 3")
    add_bullets(document, expected)
    result = document.add_table(rows=2, cols=4)
    result.style = "Table Grid"
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [
        ["Result", "PASS / FAIL / N/A", "Tester", ""],
        ["Date", "", "Defect / notes", ""],
    ]
    for row_index, row in enumerate(values):
        for col_index, value in enumerate(row):
            cell = result.cell(row_index, col_index)
            cell.text = value
            if col_index in (0, 2):
                set_cell_shading(cell, PALE_BLUE)
                for run in cell.paragraphs[0].runs:
                    set_run_font(run, 8.8, BLUE, bold=True)
            else:
                for run in cell.paragraphs[0].runs:
                    set_run_font(run, 8.8, INK)
    set_table_geometry(result, [1200, 2400, 1500, 4260])
    document.add_paragraph("")


def build_document() -> Path:
    document = Document()
    setup_styles(document)
    add_cover(document)
    add_toc(document)

    document.add_heading("1. Purpose and Test Philosophy", level=1)
    add_paragraphs(
        document,
        [
            "This guide provides the complete acceptance path for Delphi App Translation Studio: safeguard the source, build the Studio and packages, run the automated release gate, scan and automatically translate a real FMX application, generate offline JSON packs, integrate one manager component without Studio-authored target changes, deploy the packs, and verify language switching on Win32 and Win64.",
            "The recommended production path is Component Integration. The Studio generates a kit beneath its own export folder and never opens the selected target project, DPR, DPROJ, PAS, DFM, or FMX files for writing. The developer then performs a small, normal Delphi Form Designer change in a disposable test copy: one manager on the primary form and, when desired, one typed language selector. This is visible, reviewable, and reversible in Git.",
            "The automated matrix is the engineering release gate. The manual cases are user acceptance tests. Both are required for a complete sign-off because automation can prove invariants while a human must confirm visual layout, translated meaning, normal application behavior, and operational clarity.",
        ],
    )
    add_callout(
        document,
        "Critical source-protection rule.",
        "Do not place test components in the pristine GA4 repository. Copy it to a disposable working folder first. The pristine repository remains the known-good comparison point.",
    )

    document.add_heading("1.1 Scope", level=2)
    add_bullets(
        document,
        [
            "Studio host: Windows Win32 and Win64.",
            "Target frameworks: Delphi VCL and FireMonkey (FMX).",
            "Target architectures: Win32 and Win64, Debug and Release where specified.",
            "Translation providers: Google Cloud Translation Basic v2 and DeepL API Free/API Pro.",
            "Runtime operation: local JSON packs with no Internet connection and no provider key.",
            "Out of scope: macOS, iOS, Android, Linux, C++Builder, automatic control resizing, and runtime cloud translation.",
        ],
    )

    document.add_heading("1.2 Stop conditions", level=2)
    add_bullets(
        document,
        [
            "Stop if the selected target is the pristine GA4 folder rather than a disposable copy.",
            "Stop if Git reports unexplained modifications before integration begins.",
            "Stop if a provider response alters placeholders, format specifiers, or accelerator semantics and validation reports an error.",
            "Stop if Component Integration changes any target file before the developer manually opens the target in Delphi.",
            "Stop if the release validation script reports any failure; preserve the complete console output for diagnosis.",
            "Stop if a build or launch error dialog appears. Record its full text and the exact executable used.",
        ],
    )

    document.add_heading("2. Exact Folders and Files", level=1)
    add_matrix(
        document,
        ["Purpose", "Exact path", "Handling"],
        [
            ["Studio source repository", r"C:\New Delphi Projects\Delphi App Translation", "This is the active product folder used by this guide."],
            ["Pristine GA4 reference", r"C:\New Delphi Projects\Echurchsite Analytical - Pristine GA4", "Read-only reference. It was clean when this guide was generated. Do not integrate directly into it."],
            ["Recommended disposable GA4 test copy", r"C:\New Delphi Projects\Echurchsite Analytical - Component Test", "Create from the pristine reference. Perform all manual component integration here."],
            ["RAD Studio environment", r"C:\Program Files (x86)\Embarcadero\Studio\37.0\bin\rsvars.bat", "Run builds elevated so MSBuild can read the RAD Studio environment."],
            ["Studio project", r"C:\New Delphi Projects\Delphi App Translation\DelphiAppTranslationStudio.dproj", "Open this in RAD Studio 13 Florence."],
            ["Studio form", r"C:\New Delphi Projects\Delphi App Translation\source\studio\DAT.Studio.MainForm.fmx", "Designer-authored FMX user interface."],
            ["Complete release gate", r"C:\New Delphi Projects\Delphi App Translation\tools\tests\RunPhase10ReleaseValidation.ps1", "Builds and tests the supported matrix."],
            ["Component-kit output root", r"C:\New Delphi Projects\Delphi App Translation\export\component-integration", "Generated by the recommended Integration mode."],
            ["Project-local development catalogs", r"<Target>\Localization\Development", "Lossless editable JSON catalogs; not runtime deployment files."],
            ["Project-local runtime packs", r"<Target>\Localization\Languages", "Compact offline JSON packs produced by Export."],
            ["Per-user language preference", r"%LOCALAPPDATA%\<ApplicationId>\language.ini", "Stores the last selected language; does not alter the application folder."],
        ],
        [1900, 4660, 2800],
        8.6,
    )

    document.add_heading("2.1 Studio executable matrix", level=2)
    add_matrix(
        document,
        ["Build", "Exact executable"],
        [
            ["Win32 Debug", r"C:\New Delphi Projects\Delphi App Translation\bin\Win32\Debug\DelphiAppTranslationStudio.exe"],
            ["Win32 Release", r"C:\New Delphi Projects\Delphi App Translation\bin\Win32\Release\DelphiAppTranslationStudio.exe"],
            ["Win64 Debug", r"C:\New Delphi Projects\Delphi App Translation\bin\Win64\Debug\DelphiAppTranslationStudio.exe"],
            ["Win64 Release", r"C:\New Delphi Projects\Delphi App Translation\bin\Win64\Release\DelphiAppTranslationStudio.exe"],
        ],
        [1900, 7460],
        9.0,
    )
    add_callout(
        document,
        "Recommended first manual run.",
        "Use the Win32 Debug Studio executable. After the workflow succeeds, repeat the launch and essential scan checks with Win64 Debug and both Release builds.",
    )

    document.add_heading("2.2 Package source and build outputs", level=2)
    add_matrix(
        document,
        ["Package", "Source DPK", "Built output"],
        [
            ["Core runtime", r"packages\runtime\DATLanguageManagerCoreRuntime.dpk", r"bin\packages\<Platform>\<Config>\DATLanguageManagerCoreRuntime.bpl/.dcp"],
            ["VCL runtime", r"packages\runtime\DATLanguageManagerVCLRuntime.dpk", r"bin\packages\<Platform>\<Config>\DATLanguageManagerVCLRuntime.bpl/.dcp"],
            ["FMX runtime", r"packages\runtime\DATLanguageManagerFMXRuntime.dpk", r"bin\packages\<Platform>\<Config>\DATLanguageManagerFMXRuntime.bpl/.dcp"],
            ["VCL design", r"packages\design\DATLanguageManagerVCLDesign.dpk", r"bin\packages\Win32\<Config>\DATLanguageManagerVCLDesign.bpl/.dcp"],
            ["FMX design", r"packages\design\DATLanguageManagerFMXDesign.dpk", r"bin\packages\Win32\<Config>\DATLanguageManagerFMXDesign.bpl/.dcp"],
        ],
        [1700, 3500, 4160],
        8.6,
    )
    add_paragraphs(
        document,
        [
            "Runtime packages build for Win32 and Win64. The Delphi IDE loads the self-contained Win32 design package. Installation is an explicit developer action performed only through RAD Studio's Component > Install Packages > Add command. The Studio does not copy, register, or automatically install BPLs.",
        ],
    )

    document.add_heading("3. Prerequisites and Test Data", level=1)
    add_bullets(
        document,
        [
            "Windows account able to run RAD Studio and an elevated PowerShell window for builds.",
            "RAD Studio 13 Florence / toolchain 37.0 installed at the paths in Section 2.",
            "Git installed at C:\\Program Files\\Git and available for status/diff verification.",
            "A working Google Cloud Translation Basic v2 API key or DeepL API key for the live automatic-translation case.",
            "Internet access for Provider Settings, Test Connection, and Translate Automatically only.",
            "The pristine GA4 reference folder and enough free space for a disposable copy and four builds.",
            "A second target language that can be visually distinguished from English. Spanish (Spain) [es-ES] or Italian (Italy) [it-IT] is suitable.",
        ],
    )
    add_callout(
        document,
        "Credential safety.",
        "Never paste an API key into a source file, catalog, runtime pack, test report, screenshot, Git commit, or command line. Enter it only in the Studio's masked Provider Settings field. Remembered keys are stored as Windows Generic Credentials for the signed-in user.",
    )

    document.add_heading("3.1 Create the disposable GA4 test copy", level=2)
    add_steps(
        document,
        [
            r"Close WebsiteAnalytics.exe and close the GA4 project in RAD Studio.",
            r'Confirm C:\New Delphi Projects\Echurchsite Analytical - Pristine GA4 is clean with: git -C "C:\New Delphi Projects\Echurchsite Analytical - Pristine GA4" status --short.',
            r"Copy the complete pristine folder to C:\New Delphi Projects\Echurchsite Analytical - Component Test. Preserve its .git directory so Git can show every manual integration change.",
            r'Run git -C "C:\New Delphi Projects\Echurchsite Analytical - Component Test" status --short. The result must be empty before the test begins.',
            r'Record the starting commit with: git -C "C:\New Delphi Projects\Echurchsite Analytical - Component Test" rev-parse HEAD.',
        ],
    )
    add_callout(
        document,
        "Existing older test folder.",
        r"C:\New Delphi Projects\Echurchsite Analytical - Test already contains earlier automatic-source-integration changes and Localization files. Do not use it as the clean component-first acceptance target.",
    )

    document.add_heading("3.2 Establish a clean Delphi package baseline", level=2)
    add_steps(
        document,
        [
            "Start RAD Studio without opening the disposable target project or any target form.",
            "Choose Component > Install Packages.",
            "If DAT Language Manager FireMonkey design-time package is listed, select it, choose Remove, confirm the removal, and choose OK. If it is not listed, choose Cancel and continue.",
            "Close and restart RAD Studio. Confirm that no package-load error appears. Do not manually delete BPLs and do not edit the BDS registry.",
            "Leave the target project closed until TC-10 installs the freshly generated design package.",
        ],
    )
    add_callout(
        document,
        "Use Delphi's package manager only.",
        "Never use Component > Install Component for this product, never select a .dpk, and never copy package files into Delphi system folders. The accepted workflow uses the compiled Win32 design .bpl through Component > Install Packages > Add.",
    )

    document.add_heading("4. Build the Studio Manually", level=1)
    add_paragraphs(
        document,
        [
            "Run builds from an elevated PowerShell or Developer Command Prompt. The project file sends executables to bin\\<Platform>\\<Configuration> and DCUs to dcu\\<Platform>\\<Configuration>.",
        ],
    )
    document.add_heading("4.1 RAD Studio method", level=2)
    add_steps(
        document,
        [
            r"Open C:\New Delphi Projects\Delphi App Translation\DelphiAppTranslationStudio.dproj.",
            "Select Win32 / Debug and choose Project > Build DelphiAppTranslationStudio.",
            "Repeat for Win64 / Debug, Win32 / Release, and Win64 / Release.",
            "Confirm the four exact executables listed in Section 2.1 have current timestamps and nonzero sizes.",
            "Return the active platform to Win32 after the matrix unless another immediate test requires Win64.",
        ],
    )
    document.add_heading("4.2 Command-line method", level=2)
    add_code(
        document,
        [
            r'cd "C:\New Delphi Projects\Delphi App Translation"',
            r'cmd /d /c "call ""C:\Program Files (x86)\Embarcadero\Studio\37.0\bin\rsvars.bat"" && msbuild DelphiAppTranslationStudio.dproj /t:Build /p:Config=Debug /p:Platform=Win32 /v:minimal"',
            r'cmd /d /c "call ""C:\Program Files (x86)\Embarcadero\Studio\37.0\bin\rsvars.bat"" && msbuild DelphiAppTranslationStudio.dproj /t:Build /p:Config=Debug /p:Platform=Win64 /v:minimal"',
            r'cmd /d /c "call ""C:\Program Files (x86)\Embarcadero\Studio\37.0\bin\rsvars.bat"" && msbuild DelphiAppTranslationStudio.dproj /t:Build /p:Config=Release /p:Platform=Win32 /v:minimal"',
            r'cmd /d /c "call ""C:\Program Files (x86)\Embarcadero\Studio\37.0\bin\rsvars.bat"" && msbuild DelphiAppTranslationStudio.dproj /t:Build /p:Config=Release /p:Platform=Win64 /v:minimal"',
        ],
    )

    document.add_heading("5. Automated Release Gate", level=1)
    add_paragraphs(
        document,
        [
            "Run this before manual acceptance. It performs one uninterrupted Debug/Release and Win32/Win64 matrix. It builds core, VCL, FMX, runtime, design, Studio, streaming, lifecycle, selector, launch, self-localization, foundation, runtime, integration, and non-mutation tests.",
        ],
    )
    add_code(
        document,
        [
            r'cd "C:\New Delphi Projects\Delphi App Translation"',
            r'powershell.exe -ExecutionPolicy Bypass -File .\tools\tests\RunPhase10ReleaseValidation.ps1',
        ],
    )
    add_bullets(
        document,
        [
            "Expected final line: Phase 10 complete release validation passed.",
            "Expected behavior: test processes may open briefly and close automatically.",
            "The script restores the Studio's self-localization pack and preference after its Italian launch checks.",
            "Any nonzero exit or exception is a failed release gate. Do not continue to sign-off until repaired and rerun from the beginning.",
        ],
    )

    document.add_heading("5.1 What the gate proves", level=2)
    add_matrix(
        document,
        ["Area", "Evidence"],
        [
            ["Package graph", "Debug and Release runtime packages for Win32/Win64; Win32 VCL/FMX design packages; DFM/FMX streaming."],
            ["Managers", "Initialization, lifecycle discovery, stable form identity, switching generations, state preservation, cleanup, and architecture parity."],
            ["Studio", "Four builds, direct FMX form streaming, ordinary launch, maximized layout contracts, and Italian self-localization."],
            ["Translation core", "Detection, scanning, schema/provenance, catalog merge, provider contracts, validation, pack generation, and preference handling."],
            ["Safety", "Component-kit generation writes only below the export root and preserves target project/form SHA-256 hashes."],
            ["Advanced fallback", "Preview, backup, transactional apply/rollback, deployment, restore, and complete-reset regression coverage."],
        ],
        [2200, 7160],
        9.0,
    )

    document.add_page_break()
    document.add_heading("6. Manual Studio Acceptance", level=1)
    add_test_case(
        document,
        "TC-01",
        "Launch and navigation",
        "Verify the designer-authored Studio shell opens maximized and every workflow page is reachable.",
        [
            r"Run C:\New Delphi Projects\Delphi App Translation\bin\Win32\Debug\DelphiAppTranslationStudio.exe.",
            "Confirm there is no startup error dialog and the title is Delphi App Translation Studio.",
            "Select Project, Scan, Translate, Validation, Export, Integration, and Provider Settings in order.",
            "Read the bottom status card after every selection.",
            "Resize the maximized window and inspect the top, bottom, left, and right edges of each page.",
        ],
        [
            "The window starts maximized.",
            "The blue workflow selection follows the active page.",
            "The bottom status card changes to page-specific guidance and never retains a message from the previous page.",
            "Every control remains visible, aligned, and usable; nothing bleeds below or beyond the client area.",
            "Component Integration (Recommended) is the default Integration method.",
            "Provider Settings contains only DeepL and Google Cloud Translation, with no CLI-agent controls.",
        ],
    )

    add_test_case(
        document,
        "TC-02",
        "Open the disposable GA4 project",
        "Verify correct FMX and platform detection without changing the target.",
        [
            "Select Project > Open Project.",
            r"Open C:\New Delphi Projects\Echurchsite Analytical - Component Test\WebsiteAnalytics.dproj.",
            "Read the project summary before scanning.",
            r'In a separate console, run git -C "C:\New Delphi Projects\Echurchsite Analytical - Component Test" status --short.',
        ],
        [
            "Project: WebsiteAnalytics.",
            "Framework: FireMonkey.",
            "Windows targets: Win32 and Win64.",
            "The Scan Project button is enabled.",
            "Git status remains empty; opening performs no target write.",
        ],
    )

    add_test_case(
        document,
        "TC-03",
        "Scan the project",
        "Verify fast read-only extraction of FMX designer properties and resourcestrings.",
        [
            "Select Scan.",
            "Choose Scan Project and wait for Scan complete in the status panel.",
            "Record total entries, form properties, resourcestrings, files scanned, and elapsed milliseconds.",
            "Inspect representative labels, buttons, headings, memo text, list content, and any resourcestring entries.",
            "Scroll rapidly from the top to the bottom and back through the scan-result list.",
            r'Run git -C "C:\New Delphi Projects\Echurchsite Analytical - Component Test" status --short again.',
        ],
        [
            "The scan completes without an exception.",
            "The result list contains stable form/component/property keys and readable source text.",
            "Rows remain separated while scrolling; text never paints over adjacent rows.",
            "Elapsed time is reported in milliseconds. Performance is recorded, not hard-coded to a particular machine.",
            "No Localization folder is created merely by scanning.",
            "Git status remains empty.",
        ],
    )

    add_test_case(
        document,
        "TC-04",
        "Configure and test a provider",
        "Verify secure Google or DeepL configuration and a live connection without exposing the key.",
        [
            "Select Provider Settings.",
            "Choose Google Cloud Translation for a Google Basic v2 API key, or choose DeepL and the correct API Free/API Pro plan.",
            "Leave timeout at 30 seconds and strings per request at 40 for the first test.",
            "Paste the key into the masked API key field.",
            "For a persistent test, check Remember securely on this computer. For a temporary test, clear it.",
            "Choose Replace / Save Key, then Test Connection.",
            "Confirm no API key appears in the status panel, project files, or generated JSON.",
        ],
        [
            "The masked field does not display the clear key.",
            "Test Connection succeeds with a small English-to-Italian request.",
            "Remembered mode reports Windows Credential Manager storage; session-only mode reports that the key will be forgotten when the Studio closes.",
            "The target project remains unchanged.",
        ],
    )

    add_test_case(
        document,
        "TC-05",
        "Create the development catalog",
        "Verify target-language lists, locale defaults, and project-local JSON catalog creation.",
        [
            "Select Translate.",
            "Confirm source language English (United States) [en-US].",
            "Select Spanish (Spain) [es-ES], Italian (Italy) [it-IT], or another intended language from the list.",
            "Review native language name, left-to-right/right-to-left direction, short/long date, short/long time, decimal, thousands, and currency values.",
            "Choose Save.",
            "Click the blue displayed catalog path. Confirm File Explorer opens with the JSON catalog selected, then open it in a text editor and confirm it is JSON, not CSV.",
        ],
        [
            r"Catalog path: C:\New Delphi Projects\Echurchsite Analytical - Component Test\Localization\Development\WebsiteAnalytics.<locale>.translation-project.json.",
            "The file is UTF-8 JSON and includes application, framework, locale, entries, status, provenance, source checksum, and runtime classification data.",
            "CSV Out and CSV In remain optional interchange tools; they are not the normal automatic-translation path and are not runtime packs.",
        ],
    )

    add_test_case(
        document,
        "TC-06",
        "Translate automatically",
        "Verify the Studio sends eligible catalog entries to the configured provider, records returned values one by one, and auto-saves the JSON catalog.",
        [
            "On Translate, choose Translate Automatically.",
            "Read the confirmation: it must name the provider, unresolved count, protected reviewed/approved behavior, and automatic saving.",
            "Choose Yes and allow every bounded batch to finish.",
            "Inspect entries from the top, middle, and bottom of the catalog.",
            "Close and reopen the catalog or Studio, then confirm the translations remain present.",
        ],
        [
            "Eligible Needs Translation, Source Changed, and Error entries receive translated text.",
            "Returned entries are marked Machine translated with Google or DeepL origin; they are not silently approved.",
            "Reviewed, Approved, Excluded, and Obsolete entries are not overwritten.",
            "The development JSON catalog is saved automatically after a successful provider result.",
            "Provider credentials do not appear in the catalog.",
        ],
    )

    add_test_case(
        document,
        "TC-07",
        "Focused linguistic and structural review",
        "Verify review actions remain available without requiring every entry to be individually approved merely to create a draft.",
        [
            "Review visible UI strings, product names, short buttons, multiline instructions, placeholders, and accelerators.",
            "Correct a deliberately awkward machine result in Translated text and choose Apply Translation.",
            "Mark that entry Reviewed, then Approve it.",
            "For the remaining translated drafts, choose Review All, verify the count in the confirmation, and accept. Then choose Approve All, verify its Reviewed count, and accept.",
            "If the catalog has a Pascal resourcestring, confirm its runtime classification and leave Manual TranslateText wiring confirmed clear unless the target code has actually been wired.",
        ],
        [
            "An edited entry records human origin and Edited status.",
            "Reviewed and Approved are separate deliberate steps.",
            "Each bulk action operates on its stated eligible set, saves once, and leaves Excluded, Obsolete, and already Approved entries unchanged.",
            "Manual resourcestring wiring readiness remains separate from designer-property coverage.",
            "The catalog is saved after edits when a catalog filename exists.",
        ],
    )

    add_test_case(
        document,
        "TC-08",
        "Validate and export runtime JSON",
        "Verify structural defects block export and a valid catalog produces a compact offline pack.",
        [
            "Select Validation and choose Validate Catalog.",
            "Read the plain-language summary. Double-click an entry-specific warning and confirm Translate opens with that entry selected.",
            "Resolve all errors involving missing text, duplicate keys, changed source, placeholders, sequential/indexed Format arguments, or accelerator keys. Review warnings; informational runtime placeholders require no action.",
            "Treat manual resourcestring wiring as a separate readiness warning and resolve it before claiming complete runtime coverage.",
            "Select Export and choose Export Runtime Pack.",
            "Open the displayed output file and confirm it is JSON.",
        ],
        [
            "Export is blocked while validation errors exist.",
            "A valid catalog reports that the runtime pack is ready.",
            r"Runtime path: C:\New Delphi Projects\Echurchsite Analytical - Component Test\Localization\Languages\<locale>.json.",
            "The runtime pack is smaller than the development catalog and contains no API key or development history.",
        ],
    )

    document.add_page_break()
    document.add_heading("7. Component Integration Acceptance", level=1)
    add_test_case(
        document,
        "TC-09",
        "Generate the non-mutating component kit",
        "Verify the recommended mode creates a complete kit under the Studio export folder and performs zero target writes.",
        [
            r'Before generation, run git -C "C:\New Delphi Projects\Echurchsite Analytical - Component Test" status --short and record the expected Localization files only.',
            "Select Integration and leave Component Integration (Recommended) selected.",
            "Choose Build Integration Plan.",
            "Read the setup plan and choose Generate Component Kit.",
            "Select generated filenames in the left pane and inspect their text in the right pane.",
            r'Run git -C "C:\New Delphi Projects\Echurchsite Analytical - Component Test" status --short again and compare it with the pre-generation result.',
        ],
        [
            r"Kit root: C:\New Delphi Projects\Delphi App Translation\export\component-integration\WebsiteAnalytics.",
            "The kit contains README.txt, component-integration.json, Deploy-LanguagePacks.ps1, ComponentSource, and Localization\\Languages.",
            "The replaceable per-project kit contains no DelphiPackages folder, no BPL, and no automatic installer. Delphi's registered package comes from the Studio's stable bin\\packages\\Win32\\Release folder.",
            "The language folder contains the validated target pack plus an automatically generated en-US.json source pack.",
            "Target Git status is unchanged by Build Integration Plan and Generate Component Kit.",
            "Apply, Restore, Complete Reset, automatic target build, and source-authorization controls are hidden in recommended mode.",
        ],
    )

    document.add_heading("7.1 Expected component-kit contents", level=2)
    add_matrix(
        document,
        ["Path below WebsiteAnalytics kit", "Purpose"],
        [
            [r"README.txt", "Ordered Object Inspector, Search Path, deployment, and build instructions."],
            [r"component-integration.json", "Application identity, FMX framework, manager/selector classes, language metadata, and scanner form roots."],
            [r"Deploy-LanguagePacks.ps1", "Copies only JSON packs to a selected executable directory."],
            [r"Localization\Languages\en-US.json", "Generated source-language pack for deterministic return to English."],
            [r"Localization\Languages\<locale>.json", "Validated translated runtime pack."],
            [r"ComponentSource\DAT.Components.Core.pas", "Framework-neutral manager core."],
            [r"ComponentSource\DAT.Components.FMX.pas", "FMX lifecycle adapter and manager."],
            [r"ComponentSource\DAT.Components.FMX.LanguageSelector.pas", "Optional typed FMX language combo box."],
            [r"ComponentSource\DAT.Runtime.*.pas", "Pack discovery, preferences, runtime manager, and FMX property application."],
        ],
        [4100, 5260],
        8.9,
    )

    add_test_case(
        document,
        "TC-10",
        "Install the FMX design package through Delphi",
        "Use RAD Studio's standard package manager to make the manager and selector available without changing the target.",
        [
            r"In the Translation Studio Integration page, choose Show Design BPL. Confirm File Explorer selects C:\New Delphi Projects\Delphi App Translation\bin\packages\Win32\Release\DATLanguageManagerFMXDesign.bpl.",
            "Start or activate RAD Studio without opening the disposable target project or a target form.",
            "Choose Component > Install Packages, then choose Add.",
            "Browse to the exact BPL selected by Show Design BPL, select DATLanguageManagerFMXDesign.bpl, and choose Open. Never use Component > Install Component and never select DATLanguageManagerFMXDesign.dpk.",
            "Confirm DAT Language Manager FireMonkey design-time package is listed and checked. Confirm the normal Embarcadero package list remains present and Embarcadero FMX Components is listed and checked. Choose OK.",
            "Open an FMX form and confirm DAT Localization appears in the Tool Palette.",
        ],
        [
            "TDATFMXLanguageManager and TDATFMXLanguageComboBox are available.",
            "The stable Studio bin\\packages path remains the registered package location; the Studio has not copied any BPL into a Delphi system or public Bpl folder.",
            "The design BPL has no custom DAT runtime-BPL import, so it loads without a missing-module error.",
            "The complete standard package list remains available; standard controls such as TTabControl stream normally.",
            "Closing and reopening RAD Studio produces no package-load error.",
            "No package is installed for VCL unless a VCL target is being tested.",
        ],
    )
    add_callout(
        document,
        "Do not register a generated-kit path.",
        "Per-project component kits are replaceable output. Installing a BPL from a kit would make Delphi depend on a volatile path and could prevent clean kit regeneration. Always use Show Design BPL and the stable Studio bin\\packages\\Win32\\Release file.",
    )

    add_test_case(
        document,
        "TC-11",
        "Perform the visible Delphi integration",
        "Add one manager and one optional selector through normal Form Designer operations in the disposable target copy.",
        [
            r"Open C:\New Delphi Projects\Echurchsite Analytical - Component Test\WebsiteAnalytics.dproj in RAD Studio.",
            r"In Project Options, add C:\New Delphi Projects\Delphi App Translation\export\component-integration\WebsiteAnalytics\ComponentSource to the Delphi Search Path for all configurations and both Windows platforms.",
            "Open WebsiteAnalytics.MainForm.fmx (frmMainDashboard) in the FMX Form Designer.",
            "Drop one TDATFMXLanguageManager on the primary form.",
            "Set ApplicationId to WebsiteAnalytics, LanguagesFolder to Localization\\Languages, and SourceLanguage to en-US.",
            "Leave AutoLoadPreferred, AutoTranslateOwner, AutoTranslateNewForms, ReapplyOpenForms, and PreserveControlState True for the acceptance test.",
            "Place one TDATFMXLanguageComboBox in an appropriate visible header/menu area.",
            "Set its LanguageManager property to the manager. Leave AutoPopulate and ShowLanguageCode True.",
            "Save the project and inspect Git diff before building.",
        ],
        [
            "Only normal developer-authored integration changes appear: project search-path metadata if saved there, the primary form resource/component field, applicable uses entries, and the visible selector layout.",
            "No ordinary secondary form receives a manager component.",
            "No Studio-generated DPR startup call or translation unit is added in Component Integration mode.",
            "Every change is readable in Git and reversible before commit.",
        ],
    )
    add_callout(
        document,
        "Ninety-form projects.",
        "One manager on the primary form covers ordinary FMX forms through the additive before-show lifecycle adapter. The developer does not place a component on every form. Use FormIdentityMappings only for inherited or unusually renamed forms whose runtime identity differs from the scanner's form root.",
    )

    add_test_case(
        document,
        "TC-12",
        "Build GA4 Win32 and Win64",
        "Verify the component-first disposable target builds normally for both supported architectures.",
        [
            "In RAD Studio select Win32 / Debug and build WebsiteAnalytics.",
            "Select Win64 / Debug and build WebsiteAnalytics.",
            r"Confirm C:\New Delphi Projects\Echurchsite Analytical - Component Test\bin\Win32\Debug\WebsiteAnalytics.exe exists.",
            r"Confirm C:\New Delphi Projects\Echurchsite Analytical - Component Test\bin\Win64\Debug\WebsiteAnalytics.exe exists.",
            "Resolve any search-path problem by checking the ComponentSource path for the affected platform/configuration; do not copy random DCUs into the target.",
        ],
        [
            "Both builds finish with no compiler or linker errors.",
            "The target contains the manager and selector through normal compiled source/designer resources.",
            "No provider credential is linked into either executable.",
        ],
    )

    add_test_case(
        document,
        "TC-13",
        "Deploy language packs",
        "Place identical local JSON assets beside each executable without copying development catalogs or credentials.",
        [
            "Open PowerShell.",
            r"Run the kit deployment script for C:\New Delphi Projects\Echurchsite Analytical - Component Test\bin\Win32\Debug.",
            r"Run the same script for C:\New Delphi Projects\Echurchsite Analytical - Component Test\bin\Win64\Debug.",
            "Inspect both Localization\\Languages folders.",
        ],
        [
            "Each executable directory contains Localization\\Languages\\en-US.json and the translated <locale>.json.",
            "No Localization\\Development folder is deployed beside the executable.",
            "No API key, provider settings JSON, or credential artifact is deployed.",
        ],
    )
    document.add_paragraph("Use these exact commands, substituting the generated kit only if the project name changes:", style="Heading 3")
    add_code(
        document,
        [
            r'& "C:\New Delphi Projects\Delphi App Translation\export\component-integration\WebsiteAnalytics\Deploy-LanguagePacks.ps1" -ApplicationDirectory "C:\New Delphi Projects\Echurchsite Analytical - Component Test\bin\Win32\Debug"',
            r'& "C:\New Delphi Projects\Delphi App Translation\export\component-integration\WebsiteAnalytics\Deploy-LanguagePacks.ps1" -ApplicationDirectory "C:\New Delphi Projects\Echurchsite Analytical - Component Test\bin\Win64\Debug"',
        ],
    )

    document.add_heading("8. Runtime Acceptance", level=1)
    add_test_case(
        document,
        "TC-14",
        "First launch and source language",
        "Verify a fresh profile opens normally and the English source pack is available.",
        [
            r"Remove only the disposable test preference file, if present: %LOCALAPPDATA%\WebsiteAnalytics\language.ini.",
            r"Run C:\New Delphi Projects\Echurchsite Analytical - Component Test\bin\Win32\Debug\WebsiteAnalytics.exe.",
            "Inspect the title, primary controls, language selector, date-range control, and application data.",
            "Open representative secondary forms such as Settings, Property Manager, and Diagnostics.",
        ],
        [
            "The application opens without EArgumentNilException or 'An FMX form is required.'",
            "The source language is English and the selector lists canonical native names and locale codes.",
            "No malformed or mojibake language names appear.",
            "Secondary forms open normally and remain in the source language until another selection is made.",
        ],
    )

    add_test_case(
        document,
        "TC-15",
        "Immediate language switching",
        "Verify a selection applies immediately to every open form without restarting.",
        [
            "Keep the primary form and at least one secondary form open.",
            "Select the translated language in TDATFMXLanguageComboBox.",
            "Observe the application title, headings, labels, buttons, menu/list text, hints where practical, and the already-open secondary form.",
            "Open another secondary form after the language change.",
            "Switch back to en-US.",
        ],
        [
            "All currently open eligible forms refresh immediately.",
            "A form opened after the change appears translated before first paint.",
            "Switching to en-US restores deterministic English from en-US.json.",
            "No application restart is required for either direction.",
        ],
    )

    add_test_case(
        document,
        "TC-16",
        "Preserve control state and date range",
        "Verify translation refresh does not clear user input or reset the GA4 date-range selection.",
        [
            "Choose a date range other than the default, for example Yesterday or Last 30 days.",
            "Enter representative editable text in any safe test field and select part of it if available.",
            "Select the translated language, then switch back to English.",
            "Refresh or run a normal report action if it is safe in the test environment.",
        ],
        [
            "The date-range combo remains on the same logical selection and does not become blank.",
            "The reporting period does not silently change to This year or another default.",
            "Writable edit/memo content, focus, selections, and list/combo ItemIndex remain intact.",
            "Read-only instructional content is translated where cataloged.",
        ],
    )

    add_test_case(
        document,
        "TC-17",
        "Preference persistence and restart",
        "Verify the selected language is saved per user and safely applied on the next launch.",
        [
            "Select the translated language and close Website Analytics normally.",
            r"Open %LOCALAPPDATA%\WebsiteAnalytics\language.ini and confirm Selected=<locale> without editing it.",
            "Restart the same Win32 executable.",
            "Repeat the translated selection, close, and launch the Win64 executable.",
        ],
        [
            "The next launch opens in the saved language without an exception.",
            "Win32 and Win64 share the ApplicationId preference as intended.",
            "The preference file contains only language selection data and no provider key.",
            "Selecting English and closing changes the saved value back to en-US.",
        ],
    )

    add_test_case(
        document,
        "TC-18",
        "Offline runtime",
        "Prove the finished application needs no network, provider account, or Studio process.",
        [
            "Close Delphi App Translation Studio.",
            "Disconnect the test computer from the Internet or block the disposable WebsiteAnalytics executable through the test firewall policy.",
            "Launch WebsiteAnalytics, switch between English and the translated language, open secondary forms, and restart.",
            "Restore the network after the test.",
        ],
        [
            "Both languages load from local JSON packs.",
            "Switching and persistence continue to work offline.",
            "No provider dialog, API-key prompt, network delay, or cloud dependency appears in the target.",
        ],
    )

    add_test_case(
        document,
        "TC-19",
        "Missing and malformed pack behavior",
        "Verify a deployment fault fails safely and does not corrupt the application.",
        [
            "Close the target.",
            "Move the translated JSON file to a temporary folder outside Localization\\Languages; do not delete it.",
            "Launch the target and inspect the selector and source-language behavior.",
            "Restore the file, then make a disposable copy with invalid JSON and place that copy under an unrelated filename in the language folder.",
            "Launch again, confirm the invalid unrelated pack is ignored, then remove the disposable invalid file.",
        ],
        [
            "The unavailable language is not offered as a valid selection.",
            "The application remains usable in the source language according to MissingPackBehavior.",
            "A malformed/unrelated pack does not crash startup or replace valid packs.",
            "Restoring the valid file and relaunching restores the language choice.",
        ],
    )

    document.add_heading("8.1 Repeat the runtime matrix", level=2)
    add_matrix(
        document,
        ["Executable", "Required abbreviated checks", "Result"],
        [
            [r"bin\Win32\Debug\WebsiteAnalytics.exe", "TC-14 through TC-19 in full.", "PASS / FAIL"],
            [r"bin\Win64\Debug\WebsiteAnalytics.exe", "Launch, switch both directions, date range/state, restart, offline.", "PASS / FAIL"],
            [r"bin\Win32\Release\WebsiteAnalytics.exe", "After Release build/deploy: launch, switch, restart, offline.", "PASS / FAIL / N/A"],
            [r"bin\Win64\Release\WebsiteAnalytics.exe", "After Release build/deploy: launch, switch, restart, offline.", "PASS / FAIL / N/A"],
        ],
        [3800, 4060, 1500],
        8.8,
    )

    document.add_heading("9. VCL Acceptance Delta", level=1)
    add_paragraphs(
        document,
        [
            "Repeat the component workflow with a disposable VCL project when formal VCL acceptance is required. Install DATLanguageManagerVCLDesign.bpl, generate a VCL component kit, place one TDATVCLLanguageManager on the primary VCL form, and optionally place TDATVCLLanguageComboBox. The component-source path and JSON deployment pattern are otherwise the same.",
        ],
    )
    add_matrix(
        document,
        ["VCL scenario", "Expected behavior"],
        [
            ["Primary and existing forms", "Manager applies the active pack and discovers ordinary forms through its private TApplicationEvents lifecycle."],
            ["Modal form", "Manager handles the modal display boundary before normal interaction."],
            ["Dynamically created modeless form", "VCL has no public additive before-show notification for every such form. If first-paint translation must be flicker-free, call LanguageManager.ApplyToForm(NewForm) after construction and before NewForm.Show."],
            ["State protection", "Writable edits/memos, focus, selection, and list/combo indexes remain intact when switching."],
        ],
        [2900, 6460],
        9.0,
    )

    document.add_heading("10. Studio Self-Translation", level=1)
    add_paragraphs(
        document,
        [
            "The Studio can scan and translate its own project. The automated self-localization test uses samples\\StudioSelfLocalization\\it-IT.json, temporarily deploys it under the Studio project Localization\\Languages folder, writes %LOCALAPPDATA%\\DelphiAppTranslationStudio\\language.ini, launches all four Studio executables, expects the Italian title, and restores any previous files in a finally block.",
        ],
    )
    add_code(
        document,
        [
            r'cd "C:\New Delphi Projects\Delphi App Translation"',
            r'powershell.exe -ExecutionPolicy Bypass -File .\tools\tests\RunStudioSelfLocalizationSmokeTest.ps1',
        ],
    )
    add_bullets(
        document,
        [
            "Expected title: Studio di traduzione app Delphi.",
            "Expected matrix: Win32/Win64 Debug/Release all Passed=True.",
            "Expected final line: Studio self-localization smoke tests passed.",
            "The script must restore the prior Italian pack and language preference or remove its temporary files when none existed.",
        ],
    )

    document.add_heading("11. Starting Over, Restore, and Reset", level=1)
    document.add_heading("11.1 Recommended component path", level=2)
    add_paragraphs(
        document,
        [
            "Component Integration performs no automatic target write, so its clean restart procedure is deliberately ordinary and transparent. In the disposable target, close the application, use Git to inspect/revert the manually added manager, selector, uses entries, and search-path change, then remove only the generated Localization data and regenerated kit after confirming the exact paths. Keep the pristine GA4 reference untouched.",
            "Regenerating a component kit replaces generated files under the Studio export root but does not clean or modify a previously integrated target. Deployment copies the current kit JSON to each executable folder; remove obsolete locale JSON files explicitly when a language is retired.",
        ],
    )
    add_callout(
        document,
        "Do not use Complete Reset for the recommended component path.",
        "The Studio's Complete Reset control belongs to Automatic Source Integration (Advanced), because only that mode created an automatic baseline and transaction manifest. Component-mode changes are developer-authored IDE changes and should be reverted through Git or the Form Designer.",
    )

    document.add_heading("11.2 Advanced automatic-source reset test", level=2)
    add_steps(
        document,
        [
            "Use only a disposable target that previously completed Automatic Source Integration (Advanced).",
            "Open and scan it in the Studio, choose Integration, then select Automatic Source Integration (Advanced).",
            "Choose Prepare Complete Reset. Confirm the preview is read-only and identifies the original baseline.",
            "Review the listed project/source restoration and Localization cleanup actions.",
            "Check I reviewed this reset plan and authorize Complete Reset, then choose Reset Project.",
            "Confirm the Studio creates a separate SHA-256-verified Complete Reset Safety backup before mutation.",
            "Build and run the restored target, and compare Git status/diff with the expected original baseline.",
        ],
    )
    add_bullets(
        document,
        [
            "The operation refuses to guess when the automatic-integration baseline is absent.",
            "On success, original pre-integration source is restored and project-local Localization development/runtime data is removed according to the preview.",
            "The safety backup is retained.",
            "On failure, the current project is restored from the safety backup and the failure is reported.",
        ],
    )

    document.add_heading("12. Negative and Recovery Tests", level=1)
    add_matrix(
        document,
        ["Condition", "Action", "Expected result"],
        [
            ["No API key", "Choose Translate Automatically.", "Studio moves to Provider Settings and explains which provider key is missing; catalog remains usable."],
            ["Wrong DeepL plan", "Use Free key with Pro endpoint or vice versa, then Test Connection.", "Connection fails clearly without exposing the key; correct plan can be selected and retried."],
            ["Provider offline/timeout", "Disconnect during a disposable provider test.", "Bulk translation reports failure; existing catalog data is retained and can be retried."],
            ["Validation error", "Blank a translation or break a placeholder.", "Validation identifies the entry and Export remains blocked until corrected."],
            ["Wrong application pack", "Place another application's JSON beside the target.", "Discovery rejects it by applicationId and does not offer it as a valid language."],
            ["Missing component source path", "Build a disposable target with the path omitted.", "Compiler identifies missing DAT units; adding the exact generated ComponentSource path resolves the problem."],
            ["Missing deployed packs", "Launch with no Localization\\Languages beside the executable.", "Application remains usable in source text according to configured fallback; no cloud call occurs."],
            ["Corrupt preference", "Use a disposable malformed language.ini.", "Application must not crash; source/default behavior remains recoverable. Restore or remove the test file afterward."],
        ],
        [1850, 3150, 4360],
        8.5,
    )

    document.add_heading("13. Performance and Large-Project Checks", level=1)
    add_paragraphs(
        document,
        [
            "Do not impose one fixed scan-time pass threshold across all computers. Record form count, source count, entry count, files scanned, elapsed milliseconds, processor, storage type, and whether antivirus or indexing was active. Compare repeat runs on the same machine and investigate substantial regressions.",
            "For a 90-form project, the acceptance goal is one manager on the primary form, not 90 components. Scan and kit generation should remain read-only. Open a representative set of startup, modal, modeless, inherited, settings, data-entry, and report forms; verify stable identities and translations. Use FormIdentityMappings only where the runtime class/instance does not match the scanner root.",
        ],
    )
    add_matrix(
        document,
        ["Metric", "Run 1", "Run 2", "Run 3", "Notes"],
        [
            ["Forms / source files", "", "", "", ""],
            ["Translatable entries", "", "", "", ""],
            ["Scan milliseconds", "", "", "", ""],
            ["Provider strings / batches", "", "", "", ""],
            ["Provider elapsed time", "", "", "", ""],
            ["Kit generation time", "", "", "", ""],
        ],
        [2200, 1250, 1250, 1250, 3410],
        8.8,
    )

    document.add_heading("14. Defect Recording", level=1)
    add_bullets(
        document,
        [
            "Record the exact Studio and target executable path, platform, configuration, and timestamp.",
            "Record the selected project path, framework, locale, provider, and DeepL plan when applicable. Never record the API key.",
            "Capture the complete status-panel message and any Windows/RAD Studio dialog text.",
            "Attach the stable catalog key, source text, translated text, status, origin, and validation message for translation defects.",
            "Attach Git status and the smallest relevant diff for integration defects.",
            "For runtime state defects, record the control name, state before switching, selected languages, state after switching, and whether restart was involved.",
            "State whether the defect reproduces on Win32, Win64, Debug, Release, VCL, and/or FMX.",
        ],
    )
    add_matrix(
        document,
        ["Field", "Record"],
        [
            ["Defect ID / title", ""],
            ["Date / tester", ""],
            ["Studio executable", ""],
            ["Target executable", ""],
            ["Project / framework", ""],
            ["Locale / provider", ""],
            ["Steps to reproduce", ""],
            ["Expected result", ""],
            ["Actual result", ""],
            ["Attachments / logs", ""],
        ],
        [2400, 6960],
        9.0,
    )

    document.add_heading("15. Final Acceptance Checklist", level=1)
    add_matrix(
        document,
        ["Gate", "Acceptance requirement", "Result"],
        [
            ["Backup", "Pre-change product backup exists and was verified.", "PASS / FAIL"],
            ["Clean target", "Disposable GA4 copy began clean; pristine reference remained untouched.", "PASS / FAIL"],
            ["Release matrix", "RunPhase10ReleaseValidation.ps1 passed uninterrupted.", "PASS / FAIL"],
            ["UI", "All seven pages fit, navigate, and display professionally at tested resolutions.", "PASS / FAIL"],
            ["Scan", "Project detection and scan are correct, fast, and read-only.", "PASS / FAIL"],
            ["Provider", "Connection and automatic translation succeed; key remains secret.", "PASS / FAIL"],
            ["Catalog", "Development JSON auto-saves and preserves status/provenance.", "PASS / FAIL"],
            ["Validation", "Errors block export; corrected catalog passes.", "PASS / FAIL"],
            ["Runtime pack", "Compact JSON exists under target Localization\\Languages.", "PASS / FAIL"],
            ["Kit safety", "Component kit is complete and target hashes/status do not change during generation.", "PASS / FAIL"],
            ["IDE integration", "One manager and optional selector are designer-authored and reviewable.", "PASS / FAIL"],
            ["Builds", "GA4 builds Win32 and Win64; Release variants tested as required.", "PASS / FAIL"],
            ["Instant switching", "Open and new forms switch immediately and return to English.", "PASS / FAIL"],
            ["State", "Date range, editable values, selection, focus, and indexes survive switching.", "PASS / FAIL"],
            ["Persistence", "Selected locale survives a restart without startup exception.", "PASS / FAIL"],
            ["Offline", "Target switches languages and restarts with no Internet or Studio.", "PASS / FAIL"],
            ["VCL", "VCL delta and modeless-form boundary are accepted/documented.", "PASS / FAIL / N/A"],
            ["Documentation", "Guide results and defects are complete enough to reproduce.", "PASS / FAIL"],
        ],
        [1550, 6410, 1400],
        8.5,
    )
    add_matrix(
        document,
        ["Approval", "Name", "Signature", "Date"],
        [
            ["Test lead", "", "", ""],
            ["Developer", "", "", ""],
            ["Product owner", "", "", ""],
        ],
        [1900, 2400, 3160, 1900],
        9.0,
    )

    document.add_heading("16. Appendix A - Exact Test Scripts", level=1)
    add_matrix(
        document,
        ["Script", "Purpose"],
        [
            [r"tools\tests\RunPhase10ReleaseValidation.ps1", "Master supported release matrix."],
            [r"tools\tests\RunLanguageManagerPackageTests.ps1", "Debug/Release runtime/design packages and streaming/selector tests."],
            [r"tools\tests\RunLanguageManagerCoreTests.ps1", "Framework-neutral manager invariants."],
            [r"tools\tests\RunFMXLanguageManagerTests.ps1", "FMX lifecycle, switching, identities, and state."],
            [r"tools\tests\RunVCLLanguageManagerTests.ps1", "VCL lifecycle, modal/modeless boundary, switching, identities, and state."],
            [r"tools\tests\RunRuntimeSmokeTests.ps1", "Foundation, runtime, integration, build/deploy, and launch workflow."],
            [r"tools\tests\RunStudioLaunchSmokeTests.ps1", "Real Studio startup in all four configurations."],
            [r"tools\tests\RunStudioSelfLocalizationSmokeTest.ps1", "Italian Studio launch with preference/pack restoration."],
        ],
        [4200, 5160],
        8.9,
    )

    document.add_heading("17. Appendix B - What Each JSON File Is", level=1)
    add_matrix(
        document,
        ["File", "Role", "Deployment"],
        [
            [r"Localization\Development\WebsiteAnalytics.<locale>.translation-project.json", "Lossless working catalog with source text, translated text, keys, checksums, statuses, origins, locale formats, runtime classification, and review data.", "Developer source workspace only."],
            [r"Localization\Languages\<locale>.json", "Compact, validated runtime pack with application identity, language metadata, locale settings, strings, and source catalog checksum.", "Copy beside each executable."],
            [r"export\component-integration\WebsiteAnalytics\component-integration.json", "Generated kit manifest for application/framework/component/language/form-root setup.", "Developer reference; not required by runtime manager."],
            [r"%LOCALAPPDATA%\WebsiteAnalytics\language.ini", "Per-user last-selected locale.", "Created at runtime; never package as an application asset."],
        ],
        [3400, 4060, 1900],
        8.6,
    )
    add_callout(
        document,
        "JSON is the product format.",
        "CSV is optional translator interchange only. Automatic provider translation reads and updates the in-memory development catalog and saves the JSON catalog in place; the runtime consumes exported JSON packs.",
    )

    document.add_heading("18. Appendix C - Current Verified Baseline", level=1)
    add_paragraphs(
        document,
        [
            "As of August 9, 2026, the project had current Studio executables in all four supported build folders, package outputs for Win32/Win64 Debug/Release, passing component/lifecycle/streaming/runtime/Studio launch/self-localization suites, and a clean pristine GA4 reference repository on branch codex/component-manager-pilot. The older C:\\New Delphi Projects\\Echurchsite Analytical - Test folder contained prior automatic-source-integration changes and was intentionally excluded from the clean component acceptance path.",
            "This baseline is evidence, not a substitute for rerunning the tests. Record new executable timestamps, release-gate output, provider results, and target Git diffs for each acceptance session.",
        ],
    )

    finish_document(document, DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_document())
