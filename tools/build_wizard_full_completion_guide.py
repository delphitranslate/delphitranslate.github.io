from __future__ import annotations

from io import BytesIO
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_guides import (
    BLUE,
    GRAY,
    INK,
    ORANGE,
    PALE_BLUE,
    add_bullets,
    add_field,
    add_header_footer,
    add_paragraphs,
    add_steps,
    configure_page,
    finish_document,
    set_cell_shading,
    set_page_number_format,
    set_table_geometry,
    setup_styles,
)


PROJECT_ROOT = Path(__file__).resolve().parents[1]
GUIDES_DIR = PROJECT_ROOT / "docs" / "guides"
DOCX_PATH = GUIDES_DIR / (
    "Delphi App Translation Studio Wizard-to-Runtime Complete Procedure.docx"
)
ICON = (
    PROJECT_ROOT
    / "images and icons"
    / "DelphiAppTranslationStudio-Icon-Master-v2_150.png"
)
LAST_CHANGED = "August 12, 2026"


def set_run_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_cover(document: Document) -> None:
    section = document.sections[0]
    configure_page(section)
    add_header_footer(section, "", False)
    document.add_paragraph("")
    document.add_paragraph("")
    logo = document.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ICON.exists():
        logo.add_run().add_picture(str(ICON), width=Inches(1.45))
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Wizard-to-Runtime\nComplete Procedure")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Delphi App Translation Studio")
    accent = document.add_table(rows=1, cols=2)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(accent, [4680, 4680], indent=0)
    set_cell_shading(accent.cell(0, 0), BLUE)
    set_cell_shading(accent.cell(0, 1), ORANGE)
    for cell in accent.rows[0].cells:
        cell.height = Inches(0.08)
    document.add_paragraph("")
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Current pre-release test workflow\n"
        f"Last changed: {LAST_CHANGED}\n"
        "Windows - Delphi VCL and FireMonkey - Win32 and Win64\n"
        "Printable, start-to-finish operator procedure"
    )
    set_run_font(run, 11, GRAY)
    purpose = document.add_paragraph()
    purpose.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = purpose.add_run(
        "Every required transition is included: clean test copy, Wizard, "
        "in-Wizard localization review, automatic finalization, optional Studio verification, RAD "
        "Studio component placement, builds, deployment, and runtime testing."
    )
    set_run_font(run, 10, GRAY, italic=True)


def add_toc(document: Document) -> None:
    toc_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(toc_section)
    set_page_number_format(toc_section, "lowerRoman", 1)
    add_header_footer(toc_section, "Wizard-to-Runtime Complete Procedure", True)
    document.add_paragraph("Table of Contents", style="TOC Heading")
    toc = document.add_paragraph()
    toc.paragraph_format.space_after = Pt(0)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    field_run = toc.add_run()
    field_run._r.extend([begin, instruction, separate])
    contents = [
        ("1. Read This Before Starting", 1),
        ("2. Exact Product Files and Generated Locations", 2),
        ("3. Prepare a Completely Clean Test", 3),
        ("4. Start the Setup Wizard", 4),
        ("5. Single Wizard Pass - Create, Review, and Finalize", 4),
        ("6. Complete the In-Wizard Localization Review", 7),
        ("7. Automatic Resume - Finalize Reviewed Output", 8),
        ("8. Optional: Return to the Main Translation Studio", 9),
        ("9. Install the Design Package in RAD Studio", 11),
        ("10. Place and Configure the Components", 11),
        ("11. Verify Search Path and Build Configuration", 12),
        ("12. Build Win32 and Win64", 13),
        ("13. Runtime Test - First Launch and Switching", 13),
        ("14. Runtime Layout and Content Acceptance", 14),
        ("15. Diagnose Missing, Random, or Runaway Text", 15),
        ("16. Procedure After Later UI Changes", 15),
        ("17. Adding More Languages", 16),
        ("18. Final Completion Checklist", 17),
        ("19. What to Send Back After Testing", 18),
        ("20. Important Stop Conditions", 18),
    ]
    for index, (heading, page_number) in enumerate(contents):
        if index:
            toc.add_run().add_break()
        entry = toc.add_run(heading)
        set_run_font(entry, 9.2, INK)
        dots = max(4, 73 - len(heading))
        leader = toc.add_run(" " + "." * dots + " " + str(page_number))
        set_run_font(leader, 9.2, GRAY)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    toc.add_run()._r.append(end)
    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(content_section)
    set_page_number_format(content_section, "decimal", 1)
    add_header_footer(content_section, "Wizard-to-Runtime Complete Procedure", True)


def add_matrix(document, headers, rows, widths, font_size=9.0):
    if sum(widths) != 9360:
        raise ValueError("Table widths must total 9360 DXA")
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(repeat)
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, font_size, "FFFFFF", bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        cannot_split = OxmlElement("w:cantSplit")
        table.rows[-1]._tr.get_or_add_trPr().append(cannot_split)
        for index, value in enumerate(values):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                set_cell_shading(cells[index], "F6F9FC")
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(run, font_size, INK)
    set_table_geometry(table, widths)
    document.add_paragraph("")


def add_code(document, lines):
    for line in lines:
        paragraph = document.add_paragraph(style="Code Block")
        paragraph.add_run(line)


def make_information_icon() -> BytesIO:
    """Return an original orange information icon for shaded guide notes."""
    image = Image.new("RGBA", (256, 256), (0, 0, 0, 0))
    draw = ImageDraw.Draw(image)
    draw.ellipse((8, 8, 248, 248), fill=(242, 138, 27, 255))
    draw.ellipse((34, 34, 222, 222), fill=(255, 255, 255, 255))
    draw.ellipse((106, 60, 150, 104), fill=(0, 0, 0, 255))
    draw.rounded_rectangle((108, 116, 148, 200), radius=10,
                           fill=(0, 0, 0, 255))
    stream = BytesIO()
    image.save(stream, format="PNG")
    stream.seek(0)
    return stream


def add_info_callout(document, title, body):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [760, 8600])
    icon_cell, text_cell = table.rows[0].cells
    for cell in (icon_cell, text_cell):
        set_cell_shading(cell, PALE_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    icon_paragraph = icon_cell.paragraphs[0]
    icon_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    icon_paragraph.add_run().add_picture(make_information_icon(),
                                         width=Inches(0.38))
    paragraph = text_cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    title_run = paragraph.add_run(title + " ")
    set_run_font(title_run, 10, INK, bold=True)
    body_run = paragraph.add_run(body)
    set_run_font(body_run, 10, INK)
    document.add_paragraph("")


def add_numbered_steps(document, items, start=1):
    """Add a numbered list that can resume after an intervening callout."""
    numbering = document.part.numbering_part.element
    style_num_id = int(document.styles["List Number"].element.pPr.numPr.numId.val)
    base_number = numbering.xpath(f'./w:num[@w:numId="{style_num_id}"]')[0]
    abstract_num_id = base_number.find(qn("w:abstractNumId")).get(qn("w:val"))
    num_ids = [int(element.get(qn("w:numId")))
               for element in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    for item in items:
        paragraph = document.add_paragraph(item, style="List Number")
        num_properties = paragraph._p.get_or_add_pPr().get_or_add_numPr()
        num_properties.get_or_add_ilvl().val = 0
        num_properties.get_or_add_numId().val = num_id


def add_checklist(document, items):
    for item in items:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.15)
        paragraph.paragraph_format.first_line_indent = Inches(-0.15)
        paragraph.add_run("[  ] ").bold = True
        paragraph.add_run(item)


def build_document() -> Path:
    document = Document()
    setup_styles(document)
    add_cover(document)
    add_toc(document)

    document.add_heading("1. Read This Before Starting", level=1)
    add_paragraphs(document, [
        "The Setup Wizard performs the first project scan, automatic provider translation, validation, JSON export, component-kit generation, project-local dependency preparation, and direct language-pack deployment. It does not place components on a Delphi form; that remains a normal RAD Studio Form Designer operation.",
        "The current workflow uses one protected Wizard processing pass. It creates the translated development catalog, automatically opens Localization Review, waits while the developer records terminology and layout decisions, and resumes automatically when the Review Center closes. The resumed pass applies those decisions before final validation, runtime JSON export, component-kit generation, atomic dependency-folder refresh, a transient-path Win32 Release build, and direct deployment.",
        "After the single Wizard pass, the main Studio may be used for optional detailed inspection or manual correction. RAD Studio is then used to install the design package, place the manager and selector, build the target, and test runtime behavior.",
    ])
    add_info_callout(document, "Do not use an original project.",
        "Create a new disposable test folder from the pristine copy. Keep the pristine folder and the original application untouched. The Wizard creates its own ZIP backup and preserves the previous managed dependency folder before replacement, but those are additional safeguards, not substitutes for the pristine copy.")
    add_info_callout(document, "Current automatic-layout boundary.",
        "The runtime can apply accepted, checksum-backed Width, Height, WordWrap, and AutoSize rules from a language pack. It does not automatically move neighboring controls, redesign a form, change fonts, or guarantee that a wider control will not overlap another control. Any proposal left Pending, Rejected, or Manual is not applied automatically.")
    add_info_callout(document, "Current glossary boundary.",
        "An empty Project Glossary is normal before terms are approved. Built-in computer-interface terminology still operates. Catalog-derived suggestions appear after translations exist; only approved project terms become authoritative for that project and language.")

    document.add_heading("1.1 What this test should prove", level=2)
    add_checklist(document, [
        "The Wizard completes without an access violation or a stopped-processing message.",
        "The Wizard translates unresolved entries automatically with the selected API provider.",
        "The protected Wizard pass pauses for Localization Review before its definitive export.",
        "Closing Localization Review resumes the same Wizard pass and embeds accepted glossary/layout decisions.",
        "The main Studio opens the same catalog, validates it, and exports the final JSON pack.",
        "The target project contains only intentional designer changes; the Studio leaves its Pascal, DPR, DPROJ, DFM, and FMX files byte-for-byte unchanged.",
        "Win32 and Win64 builds receive the correct Localization\\Languages folder.",
        "The target starts in the expected language, switches immediately, restores layout on switching back, and remembers the user's choice after restart.",
        "Dynamic application text does not produce repeated characters.",
    ])

    document.add_heading("2. Exact Product Files and Generated Locations", level=1)
    add_matrix(document, ["Purpose", "Path or pattern"], [
        ["Recommended Studio build", r"C:\DelphiProjects\Delphi App Translation\bin\Win32\Debug\DelphiAppTranslationStudio.exe"],
        ["Win64 Studio build", r"C:\DelphiProjects\Delphi App Translation\bin\Win64\Debug\DelphiAppTranslationStudio.exe"],
        ["FMX design package", r"C:\DelphiProjects\Delphi App Translation\bin\packages\Win32\Release\DATLanguageManagerFMXDesign.bpl"],
        ["VCL design package", r"C:\DelphiProjects\Delphi App Translation\bin\packages\Win32\Release\DATLanguageManagerVCLDesign.bpl"],
        ["Development catalog", r"<Test Project>\Localization\Development\<ApplicationId>.<language>.translation-project.json"],
        ["Project glossary", r"<Test Project>\Localization\Glossaries\<ApplicationId>.<language>.glossary.json"],
        ["Runtime language pack", r"<Test Project>\Localization\Languages\<language>.json"],
        ["Component kit", r"C:\DelphiProjects\Delphi App Translation\export\component-integration\<ApplicationId>"],
        ["Localization review", r"C:\DelphiProjects\Delphi App Translation\export\localization-review\<ApplicationId>\<language>"],
        ["Saved language preference", r"%LOCALAPPDATA%\<ApplicationId>\language.ini"],
        ["Wizard safety backup", r"%USERPROFILE%\Documents\Delphi App Translation Backups\<ApplicationId>\<timestamp>.zip"],
    ], [2550, 6810], 8.8)
    add_info_callout(document, "Application ID means project name, not folder name.",
        "For C:\\...\\ExampleApp.dproj the detected ApplicationId is ExampleApp. Always use the exact Application ID displayed by the Wizard; do not type the full path and do not include .dproj or .exe.")

    document.add_heading("3. Prepare a Completely Clean Test", level=1)
    add_steps(document, [
        "Close the target application executable if it is running.",
        "Close the target project in RAD Studio. Saving is allowed before closing, but do not leave unsaved form changes.",
        "Locate the pristine project folder. Confirm that it is the known-good unlocalized source.",
        "Create a new sibling test folder by copying the entire pristine folder. Use a clear name such as <Application Name> - Translation Test.",
        "Do not copy an earlier test folder. Earlier Localization folders, dependency folders, developer-owned Search Paths, components, and language preference files would invalidate the test.",
        "Open the new test folder in File Explorer. Confirm the expected .dproj or .dpr file is present.",
        "If the pristine copy contains old bin or dcu output folders, remove those only from the new disposable test copy or use Delphi Clean before the baseline build. Never clean the pristine source as part of this test.",
        "Open the new test project's .dproj in RAD Studio, perform a baseline Win32 Debug build, and run it once in its original language.",
        "Open every important form or page and confirm that the pristine application itself does not already contain mixed-language or runaway dynamic text.",
        "Close the baseline executable and choose File > Save All in RAD Studio. The project may remain open during Wizard final processing.",
    ])
    add_info_callout(document, "No Git requirement.",
        "Git is useful evidence but is not required. A pristine copy plus a fresh disposable test copy is an accepted safety baseline. If Git is available, record git status --short now; it should be empty.")

    document.add_heading("3.1 Important: Clear a Previous Saved-Language Preference", level=2)
    add_paragraphs(document, [
        "A new project folder does not clear the per-user language preference because that file lives under Local AppData. If the same ApplicationId was tested previously, the new executable may immediately restore Spanish or another language even though the project folder is new.",
    ])
    add_steps(document, [
        "Press Windows+R.",
        "Enter %LOCALAPPDATA% and press Enter.",
        "Open the folder whose name exactly matches the Wizard's Application ID, if it exists.",
        "If language.ini exists, rename it to language.ini.pretest. Do not rename unrelated application settings.",
        "If the Application ID folder does not exist, no preference cleanup is required.",
    ])

    document.add_heading("4. Start the Setup Wizard", level=1)
    add_steps(document, [
        r"Run C:\DelphiProjects\Delphi App Translation\bin\Win32\Debug\DelphiAppTranslationStudio.exe.",
        "Wait for the Delphi App Translation Studio main form to appear.",
        "Click Start Setup Wizard on the Project page. Do not click Open Project for this initial path.",
        "Confirm the dimmed Studio background is blocked by the modal Translation Setup Wizard.",
        "Do not run a second Studio instance during this test.",
    ])

    document.add_heading("5. Single Wizard Pass - Create, Review, and Finalize", level=1)
    document.add_heading("5.1 Step 1 - Welcome", level=2)
    add_steps(document, [
        "Read the safety statement. It explains that final processing creates a backup and writes localization/configuration files but does not rewrite Pascal or form source.",
        "Click Next.",
    ])

    document.add_heading("5.2 Step 2 - Delphi Project", level=2)
    add_steps(document, [
        "Click Browse.",
        "Navigate to the new disposable test folder, not the pristine folder and not the original project folder.",
        "Select the target .dproj. Use the .dpr only when no .dproj exists.",
        "Confirm the displayed project path is the disposable test folder.",
        "Confirm the detected framework is FireMonkey or VCL as expected.",
        "Confirm Win32 and/or Win64 platforms and the form-resource count are plausible.",
        "Write down the detected Application ID. You will use this exact value in the Object Inspector later.",
        "Leave the workflow on Automatic. For a new folder it should report Create a new translation.",
        "Click Next.",
    ])

    document.add_heading("5.3 Step 3 - Deployment Destinations", level=2)
    add_paragraphs(document, [
        "The Wizard detects ordinary Win32 and Win64 build-output folders automatically. Use this page only for separate installed, portable, network, or USB application folders. The choices remain staged until authorized final processing.",
    ])
    add_steps(document, [
        "Leave the list empty when only normal Delphi build-output folders are needed.",
        "For each separate application copy, click Add Application Folder and select the full folder containing—or intended to contain—the executable. Include the correct drive letter for USB or other non-C drives.",
        "Do not select the Localization or Languages subfolder.",
        "Remove any obsolete or incorrect destination before continuing.",
        "Click Next. Cancel remains safe because the staged destinations are not saved until authorized final processing.",
    ])
    add_info_callout(document, "Automatic deployment.",
        "Final processing deploys immediately to every configured destination that is currently available. Later Studio runs repeat direct deployment to remembered destinations. An unplugged drive is skipped with a warning and does not fail processing. Ordinary IDE builds do not receive a Studio-written post-build target.")

    document.add_heading("5.4 Step 4 - Languages", level=2)
    add_steps(document, [
        "Set Source language to the language in which the Delphi application was authored, normally English (United States) [en-US].",
        "Select exactly one Target language for this pass, for example Spanish (Spain) [es-ES].",
        "Confirm Native language name is correct and uses proper characters, for example Español.",
        "Confirm Direction is Left-to-right unless the language requires right-to-left text.",
        "Review the date, time, decimal, thousands, and currency values. These locale settings are stored in the JSON pack; correct only values you deliberately want the target application to use.",
        "Confirm Automatic still reports Create a new translation, then click Next.",
    ])
    add_info_callout(document, "One target language per Wizard run.",
        "To add more languages, repeat the create/update workflow for each target language. The runtime selector later lists the source-language pack and every valid target-language JSON pack deployed beside that executable.")

    document.add_heading("5.5 Step 5 - Translation Service", level=2)
    add_steps(document, [
        "Choose Google Cloud Translation or DeepL.",
        "For DeepL, choose API Free or API Pro to match the account. Google does not use the DeepL plan field.",
        "If the provider key is already stored in Windows Credential Manager, leave the API key field blank and verify that the status reports a stored key.",
        "If no key is stored, paste the API key into the masked field.",
        "Leave Remember securely on this computer checked to store the key in Windows Credential Manager. Uncheck it only for a session-only key.",
        "Click Save / Replace Key when a new key was entered.",
        "Click Test Connection.",
        "Do not continue until the status explicitly reports that the connection test passed.",
        "Click Next.",
    ])
    add_info_callout(document, "The target application remains offline.",
        "Only the developer's Studio uses the provider and API key. The deployed target application reads local JSON packs and does not need the Internet or the API key.")

    document.add_heading("5.6 Step 6 - Scan Project", level=2)
    add_numbered_steps(document, [
        "Click Scan Project.",
        "Wait for Scan complete. Do not click Next while the scan is running.",
        "Record the total translatable entries and the new/changed/unchanged/obsolete counts.",
        "Scroll through representative items. Confirm form text and resourcestring entries belong to the selected project.",
        "Look specifically for suspicious repeated characters, data values, paths, filenames, URLs, IDs, or logging output. These should be protected or reviewed rather than translated as normal interface text.",
    ])
    add_matrix(document, ["Record", "Value observed", "Notes"], [
        ["Scan date and time", "", ""],
        ["Disposable test-project path", "", ""],
        ["Application ID", "", ""],
        ["Source and target languages", "", ""],
        ["Total translatable entries", "", ""],
        ["New / changed / unchanged / obsolete", "", ""],
        ["Suspicious or excluded text noted", "", ""],
    ], [3150, 2450, 3760], 8.6)
    add_info_callout(document, "Rescan after source changes.",
        "If you edit or save any target PAS, FMX, DFM, DPR, or DPROJ file after this scan, return to this step and scan again before final processing.")
    add_numbered_steps(document, [
        "Localization Review is intentionally unavailable at this point because layout analysis requires translated text. Continue with Next. During final processing, the Wizard translates the unresolved entries and then opens Localization Review automatically with terminology suggestions and applicable layout proposals.",
    ], start=6)

    document.add_heading("5.7 Step 7 - Delphi Component", level=2)
    add_steps(document, [
        "Read the project-specific instructions from top to bottom.",
        "Confirm the instructions show the correct target project path and exact Application ID.",
        "Click Show Design BPL. File Explorer should select DATLanguageManagerFMXDesign.bpl for an FMX target or DATLanguageManagerVCLDesign.bpl for a VCL target under bin\\packages\\Win32\\Release.",
        "Do not use Delphi's Install Component wizard. Do not select a .dpk. The later approved procedure is Component > Install Packages > Add and selection of the compiled design BPL.",
        "Close File Explorer and return to the Setup Wizard.",
        "Check Required: I understand the remaining manual RAD Studio phase.",
        "Confirm the orange reminder changes to a confirmation message.",
        "Click Next.",
    ])

    document.add_heading("5.8 Step 8 - Review and Authorize", level=2)
    add_steps(document, [
        "Read the project, Application ID, framework, target language, workflow, provider, scan count, unresolved count, integration method, deployment statement, and backup statement.",
        "Confirm the required backup box is checked. It is intentionally mandatory.",
        "Choose File > Save All in RAD Studio and close the running target application. The project itself may remain open.",
        "Check Required: target files are saved and the target application is closed.",
        "Check I reviewed these choices and authorize final processing.",
        "Click Begin Final Processing only after both confirmations are checked.",
        "After processing begins, do not try to close the Wizard or Studio. Back, Cancel, and the step rail remain disabled. The Wizard will translate, open Localization Review, and automatically resume finalization after Review closes.",
    ])

    document.add_heading("5.9 Step 9 - Processing and Completion", level=2)
    add_steps(document, [
        "Watch the progress log until it stops changing.",
        "Confirm a timestamped ZIP backup was created.",
        "Confirm unresolved entries were translated by the selected provider or resolved by terminology/translation memory.",
        "Confirm Development catalog saved appears.",
        "Confirm the log reports the number of separate application destinations saved, and confirms deployment to each destination currently available.",
        "Confirm the progress log reports that the required in-Wizard localization review is opening. The Translation Studio and Wizard remain protected while the modal Review Center is active.",
    ])

    document.add_heading("6. Complete the In-Wizard Localization Review", level=1)
    add_steps(document, [
        "Wait for the Localization Review Center to open automatically. Do not click Finish and do not start another Wizard.",
        "On Audit & Confidence, read the summary and inspect all High risk findings and representative Warning findings.",
        "Click Open Visual Review. The review package is generated automatically when the Review Center opens and is refreshed whenever a layout decision is saved. The browser shows the current translated geometry beside the proposed geometry; proposed changes are outlined in green. This preview does not alter a Delphi form.",
        "Return to the Localization Review Center.",
        "Open Terminology Suggestions. Select suggestions and read Source term, Suggested target, Context, Semantic concept, Provenance, Confidence, and Why suggested.",
        "Use Approve Selected only for a term you understand and want enforced for this project and language.",
        "Use Approve High-confidence All only after inspecting the proposed group. Provider output is not promoted automatically merely because it exists.",
        "Open Project Glossary. Confirm approved terms now appear. If a required term is absent, click New, enter Source term and Preferred translation, optionally enter Semantic concept and Developer note, leave Approved terminology checked, click Add / Update Term, and click Save Project Glossary.",
        "Open Layout Proposals. Select a proposal and read its form, control, property, current value, proposed value, reason, decision, and What happens explanation. Use Open Visual Review before accepting a group so the proposed result can be judged as a screen, not merely as numbers.",
        "For the ordinary test, click Accept All Safe Proposals only after understanding that Width, Height, WordWrap, and AutoSize rules will be applied at runtime for this language.",
        "For any proposal that would collide with another control or exceed its parent, choose Manual or Rejected and click Save.",
        "Click Close. This records the end of the review phase and returns control to the protected Wizard processing pass.",
    ])
    add_info_callout(document, "Do not start a second Wizard pass.",
        "Closing Localization Review now resumes the original processing pass. The Wizard applies the saved glossary and accepted layout rules before it performs the definitive validation, runtime-pack export, and component-kit generation.")

    document.add_heading("7. Automatic Resume - Finalize Reviewed Output", level=1)
    add_steps(document, [
        "After the Review Center closes, do not press anything while the progress log resumes automatically.",
        "Confirm the log reports Localization review closed and reviewed decisions being applied.",
        "Confirm Reviewed development catalog saved appears.",
        "Confirm the progress log lists the final localization review, layout proposal, and multilingual layout envelope paths. In File Explorer, open C:\\New Delphi Projects\\Delphi App Translation\\export\\localization-review\\<ApplicationId>\\<language> and verify that localization-review.html, layout-proposal.json, and multilingual-layout-envelope.json exist and are not zero bytes.",
        "Confirm validation passed. Warnings may remain; blocking errors must not remain.",
        "Confirm Runtime JSON pack exported appears with a path under the disposable test project's Localization\\Languages folder.",
        "Confirm Component integration kit generated appears.",
        "Confirm the managed dependency folder was refreshed, the transient-path Win32 Release build completed, direct pack deployment completed, and the target project file remained unchanged.",
        "Confirm the footer says Setup Wizard completed successfully. If it says stopped, do not continue to RAD Studio. Copy or photograph the entire progress log and footer message, close the Wizard only after processing has stopped safely, preserve the Wizard ZIP backup and disposable test folder, and send the exact STOPPED text and screenshots to the development team. Resume only after the cause is corrected and new test instructions are provided.",
        "Click the underlined blue component-kit path or Open Kit Folder. Confirm ComponentSource, Localization, component-integration.json, Deploy-LanguagePacks.ps1, README.txt, and Wizard-Completion-Report.txt exist.",
        "Read the repeat/troubleshooting explanation. No command or deployment button is required during a normal successful run. Final processing has already deployed detected build outputs and every available application destination entered earlier.",
        "Confirm the completion report states that review was completed inside the same Wizard processing pass before final export.",
        "Click Finish once. The Wizard closes and returns to the Studio main form.",
    ])

    document.add_heading("8. Optional: Return to the Main Translation Studio", level=1)
    add_paragraphs(document, [
        "The single Wizard pass has already translated, reviewed, validated, exported, and refreshed the project. If it completed successfully, validation reported no errors, and no additional translation, glossary, locale, or layout correction is needed, Section 8 is optional. You may stop using the Translation Studio here and proceed directly to Section 9 for the required RAD Studio installation and component-placement work.",
        "Use Section 8 when you want detailed catalog inspection, a manual translation correction, another validation report, or a final runtime-pack and component-kit refresh. Do not start a second translation from scratch. Any correction made here must be saved, validated, exported, and followed by the component-kit refresh in Section 8.4 before continuing to RAD Studio.",
    ])
    document.add_heading("8.1 Open the matching project and catalog", level=2)
    add_steps(document, [
        "On the Studio left rail, click 1 Project.",
        "Click Open Project and select the same disposable test .dproj.",
        "Click 3 Translate.",
        "Click Open.",
        "Open the development catalog from <Test Project>\\Localization\\Development. Its name is <ApplicationId>.<language>.translation-project.json.",
        "Confirm the target language, native language name, locale formats, entry count, translated count, status, and origin fields are populated.",
        "Click 2 Scan, then click Scan Project. This proves that the catalog still matches the saved Delphi sources.",
        "Return to 3 Translate. Confirm reviewed/approved translations and existing provider results were preserved.",
    ])

    document.add_heading("8.2 Inspect and correct translations", level=2)
    add_steps(document, [
        "Select entries with short or ambiguous interface text, unknown context, provider-basic confidence, or suspicious ownership.",
        "Read Source text, the context hint, runtime role, origin, and Translated text.",
        "Correct an inaccurate translation in Translated text and click Apply Translation. This changes the development JSON catalog, not Delphi source.",
        "Use Mark Reviewed only after a human has actually reviewed that entry. Use Approve only after it is already Reviewed.",
        "Review All and Approve All are catalog-wide decisions. Do not use them merely to remove warnings; use them only when one qualified review decision truly applies to the entire displayed group.",
        "Click Save after completing edits. Confirm Development catalog saved.",
        "Click Translate Automatically only if the readiness line or confirmation dialog reports unresolved entries. The command sends only eligible unresolved entries; reviewed and approved entries remain unchanged.",
    ])

    document.add_heading("8.3 Validate and export the final runtime pack", level=2)
    add_steps(document, [
        "Click 4 Validation.",
        "Click Run Validation.",
        "If errors are reported, double-click each error, correct the referenced translation or setting, save, and run validation again. Do not export with errors.",
        "Warnings do not block export. Review warnings about identical source/translation, placeholders, accelerator counts, runtime Translate calls, and suspicious text; document why any accepted warning is safe.",
        "When the summary reports 0 errors, click 5 Export.",
        "Click Export Runtime Pack.",
        "Confirm the runtime entry count and click the blue output path.",
        "In File Explorer, confirm the final <language>.json exists under the disposable test project's Localization\\Languages folder.",
        "Open the JSON in a text-safe viewer. Confirm schemaVersion is 3 and confirm a layout array is present when safe layout proposals were accepted.",
    ])

    document.add_heading("8.4 Refresh the component kit after the final export", level=2)
    add_steps(document, [
        "Click 6 Integration.",
        "Confirm Integration method is Component Integration (Recommended).",
        "Click Build Integration Plan.",
        "Confirm the target framework, translated pack count, generated English pack statement, manager class, and selector class are correct.",
        "Click Generate Component Kit.",
        "Click Open Kit Folder and confirm the kit was refreshed. This step is important because the kit must contain the same final runtime JSON pack that you just exported.",
        "Click Show Design BPL and leave File Explorer open with the exact Win32 Release design BPL selected.",
    ])

    document.add_heading("9. Install the Design Package in RAD Studio", level=1)
    add_info_callout(document, "Manual installation is intentional.",
        "The product does not automatically register a design-time BPL. Manual installation through RAD Studio's approved package dialog avoids changing the IDE behind the developer's back.")
    add_steps(document, [
        "Start RAD Studio without opening the target project or target form.",
        "Choose Component > Install Packages.",
        "If DAT Language Manager FireMonkey design-time package or the VCL equivalent is already listed and checked, do not add a duplicate. Click OK and continue to the next section.",
        "Otherwise click Add.",
        "Browse to the exact BPL selected by Show Design BPL: DATLanguageManagerFMXDesign.bpl for FMX or DATLanguageManagerVCLDesign.bpl for VCL.",
        "Select the .bpl and click Open. Do not choose a .dpk and do not use Component > Install Component.",
        "Confirm the DAT design-time package appears in the Design packages list and is checked.",
        "Click OK.",
        "Create or open a harmless blank form if necessary and confirm the Tool Palette contains DAT Localization with both the language manager and language combo box.",
    ])

    document.add_heading("10. Place and Configure the Components", level=1)
    add_steps(document, [
        "Open the disposable target project's .dproj in RAD Studio.",
        "Open the primary form in the Form Designer. If Delphi opened it automatically, simply confirm that the primary form is the active designer.",
        "From DAT Localization, place one TDATFMXLanguageManager for FMX or TDATVCLLanguageManager for VCL on the primary form.",
        "Select the manager in the Object Inspector.",
        "Set ApplicationId to the exact value displayed by the Wizard.",
        "Leave LanguagesFolder as Localization\\Languages.",
        "Set SourceLanguage to the source language code used by the Wizard, normally en-US.",
        "For a clean first-launch test, set AutoLoadPreferred to False. This forces the source language initially. After first-launch behavior passes, set it back to True to test saved-user preference behavior.",
        "Leave AutoTranslateOwner, AutoTranslateNewForms, ReapplyOpenForms, and PreserveControlState enabled unless this test intentionally exercises another setting.",
        "Place one TDATFMXLanguageComboBox or TDATVCLLanguageComboBox on the visible primary form. The selector is required unless the application supplies a connected Language menu.",
        "Select the combo box. In the Object Inspector, set LanguageManager to the manager component you just placed. Do not leave it blank.",
        "Position and size the combo box where it is visible and does not cover existing controls. Add a normal designer-authored label such as Language: if the application needs one.",
        "Choose File > Save All. This save is essential; it writes the two designer components and their properties to the form resource and updates the project metadata as required by Delphi.",
    ])
    add_info_callout(document, "One manager, not one per form.",
        "Place the manager on the primary form only. It observes and translates other open or newly created forms. Ordinary secondary forms do not need another manager component.")

    document.add_heading("11. Verify Search Path and Build Configuration", level=1)
    add_steps(document, [
        "Choose Project > Options.",
        "Select Building > Delphi Compiler > Search path.",
        "Inspect Value from All configurations. It should contain $(PROJECTDIR)\\dependencies\\DelphiAppTranslation\\source exactly once.",
        "Use the Target selector to inspect Debug/Release and Windows 32-bit/64-bit as applicable. The developer-owned entry should apply to every configuration and platform that builds the target.",
        "If the path is absent, add that exact relative path once without removing or replacing any existing Search Path entry. This is the one required developer-owned Project Options step; the Studio deliberately does not edit the DPROJ.",
        "Click OK to save the Project Options change, then use File > Save All.",
    ])

    document.add_heading("12. Build Win32 and Win64", level=1)
    add_steps(document, [
        "Select Win32 and Debug in Project Manager, then choose Project > Build <ApplicationId>.",
        "Confirm the build completes with zero fatal errors.",
        "Locate the Win32 Debug executable using Project > Options > Building > Delphi Compiler > Output directory if the project uses a nonstandard path.",
        "Beside the executable, confirm Localization\\Languages exists and contains en-US.json plus the target-language JSON file.",
        "Repeat for Win64 Debug.",
        "If release testing is required, repeat for Win32 Release and Win64 Release.",
        "After every build, verify the executable's own folder contains Localization\\Languages. The project-root Localization folder alone is not sufficient for runtime discovery.",
    ])
    add_info_callout(document, "The .rsm file is normal.",
        "A Delphi build may create both <ApplicationId>.exe and <ApplicationId>.rsm. The .exe is the application. The .rsm contains debug symbol information and is not a language pack.")

    document.add_heading("12.1 Manual deployment fallback", level=2)
    add_paragraphs(document, [
        "Use this only if direct Studio deployment did not create the language folder. Substitute the actual dependency and executable folder paths shown by the Wizard. Ordinary IDE builds do not contain a Studio-written post-build target.",
    ])
    add_code(document, [
        '& "C:\\Windows\\System32\\WindowsPowerShell\\v1.0\\powershell.exe" -NoProfile -ExecutionPolicy Bypass -File "C:\\New Delphi Projects\\Delphi App Translation\\export\\component-integration\\<ApplicationId>\\Deploy-LanguagePacks.ps1" -ApplicationDirectory "<Folder containing ApplicationId.exe>"',
    ])
    add_paragraphs(document, [
        "A successful command reports Language packs deployed to <application folder>\\Localization\\Languages. The PowerShell process should exit when the command finishes.",
    ])

    document.add_heading("12.2 Portable or USB executable", level=2)
    add_steps(document, [
        "Before final processing, enter the desired portable folder or USB drive on the Wizard's Deployment Destinations page, including the drive letter such as F:\\PortableApp.",
        "Copy or build the final executable into that configured folder. Final processing and future builds deploy its packs automatically whenever the destination is available.",
        "Confirm <Portable Folder>\\Localization\\Languages contains en-US.json and every target-language JSON file.",
        "Use Deploy New App Folder or the fallback command only for a new, temporary, or troubleshooting destination that was not configured earlier.",
        "Keep the manager's LanguagesFolder property relative as Localization\\Languages. Do not put F: or another drive letter in the component property.",
    ])

    document.add_heading("13. Runtime Test - First Launch and Switching", level=1)
    add_steps(document, [
        "Confirm the target application is closed.",
        "Confirm AutoLoadPreferred is False for this first-launch test, or confirm the prior language.ini was renamed as described in Section 3.1.",
        "Run the Win32 Debug executable directly from its output folder.",
        "Confirm the application initially displays its source language and the selector shows the source language.",
        "Open the selector. Confirm it contains the source language plus one item for every deployed valid target pack. Fifteen target packs should produce sixteen choices when the source pack is included.",
        "Select the target language.",
        "Confirm visible designer text changes immediately without restarting the application.",
        "Open each secondary form and visit each page or tab. Confirm forms created after the selection also appear in the active language.",
        "Return to the primary form and select the source language.",
        "Confirm source text and original layout values are restored.",
        "Repeat source-to-target-to-source switching at least ten times. Confirm controls do not grow cumulatively and the application does not lose dates, selections, focus, connection state, or editable values.",
        "Close the executable.",
        "Set AutoLoadPreferred back to True in the Object Inspector, Save All, and rebuild the configuration.",
        "Run the rebuilt executable, select the target language, close it, and run it again.",
        "Confirm it restores the user's last selected language. This is expected saved-preference behavior, not corruption of the source application.",
        "Repeat representative switching and restart checks with Win64 Debug and required Release builds.",
    ])

    document.add_heading("14. Runtime Layout and Content Acceptance", level=1)
    document.add_heading("14.1 What accepted layout rules should do", level=2)
    add_checklist(document, [
        "Accepted Width or Height rules apply only when their target language is active.",
        "Accepted WordWrap or AutoSize rules apply only to the identified control and language.",
        "Long label and check-box text may use WordWrap with a calculated Height instead of uncontrolled horizontal growth.",
        "Grid column proposals enlarge translated headings conservatively without resizing the entire grid beyond its parent.",
        "Switching away restores the original property value before another language is applied.",
        "Repeated language changes do not compound Width or Height.",
        "No target PAS, FMX, DFM, DPR, or DPROJ form layout is rewritten by a runtime layout rule.",
    ])
    document.add_heading("14.2 What still requires human inspection", level=2)
    add_checklist(document, [
        "A wider translated control does not overlap its neighbor or leave its parent container.",
        "Wrapped labels have enough height and remain aligned with related controls.",
        "Buttons, tabs, menu items, column headers, charts, grids, and status areas remain readable.",
        "Dynamic text produced by application code uses Translate or FormatTemplate where required.",
        "The translation is correct in application context, not merely a valid dictionary translation.",
        "Form titles, secondary forms, help text, and runtime status messages have each been considered separately.",
    ])

    document.add_heading("15. Diagnose Missing, Random, or Runaway Text", level=1)
    add_steps(document, [
        "Record the form, page, control, active language, exact displayed text, and a screenshot.",
        "Switch to the source language. Determine whether the same bad text exists in the pristine baseline application.",
        "If the problem exists in the pristine baseline, classify it as a target-application defect rather than a translation defect.",
        "If the source language is correct but the target language is wrong, search the development catalog for the control key or source text.",
        "If the entry exists, inspect its source ownership, runtime role, context, translation, and status. Correct the translation or runtime classification as appropriate.",
        "If the entry does not exist, save all target files, close the target project, reopen it in the Studio, scan again, and record whether the scan count increases.",
        "For runtime-generated text, verify that the target application uses the manager's Translate or FormatTemplate API with the catalog key. Static form scanning cannot translate arbitrary text assembled later by application code.",
        "For repeated o characters or other actively growing output, stop the executable. Record whether the source-language run also reproduces it. A JSON translation pack supplies a fixed string; unbounded growth usually indicates an application timer, refresh, append, or data-writing loop and must be isolated separately.",
        "After any catalog correction, Save, Run Validation, Export Runtime Pack, regenerate the component kit, rebuild/deploy, and retest.",
    ])

    document.add_heading("16. Procedure After Later UI Changes", level=1)
    add_info_callout(document, "Batch changes when practical.",
        "Add or revise labels, buttons, menus, headings, hints, and resourcestrings in one saved batch when possible. This reduces repeated scans and provider calls, but the same procedure works for one change.")
    add_steps(document, [
        "In RAD Studio, make the designer or source changes in the target application.",
        "Choose File > Save All.",
        "Close the running target executable.",
        "Save all target files in RAD Studio and close the running target application before any Wizard final-processing pass. The project may remain open.",
        "Open the project in the Translation Studio.",
        "Open the existing development catalog before scanning.",
        "Run Scan Project.",
        "Confirm new entries are New, revised source text is Source Changed, unchanged completed translations remain preserved, and removed items become Obsolete.",
        "Run Translate Automatically. Confirm only eligible unresolved/new/changed entries are offered to the provider.",
        "Review the new translation and any new terminology/layout proposal.",
        "Save the catalog, run validation, and resolve all errors.",
        "Export Runtime Pack.",
        "Regenerate the Component Kit so its Localization folder contains the final pack.",
        "Build the target after adding $(PROJECTDIR)\\dependencies\\DelphiAppTranslation\\source once through Project Options; then use the Studio's deploy action if the executable folder needs a refreshed Localization\\Languages pack set.",
        "Run the target and retest the changed form in the source and target language.",
    ])

    document.add_heading("17. Adding More Languages", level=1)
    add_paragraphs(document, [
        "Use this procedure when the application already has working localization and you want to add one new target language. Do not recreate the project from its pristine copy, remove existing language packs, reinstall the design package, or place additional localization components. One Wizard run adds one target language while preserving the languages already present.",
    ])
    add_steps(document, [
        "Open the existing approved working project in the Translation Studio. Keep its current Localization folder, language packs, project glossary, language manager, and connected selector intact.",
        "Start the Setup Wizard and select the same project file used for the existing translation. Confirm that the detected Application ID exactly matches the ApplicationId already assigned to the language manager.",
        "Choose the same source language used by the existing packs. In Target language, choose only the new language being added.",
        "Leave Workflow set to Automatic (Recommended). Because this target language does not yet have a development catalog, the Wizard should report that it will create a new translation. Existing catalogs and runtime packs for other languages remain unchanged.",
        "Confirm the translation provider and test its connection. Scan the project and review the scan total before continuing.",
        "Complete the component-information and authorization steps, then click Begin Final Processing.",
        "When Localization Review opens automatically, review terminology and layout proposals for the new language. Save the desired decisions and click Close.",
        "Wait while the same Wizard pass resumes automatically. Confirm that final validation passes, the new runtime JSON pack is exported, the component kit and project-local dependency folder are refreshed, direct deployment completes, and the footer reports successful completion.",
        "Click Finish. A separate Studio validation/export or second Wizard pass is not required when this single-pass completion succeeds. Use the main Studio only if a manual correction or detailed inspection is needed.",
        "Do not reinstall the design package and do not place another manager or selector. The components already in the target application serve every language pack that shares the same Application ID.",
        "Build the required Win32 and Win64 configurations. Confirm the new language JSON file appears under Localization\\Languages beside each executable and in every available application destination configured on the Wizard Deployment page.",
        "Run the application. Open the existing language selector and confirm the new language appears alongside the source language and all earlier target languages.",
        "Select the new language and test its translation, layout, secondary forms, source-language restoration, repeated switching, and saved preference independently. Then verify that previously installed languages still work.",
    ])

    document.add_heading("18. Final Completion Checklist", level=1)
    add_checklist(document, [
        "Disposable test folder was copied from pristine and the pristine/original folders remained untouched.",
        "The single Wizard pass translated unresolved entries and opened Localization Review before final export.",
        "Localization audit, terminology suggestions, glossary, and layout proposals were reviewed.",
        "The same Wizard pass resumed automatically and applied saved review decisions.",
        "If an optional Main Studio correction was made, the matching catalog was saved, validated with zero errors, exported, and followed by a component-kit refresh.",
        "If no optional Studio correction was needed, the single-pass Wizard's final runtime pack and refreshed component kit were used unchanged.",
        "Correct Win32 Release design BPL was installed through Component > Install Packages > Add.",
        "One manager and one connected visible selector were placed and saved on the primary form.",
        "ApplicationId, LanguagesFolder, SourceLanguage, and LanguageManager properties were verified in Object Inspector.",
        "Project Options contains $(PROJECTDIR)\\dependencies\\DelphiAppTranslation\\source exactly once for all required configurations/platforms; the target DPROJ remains unchanged by the Studio.",
        "Win32 and Win64 builds completed.",
        "Each executable folder contains Localization\\Languages with en-US.json and all target packs.",
        "First launch used the source language under the controlled test condition.",
        "Immediate switching, secondary forms, source restoration, repeated switching, and restart preference passed.",
        "Accepted safe layout rules worked without cumulative growth or unacceptable overlap.",
        "No missing/random translation or runaway dynamic-text defect remains unexplained.",
    ])

    document.add_heading("19. What to Send Back After Testing", level=1)
    add_bullets(document, [
        "Disposable test-folder path and selected .dproj path.",
        "Application ID, framework, source language, target language, and provider.",
        "Single-pass scan total, unresolved count, and automatic-resume result.",
        "Wizard completion log or the exact STOPPED message.",
        "Validation error/warning/information counts.",
        "Number of glossary terms approved and layout proposals accepted/manual/rejected.",
        "Whether the final runtime JSON has schemaVersion 3 and a layout array.",
        "Win32/Win64 build results and exact executable/output folders.",
        "Contents of each executable's Localization\\Languages folder.",
        "First-launch language, selector choices, restart language, and ten-cycle switching result.",
        "Screenshots and exact control/form names for every untranslated, mistranslated, truncated, overlapping, or runaway item.",
    ])

    document.add_heading("20. Important Stop Conditions", level=1)
    add_bullets(document, [
        "Stop if the selected project path is the pristine or original application.",
        "Stop if Wizard final processing reports STOPPED or validation reports blocking errors.",
        "Stop if the design package fails to load; do not remove unrelated Delphi packages.",
        "Stop if the Wizard Search Path is absent rather than masking the defect with an unrelated global path.",
        "Stop if runtime packs beside the executable do not match the Application ID or selected language.",
        "Stop a target executable that generates unbounded repeated text; preserve evidence before restarting.",
        "Do not distribute or publish this pre-release test output as a production localization solution.",
    ])

    finish_document(document, DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_document())
